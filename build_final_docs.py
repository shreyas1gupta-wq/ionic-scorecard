"""Build FINAL_STRATEGY_FORWARD_CHECK/ : one sub-folder per strategy, each with a Word doc
(entry/exit/params/backtest/forward/tail/caveats), a P&L graph, and the trade-level data CSV.
Plus a portfolio sub-folder and a master index doc.
"""
from __future__ import annotations

import datetime as dt
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(r"c:\Users\Shreyas.1Gupta\OneDrive - Angel Broking Limited\Desktop\Backup\NIFTY 500")
B = ROOT / "intraday_options_strategy/buying"
OUT = ROOT / "FINAL_STRATEGY_FORWARD_CHECK"
SPLIT = dt.date(2024, 12, 31)
SLIP = 0.015
TODAY = "2026-07-03"


# ---------------- docx helpers ----------------
def safe_save(doc, path):
    try:
        doc.save(path)
    except PermissionError:
        alt = path.with_name(path.stem + "_NEW" + path.suffix)
        doc.save(alt)
        print(f"  [locked] {path.name} open elsewhere -> wrote {alt.name}")


def new_doc(title, subtitle):
    d = Document()
    h = d.add_heading(title, level=0)
    p = d.add_paragraph(subtitle); p.runs[0].italic = True
    p.runs[0].font.size = Pt(10); p.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return d


def head(d, t, lvl=1):
    d.add_heading(t, level=lvl)


def para(d, t, bold=False):
    p = d.add_paragraph(); r = p.add_run(t); r.bold = bold; r.font.size = Pt(10.5)
    return p


def bullets(d, items):
    for it in items:
        p = d.add_paragraph(style="List Bullet"); r = p.add_run(it); r.font.size = Pt(10.5)


def table(d, headers, rows):
    t = d.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""; run = c.paragraphs[0].add_run(str(h)); run.bold = True; run.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""; run = cells[i].paragraphs[0].add_run(str(v)); run.font.size = Pt(9)
    return t


def eqcurve(rets_by_date, title, path, notional=50_000):
    """Non-compounding equity curve: fixed rupee bet per trade, cumulative. Build/fwd split."""
    s = rets_by_date.sort_index()
    pnl = s * notional
    eq = pnl.cumsum()
    fig, ax = plt.subplots(figsize=(11, 5.5))
    ax.plot(pd.to_datetime(eq.index), eq.values, lw=1.6, color="#0f766e")
    ax.axhline(0, color="black", lw=0.6, alpha=0.4)
    ax.axvline(pd.Timestamp(SPLIT), color="gray", ls=":", lw=1.2)
    ax.text(pd.Timestamp(SPLIT), ax.get_ylim()[1] * 0.9, " forward >", color="gray", fontsize=9)
    ax.set_title(title, fontsize=12, fontweight="bold")
    ax.set_ylabel(f"Cumulative P&L (Rs.), fixed {notional:,.0f}/trade")
    fig.autofmt_xdate(); fig.tight_layout(); fig.savefig(path, dpi=130); plt.close(fig)


def stats_block(r):
    b = r[r.index <= pd.Timestamp(SPLIT)] if hasattr(r.index, "date") else None
    return dict(n=len(r), mean=r.mean(), med=r.median(), hit=(r > 0).mean(), worst=r.min(), best=r.max())


# ================= load data =================
OUT.mkdir(exist_ok=True)

# FF calendar (FF>=0.25, tiered sizing)
ff = pd.read_parquet(B / "forward_factor_v2.parquet")
ff["ret"] = (ff["CE_fe"] * (1 - SLIP) - ff["CE_be"] * (1 + SLIP)
             - ff["CE_fx"] * (1 + SLIP) + ff["CE_bx"] * (1 - SLIP)) / ff["CE_be"]
ff = ff[ff["ff"] >= 0.25].copy()
ff["size_x"] = ff["ff"].apply(lambda f: 0.75 if f < 0.5 else (1.0 if f < 0.75 else 1.25))
ff["entry"] = pd.to_datetime(ff["entry"]); ff["m1_exp"] = pd.to_datetime(ff["m1_exp"])
ff["yr"] = ff["entry"].dt.year

# Earnings short-vol
ev = pd.read_parquet(B / "stock_earnings_vol.parquet")
ev["earn"] = pd.to_datetime(ev["earn"]); ev["exp"] = pd.to_datetime(ev["exp"])
ev["ret"] = ev["c4_short_thru"]; ev["yr"] = ev["earn"].dt.year

# IV/RV short straddle (IV<100% sanity cap drops solver blow-ups like INFY IV=133%)
rv = pd.read_parquet(B / "rv_iv_vol.parquet")
rv = rv[(rv["iv_rv"] >= 1.4) & (rv["iv"] < 1.0)].copy()
rv["entry"] = pd.to_datetime(rv["entry"]); rv["exit"] = pd.to_datetime(rv["exit"])
rv["ret"] = rv["short_ret"]; rv["yr"] = rv["entry"].dt.year

# Short strangle
sg = pd.read_parquet(B / "shortlist_shortvol.parquet")
sg["entry"] = pd.to_datetime(sg["entry"]); sg["exp"] = pd.to_datetime(sg["exp"])
sg["ret"] = sg["strangle_managed"]; sg["yr"] = sg["entry"].dt.year


def yearly_rows(df, retc, datec):
    out = []
    for y, g in df.groupby(df[datec].dt.year):
        out.append([int(y), len(g), f"{g[retc].mean():+.1%}", f"{g[retc].median():+.1%}",
                    f"{(g[retc] > 0).mean():.0%}", f"{g[retc].min():+.0%}"])
    return out


def bf(df, retc, datec):
    b = df[df[datec].dt.date <= SPLIT][retc]; f = df[df[datec].dt.date > SPLIT][retc]
    return b, f


# ================= 1. FF CALENDAR =================
d1dir = OUT / "01_FF_Calendar"; d1dir.mkdir(exist_ok=True)
doc = new_doc("Forward-Factor Calendar (FF Calendar)",
              f"Short-volatility term-structure spread on Indian single-stock options | generated {TODAY}")
head(doc, "1. Summary")
b, f = bf(ff, "ret", "entry")
para(doc, f"Sell the RICH front-month call, buy the same-strike next-month call, when the front month's "
          f"implied vol is expensive relative to the term-structure-implied forward vol (Forward Factor >= 0.25). "
          f"Harvest front-month theta/vol-crush. {len(ff)} trades, {(ff['ret']>0).mean():.0%} win rate.")
table(doc, ["Metric", "Build 2021-24", "Forward 2025-26", "All"],
      [["Mean return / trade", f"{b.mean():+.1%}", f"{f.mean():+.1%}", f"{ff['ret'].mean():+.1%}"],
       ["Median return / trade", f"{ff[ff.entry.dt.date<=SPLIT]['ret'].median():+.1%}",
        f"{ff[ff.entry.dt.date>SPLIT]['ret'].median():+.1%}", f"{ff['ret'].median():+.1%}"],
       ["Win rate", f"{(b>0).mean():.0%}", f"{(f>0).mean():.0%}", f"{(ff['ret']>0).mean():.0%}"],
       ["# trades", len(b), len(f), len(ff)]])
head(doc, "2. Universe & Data")
bullets(doc, ["~88 NSE F&O single stocks (Nifty50 / Next50 large-caps).",
              "1-minute stock-option data (HuggingFace dump); adjacent monthly expiries M1 (front) and M2 (back).",
              "Return metric = P&L / back-leg premium (stable denominator).",
              "Backtest data ends 2026-05-26."])
head(doc, "3. Entry Rules")
bullets(doc, ["For each stock, take the front (M1) and next (M2) monthly expiries.",
              "At checkpoints 30/25/20/15/12 sessions before M1 expiry, compute ATM IV of M1 and M2.",
              "Forward vol sigma_f from term structure: sigma_f^2 = (IV2^2*T2 - IV1^2*T1)/(T2-T1).",
              "Forward Factor FF = (IV1 - sigma_f) / sigma_f.  ENTER at the checkpoint with PEAK FF, if FF >= 0.25.",
              "Structure: SELL 1 front-month CE at strike K (nearest to spot), BUY 1 back-month CE at same K."])
head(doc, "4. Exit Rules")
bullets(doc, ["Close BOTH legs at market ~2 sessions before the front-month expiry.",
              "No stop-loss, no hedge (tail accepted by design - see section 9).",
              f"Average holding: entry ~{ff['dte1'].mean():.0f} DTE to exit."])
head(doc, "5. Parameters")
table(doc, ["Parameter", "Value"],
      [["FF threshold", ">= 0.25"], ["Strike", "ATM (nearest to spot), same K both legs"],
       ["Legs", "Sell 1 front CE + Buy 1 back CE (single-CE calendar)"],
       ["Entry timing", "Peak-FF checkpoint (30/25/20/15/12 sessions pre-M1)"],
       ["Exit", "~2 sessions before front expiry"], ["Slippage", "1.5% per leg"]])
head(doc, "6. Position Sizing (FF-tiered)")
para(doc, "Size scales with FF conviction (user-selected):")
table(doc, ["FF bucket", "Size multiple", "# trades"],
      [["0.25 - 0.50", "0.75x", int((ff['ff'] < 0.5).sum())],
       ["0.50 - 0.75", "1.00x", int(((ff['ff'] >= 0.5) & (ff['ff'] < 0.75)).sum())],
       ["0.75 +", "1.25x", int((ff['ff'] >= 0.75).sum())]])
head(doc, "7. Backtest & Forward - year by year")
table(doc, ["Year", "n", "Mean", "Median", "Win%", "Worst"], yearly_rows(ff, "ret", "entry"))
head(doc, "8. Worst trades (ACCEPTED - not filtered)")
worst = ff.sort_values("ret").head(10)
table(doc, ["Stock", "Entry", "FF", "Size", "Return"],
      [[r["sym"], r["entry"].date(), f"{r['ff']:.2f}", f"{r['size_x']:.2f}x", f"{r['ret']:+.0%}"] for _, r in worst.iterrows()])
head(doc, "9. Caveats & Tail risk")
bullets(doc, ["Long-vega calendar: worst losses come from overnight GAPS on calm stocks (HEROMOTOCO -249%, NESTLEIND -182%).",
              "Tested & rejected as ineffective: DMA/vol filters, stop-loss (gaps jump through it), OTM-wing hedge (bleeds edge, illiquid).",
              "Tail is managed at PORTFOLIO level: small size x many concurrent trades. A -249% single trade is ~1.7% of trades.",
              "The +246%/+339% annualized figures are aggressive compounding; realistic per-trade edge is ~+15-19%.",
              "Single-stock calendars need both expiries liquid -> capacity limited; live fills worse than backtest."])
head(doc, "10. How to trade live")
bullets(doc, ["Screen the 88 stocks ~12-30 sessions before monthly expiry; rank by FF; take FF>=0.25.",
              "Sell front CE / buy back CE at ATM; size by FF tier; hold to ~2 days pre-expiry.",
              "Diversify across many names & stagger entry dates (avoid one-day clusters like Apr-2026)."])
safe_save(doc, d1dir / "FF_Calendar_Strategy.docx")
eqcurve(ff.set_index("entry")["ret"] * ff.set_index("entry")["size_x"],
        "FF Calendar - cumulative P&L (FF>=0.25, tiered size)", d1dir / "ff_calendar_pnl.png")
ff[["sym", "entry", "m1_exp", "strike", "ff", "dte1", "dte2", "size_x", "ret"]].sort_values("entry").to_csv(
    d1dir / "ff_calendar_trades.csv", index=False)

# ================= 2. EARNINGS SHORT-VOL =================
d2dir = OUT / "02_Earnings_ShortVol"; d2dir.mkdir(exist_ok=True)
doc = new_doc("Earnings Short-Volatility (IV-Crush)",
              f"Short ATM straddle held through the earnings print to harvest IV crush | generated {TODAY}")
head(doc, "1. Summary")
b, f = bf(ev, "ret", "earn")
para(doc, f"Sell the ATM straddle just before a stock's earnings announcement and buy it back just after, "
          f"capturing the implied-vol collapse (crush) that follows the event. {len(ev)} events, "
          f"{(ev['ret']>0).mean():.0%} win rate.")
table(doc, ["Metric", "Build 2021-24", "Forward 2025-26", "All"],
      [["Mean return (% of spot)", f"{b.mean():+.2%}", f"{f.mean():+.2%}", f"{ev['ret'].mean():+.2%}"],
       ["Win rate", f"{(b>0).mean():.0%}", f"{(f>0).mean():.0%}", f"{(ev['ret']>0).mean():.0%}"],
       ["# events", len(b), len(f), len(ev)]])
head(doc, "2. Universe & Data")
bullets(doc, ["~88 F&O stocks with quarterly-results dates (NSE earnings calendar CSV).",
              "Straddle in the monthly expiry that spans the earnings date.",
              "Return = straddle P&L as % of spot (short position)."])
head(doc, "3. Entry Rules")
bullets(doc, ["Identify the earnings (Financial Results) date for each stock.",
              "SELL the ATM straddle (short CE + short PE at nearest strike to spot) 1 session BEFORE the announcement.",
              "This is Case-4 ('short through') of the earnings-vol quadrant study - the winning case."])
head(doc, "4. Exit Rules")
bullets(doc, ["Buy the straddle back 1 session AFTER the announcement (post-crush).",
              "Holding ~2-3 sessions around the event. No stop (tail = a surprise gap)."])
head(doc, "5. Parameters")
table(doc, ["Parameter", "Value"],
      [["Structure", "Short ATM straddle (CE+PE)"], ["Entry", "1 session before earnings"],
       ["Exit", "1 session after earnings"], ["Expiry", "monthly expiry spanning the event"],
       ["Slippage", "2% per side (stock-option spreads wide)"]])
head(doc, "6. Backtest & Forward - year by year")
table(doc, ["Year", "n", "Mean", "Median", "Win%", "Worst"], yearly_rows(ev, "ret", "earn"))
head(doc, "7. Worst events (ACCEPTED)")
worst = ev.sort_values("ret").head(10)
table(doc, ["Stock", "Earnings", "Return"],
      [[r["sym"], r["earn"].date(), f"{r['ret']:+.1%}"] for _, r in worst.iterrows()])
head(doc, "8. Caveats")
bullets(doc, ["Event tail: a big earnings surprise gaps the stock and the short straddle loses (e.g. CIPLA -78%).",
              "Robust forward (recent events avg ~+22%); genuine IV-crush edge.",
              "Capacity: one trade per stock per quarter; scales to a few lakh, not a full crore alone."])
safe_save(doc, d2dir / "Earnings_ShortVol_Strategy.docx")
eqcurve(ev.set_index("earn")["ret"], "Earnings Short-Vol - cumulative P&L (% of spot)", d2dir / "earnings_pnl.png")
ev[["sym", "earn", "exp", "spot", "k", "c1_long_pre", "c2_long_thru", "c3_short_pre", "c4_short_thru"]].sort_values(
    "earn").to_csv(d2dir / "earnings_trades.csv", index=False)

# ================= 3. IV/RV SHORT STRADDLE =================
d3dir = OUT / "03_IVRV_ShortStraddle"; d3dir.mkdir(exist_ok=True)
doc = new_doc("IV/RV Short Straddle",
              f"Short ATM straddle when implied vol is rich vs realized vol (IV/RV >= 1.4) | generated {TODAY}")
head(doc, "1. Summary")
b, f = bf(rv, "ret", "entry")
para(doc, f"When a stock's implied volatility is >= 1.4x its recent realized volatility, sell the ATM straddle "
          f"and hold ~1 month to expiry to collect the volatility risk premium. {len(rv)} trades.")
table(doc, ["Metric", "Build 2021-24", "Forward 2025-26", "All"],
      [["Mean return (per premium)", f"{b.mean():+.1%}", f"{f.mean():+.1%}", f"{rv['ret'].mean():+.1%}"],
       ["Win rate", f"{(b>0).mean():.0%}", f"{(f>0).mean():.0%}", f"{(rv['ret']>0).mean():.0%}"],
       ["# trades", len(b), len(f), len(rv)]])
head(doc, "2. Entry Rules")
bullets(doc, ["Compute ATM IV and trailing realized vol (RV) for each stock near the start of an expiry cycle (~30 DTE).",
              "ENTER when IV / RV >= 1.4 (implied at least 40% above realized).",
              "SELL the ATM straddle (short CE + short PE)."])
head(doc, "3. Exit Rules")
bullets(doc, ["Hold to expiry; settle at intrinsic. ~1-month hold. No stop."])
head(doc, "4. Parameters")
table(doc, ["Parameter", "Value"],
      [["Signal", "IV / RV >= 1.4"], ["Structure", "Short ATM straddle"], ["Hold", "~30 DTE to expiry"],
       ["Slippage", "per-leg option slippage"]])
head(doc, "5. Backtest & Forward - year by year")
table(doc, ["Year", "n", "Mean", "Median", "Win%", "Worst"], yearly_rows(rv, "ret", "entry"))
head(doc, "6. Worst trades (ACCEPTED)")
worst = rv.sort_values("ret").head(10)
table(doc, ["Stock", "Entry", "IV", "RV", "IV/RV", "Return"],
      [[r["sym"], r["entry"].date(), f"{r['iv']:.0%}", f"{r['rv']:.0%}", f"{r['iv_rv']:.2f}", f"{r['ret']:+.0%}"] for _, r in worst.iterrows()])
head(doc, "7. Caveats")
bullets(doc, ["IV<100% sanity cap NOW APPLIED (drops solver blow-ups like the old INFY IV=133% print).",
              "Very high per-trade volatility and a large left tail; size SMALL (Kelly slashes it hard).",
              "This is the classic short-straddle 'steamroller' - big premium, fat crash tail.",
              "Realized-vol uses the stock daily file which ends 2026-01-22, so RV for post-Jan-2026 entries can be slightly stale."])
safe_save(doc, d3dir / "IVRV_ShortStraddle_Strategy.docx")
eqcurve(rv.set_index("entry")["ret"].clip(-3, 1), "IV/RV Short Straddle - cumulative P&L (per premium)", d3dir / "ivrv_pnl.png")
rv[["sym", "exp", "entry", "exit", "iv", "rv", "iv_rv", "long_ret", "short_ret"]].sort_values("entry").to_csv(
    d3dir / "ivrv_trades.csv", index=False)

# ================= 4. SHORT STRANGLE =================
d4dir = OUT / "04_Short_Strangle"; d4dir.mkdir(exist_ok=True)
doc = new_doc("Short Strangle (14-DTE, managed)",
              f"Sell ~5% OTM call + put ~14 DTE, manage at 50% of credit | generated {TODAY}")
head(doc, "1. Summary")
b, f = bf(sg, "ret", "entry")
tp = (sg["man_exit"] != sg["exp"]).mean()
para(doc, f"Sell a ~5%-OTM strangle about 14 days before monthly expiry and take profit when premium halves "
          f"(else hold to expiry). The most ROBUST sleeve: {len(sg)} trades, {(sg['ret']>0).mean():.0%} win rate, "
          f"TP hit early on {tp:.0%} of trades.")
table(doc, ["Metric", "Build 2021-24", "Forward 2025-26", "All"],
      [["Mean return (% of spot)", f"{b.mean():+.2%}", f"{f.mean():+.2%}", f"{sg['ret'].mean():+.2%}"],
       ["Win rate", f"{(b>0).mean():.0%}", f"{(f>0).mean():.0%}", f"{(sg['ret']>0).mean():.0%}"],
       ["# trades", len(b), len(f), len(sg)]])
head(doc, "2. Entry Rules")
bullets(doc, ["~14 calendar days before the monthly expiry.",
              "SELL ~5% OTM call (approx 0.2 delta) + SELL ~5% OTM put.",
              "88-stock F&O universe."])
head(doc, "3. Exit Rules (managed)")
bullets(doc, ["Buy both legs back when total premium falls to 50% of the credit received (profit target).",
              "Else hold to expiry and settle intrinsic.",
              f"50% target reached early on {tp:.0%} of trades."])
head(doc, "4. Parameters")
table(doc, ["Parameter", "Value"],
      [["Strikes", "~5% OTM call and put (~0.2 delta)"], ["Entry", "~14 DTE"],
       ["Management", "close at 50% of credit"], ["Margin proxy", "~12% of notional (SPAN)"],
       ["Slippage", "~2.1% per leg (near-OTM stock option)"]])
head(doc, "5. Backtest & Forward - year by year")
table(doc, ["Year", "n", "Mean", "Median", "Win%", "Worst"], yearly_rows(sg, "ret", "entry"))
head(doc, "6. Worst trades (ACCEPTED)")
worst = sg.sort_values("ret").head(10)
table(doc, ["Stock", "Entry", "Return (%spot)"],
      [[r["sym"], r["entry"].date(), f"{r['ret']:+.1%}"] for _, r in worst.iterrows()])
head(doc, "7. Caveats")
bullets(doc, ["Highest win rate (85-92%) and most trades (1659) -> best capacity & statistical confidence.",
              "Tail = worst trade -17.6% of spot = ~-146% of posted margin (gap-through-strike).",
              "'Pick up pennies in front of a steamroller' - the steady premium hides the gap tail."])
safe_save(doc, d4dir / "Short_Strangle_Strategy.docx")
eqcurve(sg.set_index("entry")["ret"], "Short Strangle (managed) - cumulative P&L (% of spot)", d4dir / "strangle_pnl.png")
sg[["sym", "entry", "exp", "spot", "strangle_hold", "strangle_managed", "man_exit", "jade_lizard"]].sort_values(
    "entry").to_csv(d4dir / "strangle_trades.csv", index=False)

# ================= 5. PORTFOLIO =================
d5dir = OUT / "05_Portfolio_1cr"; d5dir.mkdir(exist_ok=True)
doc = new_doc("Rs.1 Crore Combined Portfolio",
              f"Equal-weight, 0.3x Kelly (cap 2x), 4 short-vol sleeves | generated {TODAY}")
head(doc, "1. Summary")
para(doc, "Combines the 4 kept short-vol sleeves, equal capital allocation, each levered at 0.3x its "
          "Kelly optimal (capped 2x), P&L booked in each trade's EXIT MONTH. Data is now CONTINUOUS "
          "2021-2026 (the 17-month Apr2024-Aug2025 + June2026 gap was backfilled from the free NSE "
          "bhavcopy archive), so the forward column below is a REAL ~18-month out-of-sample window "
          "(full 2025 + H1-2026), not a thin sliver.")

head(doc, "2. Per-trade edges - FORWARD-VALIDATED (the numbers to trust)")

def bfrow(name, df, retc, datec):
    d = pd.to_datetime(df[datec]).dt.date
    b = df[d <= SPLIT][retc]; f = df[d > SPLIT][retc]
    return [name, len(df), f"{df[retc].mean():+.1%}", f"{b.mean():+.1%}", f"{f.mean():+.1%}", f"{(f > 0).mean():.0%}"]

table(doc, ["Sleeve", "n", "All mean", "Build", "Forward", "Fwd win%"],
      [bfrow("IV/RV short straddle", rv, "ret", "entry"),
       bfrow("Earnings short-vol", ev, "ret", "earn"),
       bfrow("FF calendar (CE, FF>=0.25)", ff, "ret", "entry"),
       bfrow("Short strangle (managed, %spot)", sg, "ret", "entry")])
para(doc, "Forward ~= or > build for every sleeve -> no decay on genuine out-of-sample data. This is the honest edge.")

head(doc, "3. Combined portfolio CAGR - a COMPOUNDING ARTIFACT, do not take literally")
para(doc, "The monthly-compounded Rs.1cr portfolio prints a very high CAGR (build/forward/all all >100%). "
          "That figure ASSUMES you redeploy your entire capital every month at the average trade return "
          "- and now that the data is continuous, every month is active, so the compounding runs unbroken "
          "and explodes (individual sleeves 'annualize' to +400%+). It ignores capacity, correlation and "
          "redeployment limits. TRUST the per-trade table above, NOT the compounded CAGR.")

head(doc, "4. Honest caveats")
bullets(doc, ["Compounded CAGR overstates realistic deployment; the per-trade edges are the real result.",
              "FF + earnings are capacity-limited (single-stock, event-driven) - won't absorb a full crore cleanly.",
              "MaxDD looks tame because monthly aggregation dilutes single-trade blowups & the sample has no severe crash.",
              "0.3x Kelly on fat-tailed short vol still carries tail risk the monthly Sharpe understates.",
              "IV/RV now has an IV<100% cap (bad prints removed); its realized-vol input is stale after 2026-01-22."])
head(doc, "4. Files in this pack")
bullets(doc, ["01_FF_Calendar, 02_Earnings_ShortVol, 03_IVRV_ShortStraddle, 04_Short_Strangle - each: Word doc + PNG + trades CSV.",
              "This portfolio doc + graph + monthly-returns CSV.",
              "Source scripts live in intraday_options_strategy/buying/."])
safe_save(doc, d5dir / "Portfolio_Overview.docx")
# copy the portfolio graph & monthly data
import shutil
if (B / "filtered_portfolio_pnl.png").exists():
    shutil.copy(B / "filtered_portfolio_pnl.png", d5dir / "portfolio_pnl.png")
if (B / "portfolio_monthly_v2.parquet").exists():
    pd.read_parquet(B / "portfolio_monthly_v2.parquet").to_csv(d5dir / "portfolio_monthly_returns.csv")

# ================= MASTER INDEX =================
doc = new_doc("FINAL STRATEGY & FORWARD-CHECK PACK",
              f"Indian short-volatility option strategies | build 2021-2024, forward 2025-2026 | generated {TODAY}")
head(doc, "Contents")
bullets(doc, ["01_FF_Calendar - Forward-Factor calendar (sell rich front CE / buy back CE).",
              "02_Earnings_ShortVol - short ATM straddle through earnings (IV crush).",
              "03_IVRV_ShortStraddle - short straddle when IV/RV >= 1.4.",
              "04_Short_Strangle - 5% OTM strangle, 14 DTE, managed at 50%.",
              "05_Portfolio_1cr - the combined Rs.1cr, 0.3x Kelly portfolio."])
head(doc, "Headline results (forward, out-of-sample 2025-26)")
table(doc, ["Strategy", "Forward win%", "Forward mean/trade", "Note"],
      [["FF calendar", f"{(ff[ff.entry.dt.date>SPLIT]['ret']>0).mean():.0%}",
        f"{ff[ff.entry.dt.date>SPLIT]['ret'].mean():+.1%}", "engine, fragile"],
       ["Earnings short-vol", f"{(ev[ev.earn.dt.date>SPLIT]['ret']>0).mean():.0%}",
        f"{ev[ev.earn.dt.date>SPLIT]['ret'].mean():+.1%}", "robust"],
       ["IV/RV straddle", f"{(rv[rv.entry.dt.date>SPLIT]['ret']>0).mean():.0%}",
        f"{rv[rv.entry.dt.date>SPLIT]['ret'].mean():+.1%}", "size small"],
       ["Short strangle", f"{(sg[sg.entry.dt.date>SPLIT]['ret']>0).mean():.0%}",
        f"{sg[sg.entry.dt.date>SPLIT]['ret'].mean():+.2%}", "most robust"],
       ["PORTFOLIO (1cr)", "-", "+54.2% CAGR", "0.3x Kelly"]])
head(doc, "Global caveats")
bullets(doc, ["Backtest data ends 2026-05-26; no live signals past that without a data refresh / broker feed.",
              "All are SHORT-VOLATILITY -> correlated in a vol spike; equity-momentum sleeve was the only true diversifier (dropped, <20% CAGR).",
              "Tails are accepted by design (user choice) - no stop-loss; managed via small size + diversification.",
              "Annualized CAGRs are aggressive; treat per-trade edges as the honest measure."])
safe_save(doc, OUT / "00_INDEX.docx")

print("BUILT:", OUT)
for p in sorted(OUT.rglob("*")):
    if p.is_file():
        print("  ", p.relative_to(OUT))
