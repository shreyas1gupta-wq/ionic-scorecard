# -*- coding: utf-8 -*-
"""Addendum report: winsorization + complete-market median PE + small-cap."""
import os, json
import numpy as np, pandas as pd
import matplotlib; matplotlib.use("Agg"); import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT=os.path.dirname(os.path.abspath(__file__))
RES=os.path.join(OUT,"results"); CH=os.path.join(OUT,"charts"); os.makedirs(CH,exist_ok=True)
def rd(n): return pd.read_csv(os.path.join(RES,n))
info=json.load(open(os.path.join(RES,"regime_info_v2.json")))
cmp=rd("valuation_breadth_compare.csv"); cmp[cmp.columns[0]]=pd.to_datetime(cmp[cmp.columns[0]]); cmp=cmp.set_index(cmp.columns[0])
BLUE="#1f3b6f"; RED="#b3202c"; GRN="#1a7a3c"; ORA="#d98a00"; GREY="#888"
plt.rcParams.update({"font.size":10,"axes.grid":True,"grid.alpha":0.25,"figure.dpi":130})

# chart 1: breadth valuation over time
fig,ax=plt.subplots(figsize=(8,3.4))
ax.plot(cmp.index,cmp["median_stocklevel"],color=RED,lw=1.6,label="TRUE median stock PE (~1,100 names)")
ax.plot(cmp.index,cmp["nifty50_capwt"],color=BLUE,lw=1.2,label="NIFTY 50 cap-weighted PE")
ax.plot(cmp.index,cmp["nifty500_capwt"],color=GRN,lw=1.0,ls="--",label="NIFTY 500 cap-weighted PE")
ax.set_title("Large-cap bias: the typical stock vs the headline index (India)")
ax.set_ylabel("Trailing P/E"); ax.legend(fontsize=8); ax.set_ylim(5,45)
fig.tight_layout(); fig.savefig(os.path.join(CH,"breadth_valuation.png")); plt.close()

# chart 2: winsorized fwd returns by regime, broad + smallcap
b=rd("regime_stats_wins_INDIA_BROAD.csv"); s=rd("regime_stats_wins_INDIA_SMALLCAP.csv")
fig,axs=plt.subplots(1,2,figsize=(9,3.4))
for ax,d,ttl in [(axs[0],b,"Broad market (median-PE regime)"),(axs[1],s,"Small-cap (Nifty Smallcap 250)")]:
    d=d[d.regime!="ALL"].set_index("regime").reindex(["CHEAP","FAIR","RICH"])
    x=np.arange(3); w=0.38
    ax.bar(x-w/2,d["fwd12m_mean_w"]*100,w,color=BLUE,label="Fwd 12m (winsor) %")
    ax.bar(x+w/2,d["ann_vol_w"]*100,w,color=RED,alpha=.7,label="Ann vol (winsor) %")
    ax.set_xticks(x); ax.set_xticklabels(["CHEAP","FAIR","RICH"]); ax.set_title(ttl,fontsize=9); ax.axhline(0,color="k",lw=.6); ax.legend(fontsize=7)
fig.suptitle("Winsorized forward return vs vol — broad & small-cap (RICH = weak forward, as it should)",y=1.02)
fig.tight_layout(); fig.savefig(os.path.join(CH,"regime_fwd_winsor.png"),bbox_inches="tight"); plt.close()

# chart 3: winsor effect on 'worst'
u=rd("regime_stats_wins_US.csv"); u=u[u.regime!="ALL"].set_index("regime").reindex(["CHEAP","FAIR","RICH"])
fig,ax=plt.subplots(figsize=(6.4,3.2)); x=np.arange(3); w=0.38
ax.bar(x-w/2,u["fwd12m_worst_raw"]*100,w,color=GREY,label="Raw worst (min)")
ax.bar(x+w/2,u["fwd12m_worst_w"]*100,w,color=BLUE,label="Winsorized worst (2.5%)")
ax.set_xticks(x); ax.set_xticklabels(["CHEAP","FAIR","RICH"]); ax.set_ylabel("Fwd 12m worst %")
ax.set_title("U.S.: winsorization removes single-obs bias in the tail estimate"); ax.legend(fontsize=8)
fig.tight_layout(); fig.savefig(os.path.join(CH,"winsor_effect.png")); plt.close()
print("charts v2 done")

# ---------------- DOC ----------------
doc=Document(); st=doc.styles["Normal"]; st.font.name="Calibri"; st.font.size=Pt(10.5)
def H(t,l=1):
    p=doc.add_heading(t,l)
    for r in p.runs: r.font.color.rgb=RGBColor(0x1f,0x3b,0x6f)
def P(t,b=False,it=False,sz=10.5,bullet=False):
    p=doc.add_paragraph(style="List Bullet" if bullet else None); r=p.add_run(t); r.bold=b; r.italic=it; r.font.size=Pt(sz); return p
def table(df,cols,headers=None,pct=None):
    pct=pct or []
    t=doc.add_table(rows=1,cols=len(cols)); t.style="Light Grid Accent 1"
    for i,c in enumerate(cols):
        t.rows[0].cells[i].text=(headers[i] if headers else c)
        for para in t.rows[0].cells[i].paragraphs:
            for run in para.runs: run.bold=True; run.font.size=Pt(8.5)
    for _,row in df[cols].iterrows():
        cs=t.add_row().cells
        for i,c in enumerate(cols):
            v=row[c]
            s=(f"{v*100:.1f}%" if (c in pct and isinstance(v,(int,float,np.floating))) else
               (f"{v:.2f}" if isinstance(v,(int,float,np.floating)) else str(v)))
            cs[i].text=s
            for para in cs[i].paragraphs:
                for run in para.runs: run.font.size=Pt(8.3)

tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=tp.add_run("Addendum — Bias Controls"); r.bold=True; r.font.size=Pt(18); r.font.color.rgb=RGBColor(0x1f,0x3b,0x6f)
sp=doc.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=sp.add_run("Winsorization · complete-market MEDIAN P/E · small-cap-only · (companion to HEDGING_ANALYSIS_REPORT.docx) · 2026-07-08"); r.italic=True; r.font.size=Pt(9.5)

H("A. Winsorization — removing single-observation bias",1)
P("Extreme single months (the 1931 Depression window that produced a −107% forward-12m log return, COVID) were biasing the raw min/max. All descriptive stats are re-computed on returns winsorized at the 2.5/97.5 percentiles. Point estimates barely move (medians were already robust); the extreme 'worst' compresses to a defensible number. The TRUE tail is still reported via CVaR and a raw-worst column — a hedging study must never erase the tail it exists to cover.",sz=10)
doc.add_picture(os.path.join(CH,"winsor_effect.png"),width=Inches(4.8))
P("U.S. S&P 500 by CAPE regime — winsorized (raw worst shown for contrast):",b=True,sz=9.5)
table(u.reset_index(),["regime","median_ann_ret","ann_vol_w","fwd12m_mean_w","fwd12m_worst_w","fwd12m_worst_raw"],
      headers=["Regime","Median ann ret","Ann vol (w)","Fwd12m mean (w)","Fwd12m worst (w)","Fwd12m worst RAW"],
      pct=["median_ann_ret","ann_vol_w","fwd12m_mean_w","fwd12m_worst_w","fwd12m_worst_raw"])
P("Effect: the FAIR-regime worst drops from a −107% single-obs artifact to −35%; RICH from −56% to −46%. Rankings and directional conclusions in the main report are UNCHANGED — they were driven by CVaR/Sortino/drawdown of full equity curves, not by lone extremes.",it=True,sz=9)

H("B. Complete-market MEDIAN P/E — stripping out the large-cap bias",1)
P(f"The headline NIFTY 50 P/E is cap-weighted, so a handful of mega-caps dominate it. I built a TRUE cross-sectional MEDIAN trailing P/E across ~1,100 stocks per month (point-in-time annual EPS ÷ month-end price), 2016–2026. Result: the median stock now trades at {info['median_pe_latest']:.1f}× (as of {info['median_pe_date']}) versus the cap-weighted NIFTY 50 at ~21× — the typical Indian stock is MORE expensive than the index implies.",sz=10)
doc.add_picture(os.path.join(CH,"breadth_valuation.png"),width=Inches(6.4))
P("Valuation by lens — current vs full-sample mean:",b=True,sz=9.5)
comp=pd.DataFrame({"lens":["Median stock (true)","NIFTY 50 cap-wt","NIFTY 500 cap-wt","Total Market cap-wt","Smallcap 250","Microcap 250"],
    "now":[cmp["median_stocklevel"].dropna().iloc[-1],cmp["nifty50_capwt"].dropna().iloc[-1],cmp["nifty500_capwt"].dropna().iloc[-1],
           cmp["totalmkt_capwt"].dropna().iloc[-1],cmp["smallcap250"].dropna().iloc[-1],cmp["microcap250"].dropna().iloc[-1]]})
table(comp,["lens","now"],headers=["Valuation lens","P/E now"])
P(f"REGIME FLIP: on the median-PE signal the broad market is currently RICH (median-PE {info['broad_medpe_now']:.1f} vs the 75th-pctile RICH line {info['broad_medpe_q75']:.1f}) — the OPPOSITE of the cap-weighted NIFTY 50, which reads CHEAP. This revises the main report's 'India is cheap, stay unhedged' call: the LARGE-CAPS are cheap, but the typical/median stock is in its expensive quartile.",b=True,sz=9.7)
bwins=rd("regime_stats_wins_INDIA_BROAD.csv")
P("Broad market by median-PE regime (winsorized):",b=True,sz=9.5)
table(bwins,["regime","n_months","median_ann_ret","ann_vol_w","fwd12m_mean_w","fwd12m_worst_w","pct_neg_month"],
      headers=["Regime","N mo","Median ann ret","Ann vol(w)","Fwd12m(w)","Fwd12m worst(w)","% neg mo"],
      pct=["median_ann_ret","ann_vol_w","fwd12m_mean_w","fwd12m_worst_w","pct_neg_month"])
P("Critically, measured properly the broad-market RICH regime shows WEAK forward returns (+3.6%) — the same overvaluation asymmetry as the U.S., which the cap-weighted PB/PE regime in the main report had masked (it made India's RICH look benign). Best broad hedge in RICH: quarterly ATM put / put-spread (cuts maxDD −11.5%→−4% to −10%). Best play: same convex logic as large-cap (bear/backspread); premium-selling ratios again win on average but are short the tail.",it=True,sz=9)

H("C. Small-cap-only (Nifty Smallcap 250) — the segment the indices miss",1)
swins=rd("regime_stats_wins_INDIA_SMALLCAP.csv")
P("Small-cap by own-P/E regime (winsorized) — note the far higher vol and deeper tails vs large-cap:",b=True,sz=9.5)
table(swins,["regime","n_months","median_ann_ret","ann_vol_w","fwd12m_mean_w","fwd12m_worst_w","fwd12m_worst_raw","pct_neg_month"],
      headers=["Regime","N mo","Median ann ret","Ann vol(w)","Fwd12m(w)","Fwd worst(w)","Fwd worst RAW","% neg mo"],
      pct=["median_ann_ret","ann_vol_w","fwd12m_mean_w","fwd12m_worst_w","fwd12m_worst_raw","pct_neg_month"])
for bl in [
 "Small-caps carry ~20% annual vol vs ~13% for the NIFTY 50, and forward drawdowns of −33% (RICH) to −53% (FAIR, raw) — the large-cap index simply does not show this risk. RICH small-cap has the highest concurrent return (+30%) and the weakest forward (+5%): textbook boom-bust.",
 "Hedging works and is more valuable here: a quarterly collar cuts small-cap RICH maxDD from −29% to ~−17%; the CVaR improvement per unit cost is the best of any segment.",
 "PRACTICAL LIMITATION (honesty gate): liquid exchange-traded options do NOT exist for individual small-caps or for the Smallcap 250 index in India. So the modeled small-cap option structures are ILLUSTRATIVE. The executable small-cap hedge is (a) buy NIFTY/index puts as a proxy (imperfect — beta & basis risk), (b) short index futures, or (c) simply cut small-cap exposure when the median-PE signal is RICH. This is the actionable takeaway, not a bespoke small-cap option trade.",
 "Small-cap play expectancies (ratios topping the table) rest on only ~2–3 non-overlapping annual entries in the RICH regime — treat as indicative, not certified.",
]:
    P(bl,bullet=True,sz=9.5)
doc.add_picture(os.path.join(CH,"regime_fwd_winsor.png"),width=Inches(6.4))

H("D. Revised current-scenario synthesis (three lenses)",1)
rev=pd.DataFrame([
 dict(Lens="NIFTY 50 (large-cap, cap-wt)",Signal="P/E 21 → CHEAP",Read="Mega-caps genuinely not expensive",Action="Stay long, minimal hedge"),
 dict(Lens="Median stock (breadth)",Signal="Median P/E 25.6 → RICH",Read="Typical stock IS in its expensive quartile",Action="Hedge the broad book (qtrly ATM put/put-spread); trim expensive median names"),
 dict(Lens="Small-cap (Smallcap 250)",Signal="P/E 36 → FAIR (absolute high)",Read="Not bubble (2018 was), but high vol & fat tails",Action="Index-put proxy hedge or cut exposure; no bespoke sc options exist"),
])
table(rev,["Lens","Signal","Read","Action"])
P("Bottom line the bias controls add: 'India is cheap' is a LARGE-CAP statement. Strip the cap-weight bias and the median stock is in its RICH quartile with the same weak-forward / fat-tail asymmetry as the U.S. — so a broad or small-cap Indian book DOES warrant the annual/quarterly protective programme, even while the NIFTY-50 headline looks calm. U.S. remains the most extreme (CAPE 41.8).",it=True,sz=9.5)
P("Files: engine_v2.py, build_median_pe.py; regime_stats_wins_*.csv, valuation_breadth_compare.csv, ranking_*_INDIA_BROAD/SMALLCAP.csv, regime_info_v2.json; india_market_median_pe.parquet.",it=True,sz=8.5)

path=os.path.join(OUT,"HEDGING_ANALYSIS_ADDENDUM_v2.docx"); doc.save(path); print("SAVED",path)
