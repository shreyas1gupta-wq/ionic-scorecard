# -*- coding: utf-8 -*-
import os, json
import numpy as np, pandas as pd
import matplotlib; matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT=os.path.dirname(os.path.abspath(__file__))
RES=os.path.join(OUT,"results"); CH=os.path.join(OUT,"charts"); os.makedirs(CH,exist_ok=True)
def rd(n): return pd.read_csv(os.path.join(RES,n))
info=json.load(open(os.path.join(RES,"regime_info.json")))
case=json.load(open(os.path.join(RES,"case_studies.json")))
us_m=pd.read_parquet(os.path.join(RES,"us_monthly.parquet"))
rs_us=rd("regime_stats_US.csv"); rs_in=rd("regime_stats_INDIA.csv"); rs_inpe=rd("regime_stats_INDIA_PE.csv")
hh_us=rd("ranking_hedge_US.csv"); hh_in=rd("ranking_hedge_INDIA.csv")
pp_us=rd("ranking_play_US.csv"); pp_in=rd("ranking_play_INDIA.csv")
cov_in=rd("case_covid_INDIA.csv"); cov_us=rd("case_covid_US.csv")

BLUE="#1f3b6f"; RED="#b3202c"; GRN="#1a7a3c"; GREY="#888"
plt.rcParams.update({"font.size":10,"axes.grid":True,"grid.alpha":0.25,"figure.dpi":130})

# 1) US CAPE regimes
fig,ax=plt.subplots(figsize=(8,3.4))
c=us_m[us_m.index>="1900-01-01"]
ax.plot(c.index,c["cape"],color=BLUE,lw=1.1)
ax.axhline(info["us_cape_q25"],color=GRN,ls="--",lw=1,label=f"25th pct = {info['us_cape_q25']:.1f} (CHEAP below)")
ax.axhline(info["us_cape_q75"],color=RED,ls="--",lw=1,label=f"75th pct = {info['us_cape_q75']:.1f} (RICH above)")
ax.scatter([c.index[-1]],[info["us_cape_now"]],color=RED,zorder=5)
ax.annotate(f"NOW {info['us_cape_now']:.1f}",(c.index[-1],info["us_cape_now"]),xytext=(-70,-2),
            textcoords="offset points",color=RED,fontweight="bold")
ax.set_title("U.S. Shiller CAPE with 25/50/25 valuation regimes (1900–2026)")
ax.set_ylabel("CAPE"); ax.legend(fontsize=8,loc="upper left")
fig.tight_layout(); fig.savefig(os.path.join(CH,"us_cape_regimes.png")); plt.close()

# 2) regime forward returns + vol
fig,axs=plt.subplots(1,2,figsize=(9,3.4))
for ax,rs,ttl,pc in [(axs[0],rs_us,"U.S. S&P 500 (CAPE regimes)","fwd12m_mean"),
                     (axs[1],rs_in,"India NIFTY 50 (PB regimes)","fwd12m_mean")]:
    d=rs[rs.regime!="ALL"].set_index("regime").reindex(["CHEAP","FAIR","RICH"])
    x=np.arange(3); w=0.38
    ax.bar(x-w/2,d["fwd12m_mean"]*100,w,color=BLUE,label="Fwd 12m return %")
    ax.bar(x+w/2,d["ann_vol"]*100,w,color=RED,alpha=.7,label="Annual vol %")
    ax.set_xticks(x); ax.set_xticklabels(["CHEAP","FAIR","RICH"]); ax.set_title(ttl,fontsize=9)
    ax.axhline(0,color="k",lw=.6); ax.legend(fontsize=7)
fig.suptitle("Forward 12-month return vs volatility by valuation regime",y=1.02)
fig.tight_layout(); fig.savefig(os.path.join(CH,"regime_fwd_returns.png"),bbox_inches="tight"); plt.close()

# 3) COVID India payoff (1m@trough)
d=cov_in[cov_in.scenario=="1m@trough"].copy()
key=["H_put_ATM","H_collar_95_105","H_putspread_95_85","H_backspread_1x2_100_90",
     "P_longput_97","P_bearspread_95_85","P_backspread_1x2_97_90","P_ratio_3x2_95_85"]
d=d[d.struct.isin(key)].set_index("struct").reindex(key)
fig,ax=plt.subplots(figsize=(8.4,3.6))
cols=[GRN if v>=0 else RED for v in d["combined"]]
ax.bar(range(len(d)),d["combined"]*100,color=cols)
ax.axhline(case["covid_india_entry"]["trough_dd"]*100,color=GREY,ls="--",lw=1,
           label=f"Unhedged NIFTY at trough = {case['covid_india_entry']['trough_dd']*100:.0f}%")
ax.set_xticks(range(len(d))); ax.set_xticklabels([k.replace("H_","").replace("P_","") for k in d.index],rotation=30,ha="right",fontsize=8)
ax.set_ylabel("Outcome % (hedges=incl. long index; plays=standalone)")
ax.set_title(f"COVID crash: NIFTY {case['covid_india_entry']['nifty']:.0f}→{case['covid_india_entry']['trough']:.0f} "
             f"({case['covid_india_entry']['trough_dd']*100:.0f}%), iVIX {case['covid_india_entry']['ivix']:.0f} at entry")
ax.legend(fontsize=8); fig.tight_layout(); fig.savefig(os.path.join(CH,"covid_india_payoff.png")); plt.close()

# 4) US RICH hedge efficiency frontier
h=hh_us[(hh_us.regime_filter=="RICH")].copy()
fig,ax=plt.subplots(figsize=(7.2,4))
ax.scatter(h["ret_sacrificed"]*100,h["dd_improve"]*100,c=BLUE,s=22)
for _,r in h.iterrows():
    if r["dd_improve"]>0.25 or r["sortino"]>2.0:
        ax.annotate(r["struct"].replace("H_",""),(r["ret_sacrificed"]*100,r["dd_improve"]*100),
                    fontsize=6.5,xytext=(3,2),textcoords="offset points")
ax.set_xlabel("Annual return sacrificed (pp)"); ax.set_ylabel("Max-drawdown improvement (pp)")
ax.set_title("U.S. RICH regime: hedge efficiency (upper-left = best protection per unit return)")
fig.tight_layout(); fig.savefig(os.path.join(CH,"hedge_frontier_us_rich.png")); plt.close()

# 5) last 2y cumulative
l=case["last2y"]
fig,axs=plt.subplots(1,2,figsize=(9,3.4))
for ax,mk in zip(axs,["US","INDIA"]):
    r=l[mk]
    labels=["Unhedged","Collar","PutSpread","ProtPut95","Play:Bear","Play:Back1x2","Play:LongPut"]
    vals=[r["hedge_collar_1m"]["cum_unhedged"],r["hedge_collar_1m"]["cum_combined"],
          r["hedge_putspread_1m"]["cum_combined"],r["hedge_protput95_1m"]["cum_combined"],
          r["play_bearspread_1m"]["cum_pnl"],r["play_backspread1x2_1m"]["cum_pnl"],r["play_longput95_1m"]["cum_pnl"]]
    cols=[GREY]+[BLUE]*3+[RED]*3
    ax.bar(range(len(vals)),[v*100 for v in vals],color=cols)
    ax.set_xticks(range(len(vals))); ax.set_xticklabels(labels,rotation=35,ha="right",fontsize=7)
    ax.axhline(0,color="k",lw=.6); ax.set_title(f"{mk}: 2024-07→2026-07 cumulative %",fontsize=9)
fig.suptitle("Last 2 years — no crash came: what each overlay actually returned",y=1.02)
fig.tight_layout(); fig.savefig(os.path.join(CH,"last2y_cum.png"),bbox_inches="tight"); plt.close()
print("charts done")

# ================= WORD DOC =================
doc=Document()
st=doc.styles["Normal"]; st.font.name="Calibri"; st.font.size=Pt(10.5)
def H(t,l=1):
    p=doc.add_heading(t,l)
    for r in p.runs: r.font.color.rgb=RGBColor(0x1f,0x3b,0x6f)
    return p
def P(t,b=False,it=False,sz=10.5,color=None):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=b; r.italic=it; r.font.size=Pt(sz)
    if color: r.font.color.rgb=color
    return p
def table(df,cols,headers=None,pct=None,fmt=None,widths=None):
    pct=pct or []; fmt=fmt or {}
    t=doc.add_table(rows=1,cols=len(cols)); t.style="Light Grid Accent 1"
    hdr=t.rows[0].cells
    for i,c in enumerate(cols):
        hdr[i].text=(headers[i] if headers else c)
        for para in hdr[i].paragraphs:
            for run in para.runs: run.bold=True; run.font.size=Pt(8.5)
    for _,row in df[cols].iterrows():
        cells=t.add_row().cells
        for i,c in enumerate(cols):
            v=row[c]
            if c in pct and isinstance(v,(int,float,np.floating)): s=f"{v*100:.1f}%"
            elif c in fmt and isinstance(v,(int,float,np.floating)): s=fmt[c].format(v)
            elif isinstance(v,(int,float,np.floating)): s=f"{v:.2f}"
            else: s=str(v)
            cells[i].text=s
            for para in cells[i].paragraphs:
                for run in para.runs: run.font.size=Pt(8.3)
    return t

# Title
tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=tp.add_run("Valuation-Regime Hedging & Downside-Play Study"); r.bold=True; r.font.size=Pt(19); r.font.color.rgb=RGBColor(0x1f,0x3b,0x6f)
sp=doc.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=sp.add_run("S&P 500 (real Shiller CAPE, 1871–2026) & NIFTY 50 (2016–2026)  •  Historical + Monte-Carlo  •  Shreyas_Ionic_AMC R&D  •  2026-07-08")
r.italic=True; r.font.size=Pt(10); r.font.color.rgb=RGBColor(0x60,0x60,0x60)

H("1. Executive summary",1)
for b in [
 f"Where we are NOW: U.S. is in a DEEP overvalued regime — Shiller CAPE {info['us_cape_now']:.1f}, above the 75th-percentile RICH threshold of {info['us_cape_q75']:.1f} and near a 150-year high. India is NOT expensive — NIFTY P/B {info['india_pb_now']:.2f} and trailing P/E {info['india_pe_now']:.1f} both sit in the CHEAP quartile of the 2016–2026 range. The 'overvalued downside' question is therefore a U.S. question today, not an India one.",
 "The overvalued (RICH) regime is defined by ASYMMETRY, not by low average returns. U.S. RICH months have historically had strong CONCURRENT returns (bull markets stay expensive) but the weakest FORWARD 12-month return (+3.9% mean) and by far the fattest left tail (10th-percentile −20.6%, worst −56%: the 1929 / 2000 / 2008 entries).",
 "TWO opposite 'downside' objectives — do not confuse them. (a) To PROFIT from an expensive regime on average, you SELL premium (e.g. a 1×2 put ratio): +EV, ~95% win-rate, but it is SHORT the crash tail and blows up in a real decline. (b) To be PROTECTED WHEN the crash comes, you BUY convexity (put spreads / 1×2 backspreads): negative carry that bleeds in bull markets but pays 5–30× in a COVID. Because you cannot time the turn, an overvaluation-driven programme should be built on (b), sized small — selling the tail you are worried about defeats the purpose.",
 "BEST ROLLOVER HEDGE (long-index overlay): the ANNUAL COLLAR (buy ~5% OTM put, sell ~5–10% OTM call). In the U.S. RICH regime it cut max drawdown from −52% to −13% to −19% and CVaR-5% from −37% to about −5% to −7%, for only ~3–4pp/yr of return given up — the best risk-adjusted protection (Sortino ~2.4–2.8). Tenor matters enormously: ANNUAL ≫ semi-annual ≫ quarterly ≫ monthly, because monthly rolling repeatedly pays the skew/theta bill.",
 "BEST DOWNSIDE PLAY for a known-overvalued regime (the honest answer): a small, cheap CONVEX structure — a 1×2 PUT BACKSPREAD (sell 1 near put, buy 2 further puts) or a defined-risk bear PUT SPREAD. The backspread carries near-zero cost (≈ −1.3% over the last two years vs −20% for outright long puts) yet is the structure that explodes in a crash. Premium-selling ratios show the highest average expectancy but are rejected for an overvaluation mandate on tail-risk grounds.",
 "COVID validation (India, entry 19-Feb-2020, iVIX only 14): an ATM protective put turned a −37% index collapse into −1.5%; a standalone long put returned +36%; the 3:2 put ratio +31%. Convexity bought CHEAP (before the vol spike) was worth more than its multi-year carry cost in a single event.",
 "The cost of being wrong (last 2 years, 2024→2026, no crash): U.S. unhedged +40%; collar-hedged +14%; outright downside plays −20%. Insurance is expensive precisely because the house usually does not burn — which is why tenor choice and cheap convexity, not naked long puts, are the whole game.",
]:
    p=doc.add_paragraph(style="List Bullet"); p.add_run(b).font.size=Pt(9.7)

H("2. Data, method & honesty statement",1)
for b in [
 "U.S.: monthly S&P 500 price and real Robert-Shiller CAPE 1871–2026 (multpl.com); dividend yield and 10-yr Treasury as q and r; real CBOE VIX daily 1990–2026 as implied vol (pre-1990: realised vol × 1.1).",
 "India: NIFTY 50 daily 2016–2026 (NSE official) with trailing P/E, P/B, dividend yield; India VIX daily 2016–2026 as implied vol. Valuation regime uses P/B as the primary metric — it is the CAPE-analog that stays stable through the 2020–21 earnings collapse that corrupts trailing P/E (trailing P/E spiked to 42 in 2020 on collapsed earnings, not on expensive prices). Trailing-P/E regimes are shown as a cross-check only.",
 "NO real option chains exist for this span. Every option is Black-Scholes-priced at entry from the implied vol above PLUS an equity put-skew term (SPX slope 0.90, NIFTY 0.50 vol-pts per 100% moneyness), and settled at realised intrinsic at expiry. The entry-IV → realised-intrinsic gap IS the volatility risk premium / true hedge cost, so premium-selling is not free in this model. This is the 'best-estimate IV' approach you authorised.",
 "Regimes: valuation split at the 25th / 75th percentile of each metric's own history (25-50-25 → CHEAP / FAIR / RICH). Thresholds are full-sample (they use hindsight only to draw the lines, not to trade). Costs are DRAFT (8 bps spread + 2 bps fixed per leg) — NOT the firm's approved COST_STANDARDS; results shown net.",
 "Both a HISTORICAL non-overlapping rollover backtest and a regime-conditional MONTE-CARLO (block-bootstrap of that regime's monthly returns, 8,000 paths) were run for every structure × tenor × regime; MC agreement is reported to flag small-sample historical artefacts (esp. the U.S. RICH tail and India's 10-year window).",
 "Key limitations: India has only ~10 years (one real crash, COVID) so its RICH regime never saw a 2008-style event and looks benign in-sample; U.S. option realism pre-1990 relies on modelled IV; skew and term-structure are parametric, not surface-fitted. Directional conclusions are robust; exact basis points are not.",
]:
    p=doc.add_paragraph(style="List Bullet"); p.add_run(b).font.size=Pt(9.5)

H("3. The three valuation regimes — return & volatility",1)
doc.add_picture(os.path.join(CH,"us_cape_regimes.png"),width=Inches(6.6))
doc.add_picture(os.path.join(CH,"regime_fwd_returns.png"),width=Inches(6.6))
P("U.S. S&P 500 by CAPE regime (monthly, annualised):",b=True,sz=10)
table(rs_us,["regime","n_months","median_ann_ret","ann_vol","fwd12m_mean","fwd12m_median","fwd12m_p10","fwd12m_worst","pct_neg_month"],
      headers=["Regime","N mo","Median ann ret","Ann vol","Fwd12m mean","Fwd12m med","Fwd12m p10","Fwd12m worst","% neg mo"],
      pct=["median_ann_ret","ann_vol","fwd12m_mean","fwd12m_median","fwd12m_p10","fwd12m_worst","pct_neg_month"])
P("India NIFTY 50 by P/B regime (primary) — monthly, annualised:",b=True,sz=10)
table(rs_in,["regime","n_months","median_ann_ret","ann_vol","fwd12m_mean","fwd12m_median","fwd12m_p10","fwd12m_worst","pct_neg_month"],
      headers=["Regime","N mo","Median ann ret","Ann vol","Fwd12m mean","Fwd12m med","Fwd12m p10","Fwd12m worst","% neg mo"],
      pct=["median_ann_ret","ann_vol","fwd12m_mean","fwd12m_median","fwd12m_p10","fwd12m_worst","pct_neg_month"])
P("India by trailing-P/E regime (cross-check — note 2020-21 earnings-collapse contamination):",it=True,sz=9)
table(rs_inpe,["regime","n_months","median_ann_ret","ann_vol","fwd12m_mean","fwd12m_p10","fwd12m_worst"],
      headers=["Regime","N mo","Median ann ret","Ann vol","Fwd12m mean","Fwd12m p10","Fwd12m worst"],
      pct=["median_ann_ret","ann_vol","fwd12m_mean","fwd12m_p10","fwd12m_worst"])
P("Read: in the U.S., RICH = strongest concurrent return but weakest forward return and the fattest tail — the textbook overvaluation signature. In India (short, growth-market sample) even RICH stays forward-positive; CHEAP has been the highest forward-return regime (+24% fwd 12m on P/B).",it=True,sz=9)

H("4. Best ROLLOVER HEDGE (overlay on a long-index position)",1)
doc.add_picture(os.path.join(CH,"hedge_frontier_us_rich.png"),width=Inches(5.6))
P("U.S. — top hedges by Sortino, per regime (vs unhedged baseline in caption):",b=True,sz=10)
for rg in ["RICH","FAIR","CHEAP"]:
    sub=hh_us[hh_us.regime_filter==rg].head(4)
    ub=sub.iloc[0]
    P(f"{rg}  (unhedged: ann {ub['unhedged_ann']*100:.1f}%, maxDD {ub['unhedged_maxdd']*100:.0f}%, CVaR5 {ub['unhedged_cvar5']*100:.0f}%)",b=True,sz=9)
    table(sub,["struct","tenor","ann_ret","maxdd","cvar5","sortino","ret_sacrificed","dd_improve","avg_cost"],
          headers=["Structure","Tenor","Ann ret","MaxDD","CVaR5","Sortino","Ret giv.up","DD improve","Avg cost"],
          pct=["ann_ret","maxdd","cvar5","ret_sacrificed","dd_improve","avg_cost"])
P("India — top hedges by Sortino (RICH & CHEAP):",b=True,sz=10)
for rg in ["RICH","CHEAP"]:
    sub=hh_in[hh_in.regime_filter==rg].head(3)
    table(sub,["struct","tenor","ann_ret","maxdd","cvar5","sortino","ret_sacrificed","avg_cost"],
          headers=["Structure","Tenor","Ann ret","MaxDD","CVaR5","Sortino","Ret giv.up","Avg cost"],
          pct=["ann_ret","maxdd","cvar5","ret_sacrificed","avg_cost"])
P("Conclusion: the ANNUAL COLLAR is the efficient frontier of protection — it self-finances the put by selling upside, so tail risk falls 5–7× for ~3–4pp/yr. Rolling MONTHLY is the worst choice (pays skew 12× a year). If you refuse to cap upside (no short call), an annual ATM put is next-best but costs ~6pp/yr.",it=True,sz=9)

H("5. Best DOWNSIDE PLAY in the overvalued (RICH) regime",1)
P("U.S. RICH — standalone plays (mean P&L per period, % of notional; MC = Monte-Carlo cross-check):",b=True,sz=10)
sub=pp_us[pp_us.regime_filter=="RICH"].head(8)
table(sub,["struct","tenor","mean_pnl","median_pnl","winrate","worst","best","mc_mean","mc_p05"],
      headers=["Structure","Tenor","Mean P&L","Median","Win%","Worst","Best","MC mean","MC p05"],
      pct=["mean_pnl","median_pnl","winrate","worst","best","mc_mean","mc_p05"])
P("India RICH — standalone plays:",b=True,sz=10)
sub=pp_in[pp_in.regime_filter=="RICH"].head(6)
table(sub,["struct","tenor","mean_pnl","median_pnl","winrate","worst","best","mc_mean","mc_p05"],
      headers=["Structure","Tenor","Mean P&L","Median","Win%","Worst","Best","MC mean","MC p05"],
      pct=["mean_pnl","median_pnl","winrate","worst","best","mc_mean","mc_p05"])
P("The 1×2 / 2×1 put RATIOS post the highest average expectancy and win-rate — but they are SHORT the tail (worst −11% in the U.S., and they are exactly what fails in a crash). For an OVERVALUATION mandate (whose entire point is tail risk) the recommended play is a 1×2 PUT BACKSPREAD or a defined-risk BEAR PUT SPREAD: near-zero carry, defined risk, and convex when it matters (see COVID, §6).",it=True,sz=9)

H("6. Case studies",1)
P("6a. COVID crash (the tail event these structures exist for)",b=True,sz=11)
doc.add_picture(os.path.join(CH,"covid_india_payoff.png"),width=Inches(6.4))
ci=case["covid_india_entry"]; cu=case["covid_us_entry"]
P(f"India: entry {ci['date']} at NIFTY {ci['nifty']:.0f}, iVIX only {ci['ivix']:.0f}; trough {ci['trough_date']} at {ci['trough']:.0f} ({ci['trough_dd']*100:.0f}%). "
  f"U.S.: entry {cu['date']} S&P {cu['sp500']:.0f}, VIX {cu['vix']:.0f}; month-end March {cu['mar_sp500']:.0f} ({cu['mar_dd']*100:.0f}%).",sz=9.5)
P("India COVID outcomes at the trough (hedges = total incl. long index; plays = standalone), selected structures:",b=True,sz=9.5)
sel=["H_put_ATM","H_collar_95_105","H_putspread_95_85","P_longput_97","P_bearspread_95_85","P_backspread_1x2_97_90","P_ratio_3x2_95_85"]
cc=cov_in[(cov_in.scenario=="1m@trough")&(cov_in.struct.isin(sel))]
table(cc,["struct","frame","idx_ret","entry_cost","opt_pnl","combined"],
      headers=["Structure","Frame","Index ret","Entry cost","Option P&L","Combined"],
      pct=["idx_ret","entry_cost","opt_pnl","combined"])
P("The ATM protective put converted −37% into −1.5%; convex plays returned +9% to +36%. The lesson is timing of PURCHASE, not of the crash: iVIX was 14 pre-COVID, so protection was cheap going in.",it=True,sz=9)

P("6b. Last 2 years — 2024-07 → 2026-07 (the counterfactual: no crash)",b=True,sz=11)
doc.add_picture(os.path.join(CH,"last2y_cum.png"),width=Inches(6.4))
l=case["last2y"]
rows=[]
for mk in ["US","INDIA"]:
    r=l[mk]
    rows.append(dict(Market=mk,Unhedged=r["hedge_collar_1m"]["cum_unhedged"],
        Collar=r["hedge_collar_1m"]["cum_combined"],PutSpread=r["hedge_putspread_1m"]["cum_combined"],
        ProtPut95=r["hedge_protput95_1m"]["cum_combined"],
        Play_BearSpr=r["play_bearspread_1m"]["cum_pnl"],Play_Back1x2=r["play_backspread1x2_1m"]["cum_pnl"],
        Play_LongPut=r["play_longput95_1m"]["cum_pnl"]))
l2=pd.DataFrame(rows)
table(l2,list(l2.columns),pct=[c for c in l2.columns if c!="Market"])
P("With no crash, hedging cost the U.S. book ~24pp of a +40% bull (collar retained +14%); outright downside plays lost ~20%. The only near-free carry was the 1×2 backspread (−1.3%). This quantifies the premium you pay for convexity you did not (yet) need — and why the recommendation is CHEAP convexity + upside-financed collars, not naked long puts.",it=True,sz=9)

P("6c. What I would do RIGHT NOW",b=True,sz=11)
for b in [
 f"U.S. equity exposure — regime RICH (CAPE {info['us_cape_now']:.1f}). Run an ANNUAL COLLAR (buy ~5% OTM SPX/ES put, sell ~7–10% OTM call to finance it) as the core overlay: it takes the tail from ~−50% to ~−15% for ~3pp/yr. Add a SMALL 1×2 put backspread (≈0.3–0.5% of notional/yr) as a convex crash kicker. Do NOT sell premium (ratios/short calls) as your 'downside play' — that is short the exact tail the CAPE signal is warning about.",
 f"India equity exposure — regime CHEAP (P/B {info['india_pb_now']:.2f}, P/E {info['india_pe_now']:.1f}). The data says stay long and largely UNHEDGED; CHEAP has been the best forward-return regime (+24% fwd). At most hold a cheap far-OTM annual put as fire-insurance; no active downside play is warranted while India is in its cheap quartile.",
 "Trigger to escalate India hedging: P/B back above ~4.0 (the 75th-percentile RICH line) — then rotate India onto the same annual-collar programme.",
]:
    p=doc.add_paragraph(style="List Bullet"); p.add_run(b).font.size=Pt(9.7)

H("7. Final recommendations — by regime (both markets)",1)
recs=pd.DataFrame([
 dict(Regime="CHEAP (≤25th pct)",Hedge="Minimal — cheap far-OTM tail put only; stay long",Play="None (forward returns strongest); if anything, sell puts to accumulate",Tenor="Annual"),
 dict(Regime="FAIR (25–75th)",Hedge="Annual collar OR annual ATM put if upside sacred",Play="None / opportunistic bear put spread around events",Tenor="Annual > semi"),
 dict(Regime="RICH (≥75th pct)",Hedge="ANNUAL COLLAR (put ~5% OTM / short call 5–10% OTM) — core",Play="Small 1×2 PUT BACKSPREAD or defined-risk BEAR PUT SPREAD (convex, cheap)",Tenor="Annual (never monthly)"),
])
table(recs,["Regime","Hedge","Play","Tenor"])
P("Universal rules from the study: (1) buy protection by TENOR = annual, not monthly — monthly rolling pays the skew bill 12× a year; (2) FINANCE puts with short calls (collar) rather than paying outright, unless you must keep uncapped upside; (3) for an overvaluation mandate, BUY convexity, never SELL it — premium-selling ratios win most months but lose the one that matters; (4) the biggest driver of hedge P&L is the IV you PAY at entry — protection bought in calm (low VIX/iVIX) is cheap and pays hugely (COVID); chasing protection after the spike is a losing game.",it=True,sz=9)

P("Files: full grids, rankings, regime stats and case-study CSVs in \\HEDGING_ANALYSIS_20260708\\results\\; engine.py / summarize.py reproduce everything; charts in \\charts\\.",it=True,sz=8.5,color=RGBColor(0x60,0x60,0x60))

path=os.path.join(OUT,"HEDGING_ANALYSIS_REPORT.docx")
doc.save(path); print("SAVED",path)
