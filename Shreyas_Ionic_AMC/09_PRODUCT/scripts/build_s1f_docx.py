"""Build S1-F strategy Word doc (Principal-facing, D-020 human format).
v1.1 (2026-07-11): rehomed from scratchpad to 09_PRODUCT/scripts; section 4 upgraded with
A4-card REAL-settle COVID replication (supersedes model-only stress paragraph)."""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt

ROOT = Path(r"c:\Users\Shreyas.1Gupta\OneDrive - Angel Broking Limited\Desktop\Backup\NIFTY 500")
RES = ROOT / "Shreyas_Ionic_AMC/04_RND_LAB/results/SELLSIDE_20260710"
OUT = ROOT / "Shreyas_Ionic_AMC/09_PRODUCT/reports"
OUT.mkdir(parents=True, exist_ok=True)

doc = Document()
doc.add_heading("S1-F: 0DTE NIFTY ATM Short Straddle", 0)
p = doc.add_paragraph("Frozen strategy specification & evidence pack — registered for paper forward test\n"
                      "Shreyas_Ionic_AMC · 2026-07-10 · Spec v1.0 (D-030 freeze) · Evidence pack updated 2026-07-11")
p.runs[0].font.size = Pt(11)

doc.add_heading("1. The strategy in one paragraph", 1)
doc.add_paragraph(
    "On every NIFTY weekly expiry day, at 09:20, sell the at-the-money straddle (ATM call + ATM put, same-day "
    "expiry) with a 30% stop-loss on each leg; whatever survives is closed at 15:25. Two veto rules skip "
    "dangerous days (stretched 5-day RSI; a >1.5% prior-day move). Size 3-4 lots per Rs 10L at dynamic margin, "
    "halved when 3-day volatility exceeds 2x its one-year median. It earns the volatility risk premium: expiry-day "
    "option sellers are systematically overpaid for a move that usually does not come, and the stop amputates the "
    "days it does.")

doc.add_heading("2. Frozen specification", 1)
rows = [
    ("Universe / day", "NIFTY weekly options, expiry days only (from live contract data; handles Tue/Thu shifts)"),
    ("Entry", "09:20 - SELL ATM CE + ATM PE (strike = round(spot/50)x50), same-day expiry"),
    ("Stop-loss", "30% per leg on 1-min close; market exit on breach"),
    ("Exit", "15:25 all surviving legs; always flat EOD; no re-entry"),
    ("Veto F1", "Skip if daily RSI(5) at D-1 close is >=80 or <=20"),
    ("Veto F2", "Skip if |D-1 daily return| > 1.5%"),
    ("Sizing", "floor(0.75 x equity / margin) lots; margin = ~15% of notional (Rs ~2.7L/lot at 2026 levels, "
     "verify on broker calculator) -> ~3-4 lots per Rs 10L; lot 75. Runner now computes this dynamically."),
    ("Crash rule", "Halve lots when 3-day realized vol > 2x its 1-year median"),
    ("Kill criteria (forward)", "26 expiries: expectancy<=0 -> KILL; paper maxDD>15% -> KILL; fills 3+ pts/day worse than model over 13 expiries -> HALT"),
]
t = doc.add_table(rows=0, cols=2); t.style = "Light Grid Accent 1"
for a, b in rows:
    r = t.add_row().cells; r[0].text = a; r[1].text = b

doc.add_heading("3. Evidence summary", 1)
doc.add_paragraph(
    "259 real expiry days (Jun-2021 to Jun-2026), 1-minute option prices, costs = 1% slippage per fill + STT/"
    "exchange/GST + brokerage. Net +10.7 pts/day unconditional (t=3.92, PF 1.79, 69% win, both eras positive); "
    "+11.3 pts/day with vetoes (t=3.73 on 204 days). Parameter surface: 72/84 neighboring configurations "
    "positive (a plateau, not a lucky spike). Verdicts unchanged under a harsher flat-point cost model. "
    "At spec sizing with REALISTIC dynamic margin (~15% of notional; Rs 1.8L/lot in 2021 rising to Rs 2.7L/lot "
    "in 2026): Rs 10L -> Rs 18.7L over 5.0 years (13.4% CAGR), maximum drawdown -4.4%. (A flat-1.1L margin "
    "assumption showed 28-31% CAGR; that figure is superseded - the corrected margin model matches broker "
    "SPAN+exposure calculators. Return on margin deployed, ~+0.5% per expiry, is the invariant.)")
img1 = RES / "s1f_final_graph/S1F_EQUITY_DYNMARGIN.png"
if not img1.exists():
    img1 = RES / "s1f_final_graph/S1F_FINAL_EQUITY.png"
if img1.exists(): doc.add_picture(str(img1), width=Inches(6.5))
doc.add_paragraph("Equity and drawdown at corrected dynamic margin, net of costs (in-sample).").runs[0].font.size = Pt(9)
img2 = RES / "final_three/FINAL_THREE_PNL.png"
if img2.exists(): doc.add_picture(str(img2), width=Inches(6.5))
doc.add_paragraph("Per-lot cumulative P&L: S1 vs challengers S1b (ATM-50) and V2 (defense strangle).").runs[0].font.size = Pt(9)

doc.add_heading("4. Crash stress — now on REAL prices (updated 2026-07-11)", 1)
doc.add_paragraph(
    "NEW EVIDENCE (A4 experiment, pre-registered): the 2011-2021 NSE settlement archive was backfilled and the "
    "strategy's structure (short ATM straddle + 30% per-leg stop) was replicated at monthly cadence on REAL "
    "exchange prices through COVID. Result: the COVID-window drawdown was 1.05x the worst normal-era (2011-2019) "
    "stretch - against a pre-registered kill bar of 3x. The crash cycle itself (entered Feb-2020, expiring "
    "26-Mar-2020 through the bottom) lost 544 points on a ~730-point premium: the stop-loss amputated the "
    "catastrophe as designed (unstopped, the same cycle path shows five-figure point losses). Every month of "
    "2020 traded. The strategy's crash survival no longer rests on a model.")
doc.add_paragraph(
    "The earlier Black-Scholes reconstruction on 2020 minute data (validated corr 0.64) remains as supplementary "
    "evidence and agrees: roughly FLAT through Jan-2020..May-2021 under harsh IV assumptions, maxDD ~-16% at "
    "higher-than-spec sizing. Remaining honest limits: stops evaluated on 1-minute closes in the primary backtest "
    "(tick fills somewhat worse); ~165 research trials preceded this spec, so the paper forward test with "
    "pre-registered kill criteria remains the final arbiter; a gap through both stops beyond anything in 15 years "
    "of data remains the tail risk - the crash rule and modest sizing are the defenses.")

doc.add_heading("5. Paper-trading procedure", 1)
doc.add_paragraph(
    "Each expiry morning ~09:10 run 06_TRADING_DESK/paper/s1f_daily_runner.py (hardened 2026-07-11: dynamic "
    "margin sizing, expiry calendar from live contract data - first ticket Tue 2026-07-14, cron armed). It "
    "evaluates F1/F2 and the crash rule from live data, prints GO/SKIP with the exact order ticket and logs the "
    "intent BEFORE market action to s1f_paper_log.csv. Actual 09:20 fills and exits are marked in the same file; "
    "weekly reconciliation per the firm paper process. Shadow-tracked at zero size: unconditional S1 and "
    "challenger S1b (ATM-50).")

doc.add_heading("6. What was tried and killed getting here", 1)
doc.add_paragraph(
    "18 option-BUYING designs - all negative or insignificant; a 16-indicator screen measured the intraday "
    "information ceiling at ~2.4 pts vs the ~6 needed for buying. Weekly/biweekly strangles: positive but weak. "
    "Iron flies/condors: wing costs exceed the edge. PCR/most filters: no value. NEW (2026-07-11, five "
    "pre-registered cards): overnight premium REFUTED on our data (it is intraday; overnight selling is a "
    "steamroller trap); no DTE bucket beats 0DTE-with-stop (the edge is MANUFACTURED by the stop: same day/entry "
    "without SL = -1.5 pts/day vs +10.7 with); US-session gap model real (R2 0.22) but adds nothing to S1's "
    "vetoes; FII flow signal killed at the significance bar. Full evidence: 04_RND_LAB/results/.")

doc.save(OUT / "S1F_STRATEGY_PACK_20260710.docx")
print("saved ->", OUT / "S1F_STRATEGY_PACK_20260710.docx")
