# FF calendar (S-03/K-012) VERDICT ADDENDUM - supersedes the "CONDITIONAL YES" in
# FF_CALENDAR_BRIEF_2026-07-05.docx after the CIO ruling (2026-07-05).
# Numbers: results/S-03/20260705_resurrection/CAUSAL_RETEST.md + CIO_RULING.md
import os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = r"C:\Users\Shreyas.1Gupta\OneDrive - Angel Broking Limited\Desktop\Backup\NIFTY 500"
OUTD = os.path.join(ROOT, "Shreyas_Ionic_AMC", "09_PRODUCT", "reports")
PNG = os.path.join(OUTD, "ff_verdict_ladder.png")
DOCX = os.path.join(OUTD, "FF_CALENDAR_VERDICT_2026-07-05.docx")

BLUE, AQUA, YELLOW, GREEN, VIOLET, RED = "#2a78d6", "#1baf7a", "#eda100", "#008300", "#4a3aa7", "#e34948"

# ---- chart: the honest ladder (forward cohort, per Rs100 deployed) ----
labels = ["Headline\n(hindsight entry,\nflat slippage)", "Causal entry\n(leak removed)",
          "+ Fill gate\n(dead markets out)", "+ Tiered slippage\n= VERDICT",
          "2x cost stress", "Same-day entry\n(EXPLORATORY)"]
vals = [10.04, 8.75, 2.30, -0.03, -2.36, 0.99]
cols = [BLUE, BLUE, BLUE, RED, RED, YELLOW]
fig, ax = plt.subplots(figsize=(8.6, 4.2), dpi=160)
bars = ax.bar(range(len(vals)), vals, color=cols, width=0.62)
ax.axhline(0, color="#555555", lw=1)
for i, v in enumerate(vals):
    ax.text(i, v + (0.35 if v >= 0 else -0.55), f"{v:+.2f}", ha="center",
            fontsize=10, fontweight="bold", color="#222222")
ax.set_xticks(range(len(labels)))
ax.set_xticklabels(labels, fontsize=8.3)
ax.set_ylabel("Forward P&L per Rs100 premium deployed", fontsize=9)
ax.set_title("FF Calendar: where the +10 went (forward 2024-26, n=199 signals)", fontsize=11, fontweight="bold")
ax.spines[["top", "right"]].set_visible(False)
ax.set_ylim(min(vals) - 1.6, max(vals) + 1.6)
fig.tight_layout()
fig.savefig(PNG)
plt.close(fig)

# ---- docx ----
doc = Document()
st = doc.styles["Normal"]
st.font.name = "Calibri"; st.font.size = Pt(10.5)

h = doc.add_heading("FF Calendar - Final Verdict (supersedes the 'Conditional Yes' brief)", level=0)
p = doc.add_paragraph()
r = p.add_run("CIO ruling 2026-07-05: STAYS KILLED - with the FF signal handed to the Structurer as a NEW idea on a liquid vehicle. "
              "Zero capital, zero paper capital on the calendar version.")
r.bold = True; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0xE3, 0x49, 0x48)

doc.add_paragraph(
    "Your challenge was right, and it was also not enough: the review proved the FF signal is REAL "
    "(it beat 2,000 matched random portfolios at the 100th percentile - the sizing fix was genuine), "
    "but the pre-registered honest re-test showed the CALENDAR VEHICLE cannot monetize it. "
    "61% of forward signals fire on options with zero volume and mostly zero open interest - markets that do not exist. "
    "Per your limit-or-skip rule those trades are skipped; what remains, entered without hindsight and with "
    "volume-tiered slippage, earns nothing.")

doc.add_picture(PNG, width=Inches(6.7))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading("The decomposition ladder (forward cohort)", level=1)
t = doc.add_table(rows=7, cols=3); t.style = "Light Grid Accent 1"
rows = [("Step", "P&L per Rs100 deployed", "What it isolates"),
        ("Headline (hindsight peak-FF entry, flat 1.5% slippage)", "+10.04", "The number the resurrection case was built on"),
        ("Causal first-cross entry", "+8.75", "Hindsight-entry leak removed: -1.2 (real but minor)"),
        ("+ ex-ante liquidity gate", "+2.30", "Dead back-leg markets skipped: -6.45 (the killer)"),
        ("+ volume-tiered slippage = PRE-REGISTERED VERDICT", "-0.03", "Slippage on thin survivors: -2.33. Survivors: 81 trades, 51.9% win, PF 0.99 - a coin flip"),
        ("2x cost stress (firm certification bar)", "-2.36", "Fails decisively"),
        ("Same-day-close entry (EXPLORATORY, outside verdict)", "+0.99", "Optimistic-but-arguable bound; also dies at 2x")]
for i, row in enumerate(rows):
    for j, cell in enumerate(row):
        c = t.rows[i].cells[j]; c.text = cell
        if i == 0:
            for pr in c.paragraphs:
                for rn in pr.runs: rn.bold = True

doc.add_paragraph("Build side is negative too (-0.51 per Rs100): with honest execution there was never an edge in-sample either. "
                  "Per-year at the verdict step: 2021 -0.98 | 2022 -2.34 | 2023 -0.13 | 2024 +0.48 | 2025 +0.59 | 2026 -3.81.")

doc.add_heading("Why the CIO said no to even paper-tracking", level=1)
doc.add_paragraph(
    "1) Your D-031 ruling relaxes the CAPACITY bar for exceptional strategies - it does not relax the EDGE bar, and the honest edge is zero. "
    "2) The limit-or-skip convention is already priced in: skipping unfillable trades IS what the -0.03 assumes. "
    "3) Exit risk: 61% dead markets means positions you may not be able to CLOSE - the CIO vetoes on exitability alone, whatever the edge. "
    "4) The headline's worst trade (-464) was itself unfillable - even the risk numbers of the old backtest were fiction.")

doc.add_heading("What survives, and what happens next", level=1)
doc.add_paragraph(
    "SURVIVES: the FF signal itself (placebo-proven information), the equal-premium sizing lesson (now firm standard thinking), "
    "and four new laws in the knowledge base (fill-rate audits come BEFORE sizing debates; test gates against simple dropping; "
    "diff legacy script versions for introduced lookahead; pre-register the entry-timing convention). "
    "NEXT: Aakash (Structurer) receives a NEW intake - the FF signal on a liquidity-native vehicle (liquid front-month structures, futures basis, or index calendars) "
    "with 5 pre-registered kills including the full ~34-trial family honesty count. If a tradeable vehicle exists, this signal earns its way back through every gate; if not, it dies clean.")

p = doc.add_paragraph()
r = p.add_run("Full audit trail: results/S-03/20260705_resurrection/ (4 review legs + CIO_RULING.md). "
              "This addendum supersedes the verdict section of FF_CALENDAR_BRIEF_2026-07-05.docx.")
r.font.size = Pt(8.5); r.italic = True

doc.save(DOCX)
print("OK wrote", DOCX)
