"""Assemble FIRM_SYSTEM_BLUEPRINT from researched sections -> master .md + Principal .docx.
Scripts-first fallback (writer agent blocked by spend limit): mechanical assembly, exec summary
and roadmap intro authored inline, per-section improvement subsections relocated to the roadmap.
"""
import re
from pathlib import Path

ROOT = Path(r"c:\Users\Shreyas.1Gupta\OneDrive - Angel Broking Limited\Desktop\Backup\NIFTY 500")
REP = ROOT / "Shreyas_Ionic_AMC/09_PRODUCT/reports"
SEC = REP / "_blueprint_sections"

EXEC = """# SHREYAS_IONIC_AMC — Firm System Blueprint
*Prepared for the Principal — 13 July 2026. Sources: the firm's own files, read directly; six researched sections assembled verbatim below.*

## Executive Summary

**What this is.** Shreyas_Ionic_AMC is a quantamental trading and investing firm operated end-to-end by AI agents on one laptop, across two Claude accounts (a CIO/R&D desk and an execution desk), under a written constitution and 34 binding Principal rulings. Everything material — governance, research verdicts, data catalogs, session logs — lives in version-controlled files, so any session (or any future model) can resume the firm cold.

**What has been built, in four layers:**
1. **A firm** — 28 named agent-employees with roles, virtual compensation (AlphaPoints economy, approx Rs 38.35 Cr virtual payroll), an investment committee with blind memo rounds, a red team with mandatory pre-certification review, and a four-tier approval hierarchy where only LIVE capital and risk-limit loosening reach the Principal.
2. **A data estate** — 15+ verified datasets: 26 years of Indian daily equities (survivorship-controlled via 42 point-in-time index snapshots), 813M+ minute bars, a complete 15-year index-derivatives panel (2011-2026), participant-flow data, point-in-time earnings with second-precision announcement timestamps (2019+), plus US/global layers (S&P membership 1996-2026, SPX 1975+, vol suite, factors 1926+, gold/crypto minute data, USDINR 1973+).
3. **A research machine** — the firm's real moat: every experiment is frozen in its own git commit *before* it runs (provable pre-registration), passed through a static lookahead scanner, adjudicated against pre-registered placebo batteries (13 distinct controls, each born from a real in-house incident), and banked with resurrection conditions. 255 trials in the ledger; roughly 95% killed, every kill reproducible.
4. **A trading book (paper)** — honest labels: **2 certified alpha sleeves** (S1-F 0DTE index straddle; B1b FII-minus-Client futures flow), **2 labeled betas** (midsmall momentum rotation, breakout pack — both red-teamed and kept only with binding relabels), **3 forward shadows** (P6 snapback, B1c DII flow, S1-SX SENSEX Thursday). Stacked-book frontier: 15.8% CAGR / -8.1% maxDD / Sharpe 2.29 (quality point) to 35.9% / -22.1% / 1.91 (growth point), with the correlation-horizon caveat documented.

**Headline security findings (Section 6, action required):** (i) the complete Angel login secret set — API key, client ID, PIN and TOTP seed — sits in plaintext, with a second forgotten copy in an old session scratchpad; the TOTP seed defeats two-factor entirely (HIGH; fund-less account caps monetary damage). (ii) The daily option-capture task fails silently when the laptop is asleep/on battery — data Angel purges at expiry is then unrecoverable (HIGH, operational). (iii) The weekly data-snapshot backup layer exists on paper only (MEDIUM). (iv) The entire firm syncs to the employer's OneDrive tenant — an explicit accept-or-move decision is owed (MEDIUM-HIGH).

**Reading order.** Section 1 explains who decides what; Section 2 who does the work; Section 3 how research is kept honest; Section 4 what data exists and its landmines; Section 5 what runs every day and what the book holds; Section 6 the platform and its risks; the final chapter is the consolidated improvement roadmap.
"""

ROADMAP_INTRO = """# Improvement Roadmap (consolidated)

*Each section's researcher filed improvement opportunities for their own area; they are consolidated here unedited, prefaced by the ten highest-leverage items across the whole firm.*

## Top 10 across the firm (priority order)
1. **Credential hygiene (HIGH, same-day):** reference-check then delete the stale scratchpad credentials copy (`angel_cfg.py` + `.pyc`); move the canonical `creds.json` to Windows Credential Manager / DPAPI; keep one offline break-glass copy with the Principal.
2. **Capture-task hardening (HIGH, same-day):** switch `AngelDailyOptionCapture` to "run whether logged on or not" + allow-on-batteries; add a Last-Result alarm to `/eod` (the last run failed with 0x8007052B and nothing alerted).
3. **Backup layer 3 (P1):** implement the weekly snapshot of critical derived datasets that policy already mandates (zip + rotate 4, Sunday cron) and log the quarterly restore drill.
4. **Governance de-staling sweep (P1):** reconcile the stale "6 parallel agents" lines vs D-023's 3; repair the MODEL_ASSIGNMENTS broken table (11 stranded rows); complete truncated D-032; renumber KNOWLEDGE_BASE duplicates; add supersession marks (D-009 -> D-033).
5. **Trials-ledger automation (P1):** auto-rebuild TRIALS_LEDGER.csv from RUN_CARD.json files; add a freeze-hash compliance tripwire (verify each card's engine ran at its frozen commit).
6. **Monthly-correlation standing gate (P1):** make monthly/quarterly-horizon correlation a mandatory field in every sleeve verdict (the daily-corr artifact caught 2026-07-13).
7. **Skill-library hygiene (P2):** flesh out the one-line `lookahead-audit` stub (a mandatory gate deserves a real skill file); consolidate ~10 near-duplicate design skills.
8. **OneDrive tenancy decision (Principal):** explicitly accept (log in DECISIONS_LOG) or relocate the firm off the employer tenant.
9. **Data unlocks (Principal, both free):** Kaggle API key (Quandl WIKI mirror = pre-2018 US delisted prices) and Tiingo free key (2018-26 dead-name tail) — completes the US survivorship fix.
10. **Second capture site (P2):** a small cloud VM or home box as redundant Angel capture, removing the single-laptop SPOF on purge-sensitive data.
"""

order = ["01_governance.md", "02_agents.md", "03_methodology.md", "04_data.md", "05_desk_ops.md", "06_platform_security.md"]
body_parts, improve_parts = [], []
for fn in order:
    txt = (SEC / fn).read_text(encoding="utf-8")
    m = re.search(r"^### Improvement opportunities\s*$", txt, flags=re.M)
    if m:
        improve = txt[m.start():].strip()
        txt = txt[:m.start()].rstrip()
    else:
        improve = ""
    area = fn.split("_", 1)[1].replace(".md", "").replace("_", " ").title()
    body_parts.append(txt)
    if improve:
        improve_parts.append(f"## {area}\n" + "\n".join(improve.splitlines()[1:]).strip())

master = EXEC + "\n\n---\n\n" + "\n\n---\n\n".join(body_parts) + "\n\n---\n\n" + ROADMAP_INTRO + "\n\n" + "\n\n".join(improve_parts) + """

---
# How to read this firm (for a newcomer)
Start with the root `CLAUDE.md` (the constitution), then `01_COMMAND_CENTER/CURRENT_STATE.md` (what is true right now) and the last two entries of `SESSION_JOURNAL.md` (what just happened). `DECISIONS_LOG.md` holds every binding Principal ruling. Research verdicts live in `04_RND_LAB/STOCKS_PROGRAM_2026/MASTER_PLAN.md` and results folders — every number quoted anywhere must trace to a frozen card and a results file. Nothing in this firm is real money; the paper-to-live gate belongs to the Principal alone.
"""
md_path = REP / "FIRM_SYSTEM_BLUEPRINT_20260713.md"
md_path.write_text(master, encoding="utf-8")
print(f"master md: {len(master.split())} words -> {md_path.name}")

# ---------- md -> docx ----------
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Georgia"; st.font.size = Pt(10.5)

def add_md_text(p, text):
    # minimal inline md: **bold**, *italic*, `code`
    for tok in re.split(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)", text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9.5)
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            r = p.add_run(tok[1:-1]); r.italic = True
        else:
            p.add_run(tok)

lines = master.splitlines()
i = 0
while i < len(lines):
    ln = lines[i]
    if ln.startswith("|") and i + 1 < len(lines) and set(lines[i + 1].replace("|", "").replace(" ", "")) <= set("-:"):
        # table block
        rows = []
        while i < len(lines) and lines[i].startswith("|"):
            cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
            if not set("".join(cells)) <= set("-: "):
                rows.append(cells)
            i += 1
        ncol = max(len(r) for r in rows)
        t = doc.add_table(rows=len(rows), cols=ncol)
        t.style = "Light Grid Accent 1"
        for ri, r in enumerate(rows):
            for ci in range(ncol):
                cell = t.cell(ri, ci)
                cell.text = ""
                p = cell.paragraphs[0]
                add_md_text(p, r[ci] if ci < len(r) else "")
                for run in p.runs:
                    run.font.size = Pt(8.5)
                    if ri == 0:
                        run.bold = True
        doc.add_paragraph()
        continue
    if ln.startswith("# "):
        doc.add_heading(re.sub(r"[*`]", "", ln[2:]), level=0 if "Blueprint" in ln else 1)
    elif ln.startswith("## "):
        doc.add_heading(re.sub(r"[*`]", "", ln[3:]), level=2)
    elif ln.startswith("### "):
        doc.add_heading(re.sub(r"[*`]", "", ln[4:]), level=3)
    elif ln.startswith("#### "):
        doc.add_heading(re.sub(r"[*`]", "", ln[5:]), level=4)
    elif re.match(r"^\s*[-*] ", ln):
        p = doc.add_paragraph(style="List Bullet")
        add_md_text(p, re.sub(r"^\s*[-*] ", "", ln))
    elif re.match(r"^\s*\d+\. ", ln):
        p = doc.add_paragraph(style="List Number")
        add_md_text(p, re.sub(r"^\s*\d+\. ", "", ln))
    elif ln.strip() == "---":
        doc.add_page_break()
    elif ln.strip():
        p = doc.add_paragraph()
        add_md_text(p, ln.strip())
    i += 1

docx_path = REP / "FIRM_SYSTEM_BLUEPRINT_20260713.docx"
doc.save(docx_path)
print("docx saved:", docx_path.name)
