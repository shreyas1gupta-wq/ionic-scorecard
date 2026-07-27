# -*- coding: utf-8 -*-
"""transfer_in_review.py — NDPMS transfer-in portfolio review (Ionic Wealth).

Checks an incoming client CAS / Demat / CAMS / Kfintech statement against the
transfer-in acceptance checklist and produces a Transfer Report: an immediate
pre-approved list (Stocks / MF & SIF / Bonds & InvIT) plus an exceptions list
routed to the right reviewer.

Checklist implemented (as specified by RM, 2026-07-27):
  Stocks       -> beyond NIFTY 500?              exception routed to Yajash
  Debt & InvIT -> listed + credit rating > A+?    exception routed to Girdhar
  MF & SIF     -> demat-mode holding, Direct plan? exception noted for RM
  Every stock and every demat-mode MF/SIF row also carries an
  RM_purchase_history_confirmed column (blank) — RM must tick this before any
  transfer, pass or fail on the other checks (this cannot be automated: no
  purchase-history data source exists in this repo).

INPUT
  Preferred: a pre-extracted CSV/XLSX with columns (case-insensitive, extras OK)
    type/category : EQ | MF | SIF | DEBT | INVIT   (optional — inferred if absent)
    isin           : strongly preferred, exact matching
    name           : security / scheme name as printed on the statement
    units          : quantity / units held
    value_inr      : current value in INR
    rating         : credit-rating string if visible on the statement (debt/InvIT)
    plan           : Direct | Regular, if visible (MF/SIF)
    family_member  : whose account this line came from (carried through, not scored)

  Also accepted: a raw CAS/CAMS/Kfintech PDF. Extraction is HEURISTIC (line/ISIN
  scanning via PyMuPDF — no fine-grained table structure, because no sample
  statement has been used to validate column positions yet). Every PDF-extracted
  row is tagged extraction_confidence=low and MUST be spot-checked against the
  source statement page before the report is treated as final — this script
  never silently upgrades a guess to a fact (D-035 epistemic conduct).

USAGE
  python transfer_in_review.py --statement <file.pdf|csv|xlsx> --client "Name" --out <dir>
"""
from __future__ import annotations

import os
import re
import csv
import sys
import json
import argparse
import datetime as dt

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, "..", "..", ".."))  # NIFTY 500 repo root
NIFTY500_JSON = os.path.join(ROOT, "datasets", "nifty500_current_2026.json")
DEFAULT_OUT = os.path.join(HERE, "..", "reports", "transfer_reviews")

sys.path.insert(0, HERE)
from docx_style_kit import (  # noqa: E402
    apply_firm_styles, add_title_page, add_heading, add_source_table, FIRM_RUST, FIRM_NAVY,
)
from docx import Document  # noqa: E402
from docx.shared import RGBColor, Pt  # noqa: E402

# ---------------------------------------------------------------------------
# Reference data / constants
# ---------------------------------------------------------------------------
ISIN_RE = re.compile(r"\b[A-Z]{2}[A-Z0-9]{9}[0-9]\b")

# long-term ratings strictly better than "A+" (AA- and above), plus the
# short-term top-tier equivalent (A1+) commonly seen on InvIT/NCD/CP paper.
# Anything NOT in this set — including A+, A, A-, BBB.., unrated, or an
# unrecognized string — does not auto-pass; conservative by design.
RATING_PASS_SET = {"AAA", "AA+", "AA", "AA-", "A1+"}

# ordered longest-token-first so alternation doesn't short-match e.g. "AA" inside "AA+"
_RATING_TOKENS = ["AAA", "AA+", "AA-", "AA", "A1+", "A1", "BBB+", "BBB-", "BBB",
                  "BB+", "BB-", "BB", "A+", "A-", "A", "B+", "B-", "B", "C", "D"]
RATING_TOKEN_RE = re.compile(r"\b(" + "|".join(re.escape(t) for t in _RATING_TOKENS) + r")\b")

SECTION_KEYWORDS = [
    (re.compile(r"MUTUAL\s+FUND", re.I), "MF"),
    (re.compile(r"\bSIF\b|SPECIALIZED\s+INVESTMENT\s+FUND", re.I), "SIF"),
    (re.compile(r"GOVERNMENT\s+SECURIT|BOND|DEBENTURE|\bNCD\b|\bSDL\b|G-?SEC", re.I), "DEBT"),
    (re.compile(r"\bINVIT\b|\bREIT\b", re.I), "INVIT"),
    (re.compile(r"EQUIT|\bSHARES?\b", re.I), "EQ"),
]
NAME_KEYWORDS = SECTION_KEYWORDS  # same patterns, applied to the row's own name text as fallback

NUM_RE = re.compile(r"-?\d[\d,]*\.?\d*")


def _norm(s):
    return re.sub(r"[^a-z0-9]", "", str(s or "").lower())


def _to_float(s):
    try:
        return float(str(s).replace(",", "").strip())
    except Exception:
        return None


# ---------------------------------------------------------------------------
# Reference: current NIFTY 500 membership
# ---------------------------------------------------------------------------
def load_nifty500(path=NIFTY500_JSON):
    by_isin, by_name = {}, {}
    with open(path, encoding="utf-8") as f:
        data = json.load(f)
    for r in data:
        isin = str(r.get("isin") or "").strip().upper()
        if isin:
            by_isin[isin] = r
        by_name[_norm(r.get("company") or r.get("symbol"))] = r
    return by_isin, by_name


def _match_nifty500(isin, name, by_isin, by_name):
    isin = str(isin or "").strip().upper()
    if isin and isin in by_isin:
        return by_isin[isin]
    key = _norm(name)
    if key and key in by_name:
        return by_name[key]
    return None


# ---------------------------------------------------------------------------
# Input: pre-extracted CSV / XLSX (preferred path)
# ---------------------------------------------------------------------------
def load_pre_extracted(path):
    rows = []
    if path.lower().endswith((".xlsx", ".xls")):
        import openpyxl
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        ws = wb[wb.sheetnames[0]]
        head = None
        for r in ws.iter_rows(values_only=True):
            if head is None:
                head = [_norm(c) for c in r]
                continue
            if all(c is None for c in r):
                continue
            rows.append({h: v for h, v in zip(head, r)})
    else:
        with open(path, encoding="utf-8-sig", newline="") as f:
            for r in csv.DictReader(f):
                rows.append({_norm(k): v for k, v in r.items()})
    out = []
    for r in rows:
        out.append({
            "category_hint": (r.get("type") or r.get("category") or "").strip().upper(),
            "isin": (r.get("isin") or "").strip().upper(),
            "name": r.get("name") or "",
            "units": _to_float(r.get("units") or r.get("qty") or r.get("quantity")),
            "value_inr": _to_float(r.get("valueinr") or r.get("value") or r.get("marketvalue")),
            "rating_raw": r.get("rating") or "",
            "plan_raw": r.get("plan") or "",
            "listed_raw": r.get("listed") or "",
            "family_member": r.get("familymember") or r.get("family") or "",
            "extraction_confidence": "high",  # human/ops-typed input
            "source_line": "",
        })
    return out


# ---------------------------------------------------------------------------
# Input: raw CAS/CAMS/Kfintech PDF (heuristic — see module docstring)
# ---------------------------------------------------------------------------
def extract_from_pdf(path):
    import fitz  # PyMuPDF
    doc = fitz.open(path)
    out = []
    current_section = None
    for page_no, page in enumerate(doc, 1):
        text = page.get_text("text")
        for line in text.splitlines():
            line = line.strip()
            if not line:
                continue
            for pat, sect in SECTION_KEYWORDS:
                if pat.search(line):
                    current_section = sect
                    break
            m = ISIN_RE.search(line)
            if not m:
                continue
            isin = m.group()
            name_guess = line[:m.start()].strip(" -:|") or line
            nums = [_to_float(n) for n in NUM_RE.findall(line)]
            nums = [n for n in nums if n is not None]
            units_guess = nums[-2] if len(nums) >= 2 else None
            value_guess = nums[-1] if nums else None
            rating_m = RATING_TOKEN_RE.search(line)
            out.append({
                "category_hint": current_section or "",
                "isin": isin,
                "name": name_guess,
                "units": units_guess,
                "value_inr": value_guess,
                "rating_raw": rating_m.group() if rating_m else "",
                "plan_raw": "DIRECT" if re.search(r"\bDIRECT\b", line, re.I)
                            else ("REGULAR" if re.search(r"\bREGULAR\b", line, re.I) else ""),
                "listed_raw": "LISTED" if re.search(r"\bLISTED\b|\bNSE\b|\bBSE\b", line, re.I) else "",
                "family_member": "",
                "extraction_confidence": "low",
                "source_line": f"p{page_no}: {line}",
            })
    doc.close()
    return out


# ---------------------------------------------------------------------------
# Category inference + checks
# ---------------------------------------------------------------------------
CATEGORY_ALIASES = {"EQ": "EQ", "EQUITY": "EQ", "STOCK": "EQ", "SHARE": "EQ",
                    "MF": "MF", "MUTUALFUND": "MF",
                    "SIF": "SIF",
                    "DEBT": "DEBT", "BOND": "DEBT", "NCD": "DEBT", "GSEC": "DEBT",
                    "INVIT": "INVIT", "REIT": "INVIT"}


def guess_category(row):
    hint = CATEGORY_ALIASES.get(str(row["category_hint"]).strip().upper(), "")
    if hint:
        return hint
    for pat, sect in NAME_KEYWORDS:
        if pat.search(row["name"] or ""):
            return sect
    return "UNKNOWN"


def check_row(row, by_isin, by_name):
    """Returns (status, routed_to, reason) for one holding row."""
    cat = row["category"]
    if cat == "EQ":
        hit = _match_nifty500(row["isin"], row["name"], by_isin, by_name)
        if hit:
            return "Pre-Approved", None, "in current NIFTY 500"
        return "Exception", "Yajash", "beyond NIFTY 500 — review and confirm"

    if cat in ("DEBT", "INVIT"):
        rating = row["rating_raw"].strip().upper()
        listed = bool(row["listed_raw"])
        if rating in RATING_PASS_SET and listed:
            return "Pre-Approved", None, f"rating {rating}, listed"
        reasons = []
        if rating not in RATING_PASS_SET:
            reasons.append(f"rating '{rating or 'not found on statement'}' not confirmed > A+")
        if not listed:
            reasons.append("listed status not confirmed from statement")
        return "Exception", "Girdhar", "; ".join(reasons)

    if cat in ("MF", "SIF"):
        plan = row["plan_raw"].strip().upper()
        if plan == "DIRECT":
            return "Pre-Approved", None, "Direct plan, demat/ISIN holding"
        return "Exception", "RM", (f"plan is '{plan or 'not stated'}', not confirmed Direct — "
                                    f"[INFERENCE: RM to check, no direct/regular ISIN master in repo]")

    return "Exception", "RM", "category could not be determined from statement — RM to classify"


REQUIRES_PURCHASE_HISTORY = {"EQ", "MF", "SIF"}  # "stocks and FM [=MF/SIF] in demat" per RM's checklist


# ---------------------------------------------------------------------------
# Report generation
# ---------------------------------------------------------------------------
def build_report_rows(raw_rows, by_isin, by_name):
    out = []
    for r in raw_rows:
        cat = guess_category(r)
        r["category"] = cat
        status, routed_to, reason = check_row(r, by_isin, by_name)
        out.append({
            **r,
            "status": status,
            "routed_to": routed_to or "",
            "reason": reason,
            "rm_purchase_history_confirmed": "" if cat in REQUIRES_PURCHASE_HISTORY else "n/a",
        })
    return out


CSV_COLS = ["category", "isin", "name", "units", "value_inr", "rating_raw", "plan_raw",
            "listed_raw", "family_member", "status", "routed_to", "reason",
            "rm_purchase_history_confirmed", "extraction_confidence", "source_line"]


def write_csv(rows, out_path):
    with open(out_path, "w", encoding="utf-8", newline="") as f:
        w = csv.DictWriter(f, fieldnames=CSV_COLS)
        w.writeheader()
        for r in rows:
            w.writerow({k: r.get(k, "") for k in CSV_COLS})


def write_docx(rows, out_path, client_name, source_path, as_of):
    approved = [r for r in rows if r["status"] == "Pre-Approved"]
    exceptions = [r for r in rows if r["status"] != "Pre-Approved"]
    low_conf = any(r["extraction_confidence"] == "low" for r in rows)

    doc = Document()
    apply_firm_styles(doc)
    add_title_page(doc, "Transfer-In Review — Transfer Report",
                    f"Client: {client_name}", date_str=as_of,
                    author="Shreyas_Ionic_AMC / Ionic Wealth — NDPMS transfer-in desk")

    if low_conf:
        p = doc.add_paragraph()
        run = p.add_run(
            "CAUTION: this statement was parsed from a raw PDF via heuristic line/ISIN "
            "scanning (no verified table layout for this statement format). Every row below "
            "marked extraction_confidence=low must be checked against the source pages before "
            "this report is treated as final."
        )
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(FIRM_RUST)
        run.font.size = Pt(10)

    for cat, label in [("EQ", "Stocks"), ("MF", "Mutual Funds"), ("SIF", "SIF"),
                        ("DEBT", "Bonds / Debentures"), ("INVIT", "InvIT / REIT")]:
        cat_rows = [r for r in approved if r["category"] == cat]
        if not cat_rows:
            continue
        add_heading(doc, f"Pre-Approved — {label}", level=2)
        add_source_table(
            doc, f"A-{cat}", f"{label} clearing the automated checklist (pending RM purchase-history sign-off)",
            ["Name", "ISIN", "Value (INR)", "Note"],
            [[r["name"], r["isin"], f"{r['value_inr']:,.0f}" if r["value_inr"] else "", r["reason"]]
             for r in cat_rows],
            source=os.path.basename(source_path), as_of=as_of, right_align_cols={2},
        )

    if exceptions:
        add_heading(doc, "Exceptions — routed for review", level=2)
        add_source_table(
            doc, "B", "Holdings requiring named sign-off before transfer",
            ["Name", "ISIN", "Category", "Value (INR)", "Routed To", "Reason"],
            [[r["name"], r["isin"], r["category"], f"{r['value_inr']:,.0f}" if r["value_inr"] else "",
              r["routed_to"], r["reason"]] for r in exceptions],
            source=os.path.basename(source_path), as_of=as_of, right_align_cols={3},
        )

    add_heading(doc, "Action items", level=2)
    for owner, note in [
        ("Yajash", "Review every Stocks exception above (beyond NIFTY 500) and confirm in writing."),
        ("Girdhar", "Confirm listed status + credit rating for every Debt/InvIT exception above."),
        ("RM", "Confirm purchase history for every Stock and every demat-mode MF/SIF row "
               "(pre-approved AND exceptions) before any transfer proceeds — see "
               "rm_purchase_history_confirmed column in transfer_report.csv."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{owner}: ")
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(FIRM_NAVY)
        p.add_run(note)

    doc.save(out_path)


def run(statement_path, client_name, out_dir, as_of=None):
    as_of = as_of or dt.date.today().isoformat()
    os.makedirs(out_dir, exist_ok=True)

    if statement_path.lower().endswith(".pdf"):
        raw_rows = extract_from_pdf(statement_path)
    else:
        raw_rows = load_pre_extracted(statement_path)

    by_isin, by_name = load_nifty500()
    rows = build_report_rows(raw_rows, by_isin, by_name)

    csv_path = os.path.join(out_dir, "transfer_report.csv")
    docx_path = os.path.join(out_dir, f"Transfer_Report_{_norm(client_name) or 'client'}.docx")
    write_csv(rows, csv_path)
    write_docx(rows, docx_path, client_name, statement_path, as_of)

    summary = {
        "client": client_name, "as_of": as_of, "source": statement_path,
        "n_rows": len(rows),
        "n_pre_approved": sum(1 for r in rows if r["status"] == "Pre-Approved"),
        "n_exceptions": sum(1 for r in rows if r["status"] != "Pre-Approved"),
        "routed": {name: sum(1 for r in rows if r["routed_to"] == name) for name in ("Yajash", "Girdhar", "RM")},
        "low_confidence_rows": sum(1 for r in rows if r["extraction_confidence"] == "low"),
    }
    json.dump(summary, open(os.path.join(out_dir, "summary.json"), "w", encoding="utf-8"), indent=2)

    print(f"transfer_in_review: {summary['n_pre_approved']} pre-approved, "
          f"{summary['n_exceptions']} exceptions -> {csv_path}")
    if summary["low_confidence_rows"]:
        print(f"WARNING: {summary['low_confidence_rows']} rows extracted from PDF at low confidence — "
              f"verify against source pages before treating this report as final.")
    print(f"docx -> {docx_path}")
    return summary


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--statement", required=True, help="CAS/CAMS/Kfintech PDF, or a pre-extracted CSV/XLSX")
    ap.add_argument("--client", required=True)
    ap.add_argument("--out", default="")
    a = ap.parse_args()
    out_dir = a.out or os.path.join(DEFAULT_OUT, f"{_norm(a.client)}_{dt.date.today().isoformat()}")
    run(a.statement, a.client, out_dir)
