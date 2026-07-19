"""
build_portfolio_recommendations.py — Step-4 deliverables for the NDPMS
portfolio-holdings review (STOCK_SCORECARD_750, 2026-07-18).

Merges the quant layer (results/portfolio_quant.csv, 59 holdings) with the
qualitative analyst layer (results/pf_qual_<SYMBOL>.json, 59 files) into:

  1. STOCK_SCORECARD_750/results/PORTFOLIO_RECOMMENDATIONS.xlsx
     Sheets: Summary | Analyst Notes | Escalations | Methodology
  2. 09_PRODUCT/reports/PORTFOLIO_HOLDINGS_REVIEW_2026-07-18.docx
     Principal-facing memo per STYLE_GUIDE.md (docx_style_kit).

Recommendations are Sell/Hold ONLY (existing-holdings review). Research/paper
only — no live action without Principal sign-off (firm hard rule).
"""
from __future__ import annotations

import glob
import json
import re
import sys
from pathlib import Path

import pandas as pd

SCRIPTS_DIR = Path(__file__).resolve().parent
FIRM_DIR = SCRIPTS_DIR.parent.parent
RESULTS_DIR = FIRM_DIR / "04_RND_LAB" / "STOCK_SCORECARD_750" / "results"
REPORTS_DIR = FIRM_DIR / "09_PRODUCT" / "reports"
AS_OF = "2026-07-18"

sys.path.insert(0, str(SCRIPTS_DIR))
from docx_style_kit import (  # noqa: E402
    FIRM_GOLD, FIRM_INK, FIRM_NAVY, FIRM_PALETTE_MPL, FIRM_RUST, FIRM_STONE,
    FIRM_TEAL, add_heading, add_source_table, add_title_page,
    apply_firm_styles, insert_chart, source_caption_mpl, style_chart_axes,
)


# ---------------------------------------------------------------------------
# Load + merge the two layers
# ---------------------------------------------------------------------------
def load_merged() -> pd.DataFrame:
    q = pd.read_csv(RESULTS_DIR / "portfolio_quant.csv")
    keep = ["symbol", "Company Name", "sector", "value_inr", "pe_current",
            "composite_3y", "composite_1y",
            "recommendation_3y", "recommendation_1y", "recommendation"]
    q = q[keep].rename(columns={
        "Company Name": "company", "recommendation": "quant_rec",
        "recommendation_3y": "quant_rec_3y", "recommendation_1y": "quant_rec_1y",
    })

    qual_rows = []
    for p in glob.glob(str(RESULTS_DIR / "pf_qual_*.json")):
        d = json.load(open(p, encoding="utf-8"))
        qual_rows.append({
            "symbol": d["symbol"],
            "analyst_rec": d["your_recommendation"],
            "growth_3y_pct": d["expected_next_3y_growth_pct"],
            "escalation": bool(d["escalation_flag"]),
            "escalation_reason": d.get("escalation_reason") or "",
            "summary": d.get("summary", ""),
            "positive_para": d.get("positive_para", ""),
            "negative_para": d.get("negative_para", ""),
            "reverse_dcf_judgment": d.get("reverse_dcf_judgment", ""),
            "recommendation_rationale": d.get("recommendation_rationale", ""),
        })
    qual = pd.DataFrame(qual_rows)

    df = q.merge(qual, on="symbol", how="outer", indicator=True)
    bad = df[df["_merge"] != "both"]
    if len(bad):
        raise SystemExit(f"MERGE MISMATCH — not 1:1: {bad[['symbol', '_merge']].to_dict('records')}")
    df = df.drop(columns="_merge").sort_values("value_inr", ascending=False).reset_index(drop=True)
    assert len(df) == 59, f"expected 59 holdings, got {len(df)}"
    df["pct_of_book"] = df["value_inr"] / df["value_inr"].sum() * 100
    df["override"] = [
        "" if a == b else f"{b}->{a}" for a, b in zip(df["analyst_rec"], df["quant_rec"])
    ]
    return df


def first_sentences(text: str, n: int = 2) -> str:
    """First n sentences of a paragraph (for the condensed docx escalation lines)."""
    parts = re.split(r"(?<=[.!?])\s+", text.strip())
    return " ".join(parts[:n]).strip()


# ---------------------------------------------------------------------------
# Excel
# ---------------------------------------------------------------------------
def build_excel(df: pd.DataFrame, xlsx_path: Path) -> None:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    NAVY, RUST, GOLD, STONE = "1F3A5D", "A34A28", "B08D57", "5F5E57"
    hdr_font = Font(name="Bahnschrift", bold=True, color="FFFFFF", size=10)
    hdr_fill = PatternFill("solid", fgColor=NAVY)
    body_font = Font(name="Georgia", size=9)
    sell_fill = PatternFill("solid", fgColor="F2DDD3")   # rust tint
    esc_fill = PatternFill("solid", fgColor="F3EAD9")    # gold tint
    wrap = Alignment(wrap_text=True, vertical="top")

    wb = Workbook()

    # ---- Sheet 1: Summary --------------------------------------------------
    ws = wb.active
    ws.title = "Summary"
    headers = ["#", "Symbol", "Company", "Sector", "Value (Rs)", "% of book",
               "Quant 3y", "Quant 1y", "Quant rec", "Analyst rec (FINAL)",
               "Est. 3y growth (%/yr)", "Escalation", "Override (quant->analyst)"]
    ws.append(headers)
    for j, _ in enumerate(headers, 1):
        c = ws.cell(row=1, column=j)
        c.font, c.fill = hdr_font, hdr_fill
    for i, r in df.iterrows():
        ws.append([i + 1, r.symbol, r.company, r.sector, round(r.value_inr),
                   round(r.pct_of_book, 2), round(r.composite_3y, 1),
                   round(r.composite_1y, 1), r.quant_rec, r.analyst_rec,
                   r.growth_3y_pct, "ESCALATED" if r.escalation else "",
                   r.override])
        row = i + 2
        for j in range(1, len(headers) + 1):
            ws.cell(row=row, column=j).font = body_font
        if r.analyst_rec == "Sell":
            for j in range(1, len(headers) + 1):
                ws.cell(row=row, column=j).fill = sell_fill
        elif r.escalation:
            ws.cell(row=row, column=12).fill = esc_fill
    widths = [4, 12, 34, 26, 12, 9, 9, 9, 10, 16, 16, 12, 20]
    for j, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(j)].width = w
    ws.freeze_panes = "C2"
    ws.auto_filter.ref = f"A1:{get_column_letter(len(headers))}{len(df) + 1}"

    # ---- Sheet 2: Analyst Notes ---------------------------------------------
    ws2 = wb.create_sheet("Analyst Notes")
    h2 = ["Symbol", "Rec", "Summary", "Positives", "Negatives",
          "Reverse-DCF judgment", "Recommendation rationale"]
    ws2.append(h2)
    for j, _ in enumerate(h2, 1):
        c = ws2.cell(row=1, column=j)
        c.font, c.fill = hdr_font, hdr_fill
    for i, r in df.iterrows():
        ws2.append([r.symbol, r.analyst_rec, r.summary, r.positive_para,
                    r.negative_para, r.reverse_dcf_judgment,
                    r.recommendation_rationale])
        for j in range(1, len(h2) + 1):
            c = ws2.cell(row=i + 2, column=j)
            c.font, c.alignment = body_font, wrap
        ws2.row_dimensions[i + 2].height = 150
    for j, w in enumerate([12, 7, 55, 55, 55, 55, 55], 1):
        ws2.column_dimensions[get_column_letter(j)].width = w
    ws2.freeze_panes = "B2"

    # ---- Sheet 3: Escalations -----------------------------------------------
    ws3 = wb.create_sheet("Escalations")
    esc = df[df.escalation].reset_index(drop=True)
    h3 = ["Symbol", "Analyst rec", "Quant rec", "Value (Rs)",
          "Escalation reason (verbatim, for Principal adjudication)"]
    ws3.append(h3)
    for j, _ in enumerate(h3, 1):
        c = ws3.cell(row=1, column=j)
        c.font, c.fill = hdr_font, hdr_fill
    for i, r in esc.iterrows():
        ws3.append([r.symbol, r.analyst_rec, r.quant_rec, round(r.value_inr),
                    r.escalation_reason])
        for j in range(1, len(h3) + 1):
            c = ws3.cell(row=i + 2, column=j)
            c.font, c.alignment = body_font, wrap
        ws3.row_dimensions[i + 2].height = 170
        if r.analyst_rec == "Sell":
            for j in range(1, len(h3) + 1):
                ws3.cell(row=i + 2, column=j).fill = sell_fill
    for j, w in enumerate([12, 10, 10, 12, 130], 1):
        ws3.column_dimensions[get_column_letter(j)].width = w
    ws3.freeze_panes = "B2"

    # ---- Sheet 4: Methodology -----------------------------------------------
    ws4 = wb.create_sheet("Methodology")
    lines = [
        ("SOURCE OF TRUTH", "04_RND_LAB/STOCK_SCORECARD_750/FROZEN_METHODOLOGY.md (v5, 2026-07-18) is the single source of truth for scoring methodology and BOTH output formats. If this sheet and that file ever disagree, FROZEN_METHODOLOGY.md wins. This workbook is the INTERNAL/analyst view; the FROZEN client-facing format (Stock Name | Ticker | ISIN | Sector | Recommendation | Rationale, where Rationale = the analyst summary field) is produced by 09_PRODUCT/scripts/build_client_excel.py -> 09_PRODUCT/reports/CLIENT_RECOMMENDATIONS.xlsx."),
        ("Scope", "All 59 NDPMS CAS holdings. Existing-holdings review: recommendation vocabulary Sell or Hold ONLY, never Buy. A missing score would yield 'No Recommendation' (never force a call on a data gap) — did not arise on these 59, all scored and analyst-reviewed. Research/paper only — no live action without Principal sign-off."),
        ("Quant layer", "results/portfolio_quant.csv — dual-horizon composite (3Y fundamentals-tilted 63/37, 1Y technical-tilted 40/60; DCF pillar excluded, weights renormalized) vs a 300-stock reference universe. Overlay gates: Balance-Sheet Safety (RED D/E>2.5 or IntCov<1.5 caps at 40; AMBER x0.85) and Liquidity (tiered turnover bar). Financial Services/Banking/NBFC/Insurance are EXEMPT from the D/E>2.5 trigger in BOTH the gate AND the red-flag penalty counter (leverage is the business model; these get bs_flag='N/A-financial-sector', with balance-sheet health deferred to the qualitative layer). Recommendation logic: per horizon, gate RED -> Sell, else score >=40 Hold / <40 Sell; Overall = Sell if EITHER horizon says Sell. Data is as of the scrape date (price lag vs 2026-07-18 expected)."),
        ("Qual layer", "One Sonnet sector-analyst agent per stock (personas: Rohan Deshmukh industrials, Meera Krishnan financials, Priya Nair consumer, Karan Malhotra IT/telecom, Dr. Sneha Patil pharma/chemicals). Single pass per stock: ~3min-equivalent deep research (business model, earnings-quality/one-off check, sector cycle, own reverse-DCF judgment, forward expected_next_3y_growth_pct) + the agent's own ~1min self-review. The analyst's your_recommendation OVERRIDES the quant-only call entirely where research exists. Internal PIT quarterly data (datasets/earnings_pit/unified_quarterly_pit.parquet, available_date discipline) used for cross-checks. Fundamental-only language; chart/technical language is reserved for the separate technical agent's fields."),
        ("Escalation rule", "Narrow, deliberately: escalation_flag=true ONLY for genuine analytical disagreement — a real Hold-vs-Sell judgment call, or a methodology gap likely affecting other stocks. Price/fundamentals scrape-date staleness is EXPECTED and NEVER escalation-worthy on its own (analysts silently corrected to current figures). Escalations are logged for Principal adjudication, not resolved by the agent."),
        ("Growth estimates", "expected_next_3y_growth_pct is each analyst's own forward estimate ([OPINION]-class, labeled in source files), not consensus. HINDUNILVR/IRCTC/ETERNAL fields patched prose->numeric (6.5/8.5/30); full prose preserved inside reverse_dcf_judgment in their JSONs."),
        ("VINTAGE NOTE (flagged, not re-solved)", "FROZEN_METHODOLOGY.md v4 added a red flag for analyst-estimated growth <10%. portfolio_quant.csv's redflag_count/penalty/final_adj columns predate the analyst layer, so this flag is NOT reflected in the quant columns of this workbook for the 12 names now under 10% (NATIONALUM, ITC, HINDUNILVR, ASIANPAINT, GAIL, INFY, TCS, IEX, TATASTEEL, IRCTC, COCHINSHIP, DEEPAKNTR). No FINAL recommendation changes (analyst override governs); awaiting ruling on whether to recompute the quant columns."),
        ("Known open gaps", "Per FROZEN_METHODOLOGY.md 'Known open gaps' (the full authoritative list, 7 items): ROE/ROCE lookback un-reverified; M&M-class captive-NBFC ROCE distortion; demerger PE-blending (SIEMENS/TMCV/TMPV/ENRIN/ITCHOTELS class); SUZLON-class DTA-credit PAT inflation; promoter pledge % missing; no NIFTY index PE/PB series (valuation-regime concept inactive); reverse-DCF has no automated pillar weight. All to Kavya (data) / Arjun (methodology) before the 750-universe rollout."),
        ("Sources", "Per-stock research files results/pf_qual_<SYMBOL>.json (each carries its own research_sources list with URLs). Full escalation texts: results/ESCALATIONS_FOR_PRINCIPAL.md."),
        ("Run details", f"DESK-100, {AS_OF}. 51 stocks this session in 10-parallel Sonnet batches (Principal-authorized), 8 prior (LT, BAJFINANCE, HINDUNILVR, SBIN, ABB, M&M, HINDALCO, BOSCHLTD). Sonnet-only pipeline (Opus dropped for cost)."),
    ]
    ws4.append(["Item", "Detail"])
    for j in (1, 2):
        c = ws4.cell(row=1, column=j)
        c.font, c.fill = hdr_font, hdr_fill
    for i, (k, v) in enumerate(lines):
        ws4.append([k, v])
        ws4.cell(row=i + 2, column=1).font = Font(name="Georgia", size=9, bold=True)
        c = ws4.cell(row=i + 2, column=2)
        c.font, c.alignment = body_font, wrap
        ws4.row_dimensions[i + 2].height = 62
    ws4.column_dimensions["A"].width = 22
    ws4.column_dimensions["B"].width = 130

    wb.save(xlsx_path)
    print("Excel written:", xlsx_path)


# ---------------------------------------------------------------------------
# Charts (matplotlib, firm style)
# ---------------------------------------------------------------------------
def build_charts(df: pd.DataFrame, out_dir: Path) -> dict:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    paths = {}
    lakh = 1e5

    # Chart 1 — where the book's value sits by verdict
    hold_clean = df[(df.analyst_rec == "Hold") & (~df.escalation)].value_inr.sum() / lakh
    hold_esc = df[(df.analyst_rec == "Hold") & (df.escalation)].value_inr.sum() / lakh
    sell = df[df.analyst_rec == "Sell"].value_inr.sum() / lakh
    fig, ax = plt.subplots(figsize=(7.6, 3.4))
    cats = ["Hold — clean", "Hold — escalated", "Sell"]
    vals = [hold_clean, hold_esc, sell]
    cols = [f"#{FIRM_TEAL}", f"#{FIRM_GOLD}", f"#{FIRM_RUST}"]
    bars = ax.barh(cats[::-1], vals[::-1], color=cols[::-1], height=0.55)
    for b, v in zip(bars, vals[::-1]):
        n = [ (df.analyst_rec == "Sell").sum(),
              ((df.analyst_rec == "Hold") & df.escalation).sum(),
              ((df.analyst_rec == "Hold") & ~df.escalation).sum() ][list(bars).index(b)]
        ax.text(b.get_width() + max(vals) * 0.01, b.get_y() + b.get_height() / 2,
                f"Rs {v:,.1f} L · {n} stocks", va="center", fontsize=9, color=f"#{FIRM_INK}")
    ax.set_xlabel("Holding value (Rs lakh)")
    ax.set_xlim(0, max(vals) * 1.28)
    style_chart_axes(ax)
    source_caption_mpl(fig, "PORTFOLIO_QUAL_SUMMARY.csv (59 holdings)", AS_OF)
    fig.tight_layout(rect=(0, 0.06, 1, 1))
    p = out_dir / "chart1_book_by_verdict.png"
    fig.savefig(p, dpi=200, bbox_inches="tight"); plt.close(fig)
    paths["c1"] = p

    # Chart 2 — top 15 positions by value, colored by verdict
    top = df.head(15).iloc[::-1]
    fig, ax = plt.subplots(figsize=(7.6, 5.4))
    col = [f"#{FIRM_RUST}" if r == "Sell" else (f"#{FIRM_GOLD}" if e else f"#{FIRM_TEAL}")
           for r, e in zip(top.analyst_rec, top.escalation)]
    ax.barh(top.symbol, top.value_inr / lakh, color=col, height=0.62)
    for y, (v, r) in enumerate(zip(top.value_inr / lakh, top.analyst_rec)):
        ax.text(v + top.value_inr.max() / lakh * 0.012, y, r, va="center",
                fontsize=8, color=f"#{FIRM_INK}")
    ax.set_xlabel("Holding value (Rs lakh)")
    ax.set_xlim(0, top.value_inr.max() / lakh * 1.18)
    style_chart_axes(ax)
    from matplotlib.patches import Patch
    ax.legend(handles=[Patch(color=f"#{FIRM_TEAL}", label="Hold — clean"),
                       Patch(color=f"#{FIRM_GOLD}", label="Hold — escalated"),
                       Patch(color=f"#{FIRM_RUST}", label="Sell")],
              loc="lower right", frameon=False, fontsize=8)
    source_caption_mpl(fig, "PORTFOLIO_QUAL_SUMMARY.csv (top 15 of 59 by value)", AS_OF)
    fig.tight_layout(rect=(0, 0.06, 1, 1))
    p = out_dir / "chart2_top15_by_verdict.png"
    fig.savefig(p, dpi=200, bbox_inches="tight"); plt.close(fig)
    paths["c2"] = p
    return paths


# ---------------------------------------------------------------------------
# Word report
# ---------------------------------------------------------------------------
def build_docx(df: pd.DataFrame, charts: dict, docx_path: Path) -> None:
    from docx import Document

    doc = Document()
    apply_firm_styles(doc)
    add_title_page(
        doc,
        "Portfolio Holdings Review — All 59 NDPMS Holdings",
        "Existing-holdings verdicts (Sell/Hold only) · quant layer + sector-analyst qualitative layer",
        date_str=AS_OF,
        author="Research Desk (5 sector analysts) · compiled DESK-100",
    )
    p = doc.add_paragraph()
    r = p.add_run("Research/paper output only. No live action is authorized by this document; "
                  "all Sells and all 32 escalations await Principal adjudication (firm hard rule: no real-money trades, ever). "
                  "Methodology per FROZEN_METHODOLOGY.md (v5, 2026-07-18) — the single source of truth; this memo is the internal view, "
                  "the frozen client-facing sheet is CLIENT_RECOMMENDATIONS.xlsx (build_client_excel.py).")
    r.italic = True

    # ---- 1. Verdict ----------------------------------------------------------
    add_heading(doc, "1. Verdict at a glance", 1)
    total_l = df.value_inr.sum() / 1e5
    sell_df = df[df.analyst_rec == "Sell"]
    doc.add_paragraph(
        f"Across the full book of 59 holdings (Rs {total_l:,.1f} lakh), the analyst layer lands at "
        f"48 Hold and 11 Sell. The Sells are concentrated in the tail of the book: they represent "
        f"Rs {sell_df.value_inr.sum() / 1e5:,.1f} lakh ({sell_df.value_inr.sum() / df.value_inr.sum() * 100:.1f}% of value) "
        f"but 11 of 59 names. 32 holdings carry an escalation flag — a genuine judgment call or a "
        f"quant-invisible development the desk believes the Principal should rule on personally. "
        f"The desk overrode the quant layer on 17 names: 11 quant-Sells rescued to Hold where the weakness "
        f"was one-off, cyclical, or already de-rated (LT, HINDUNILVR, RELIANCE, VBL, BAJAJHFL, ITC, ULTRACEMCO, "
        f"PPLPHARMA, GAIL, CMSINFO, ITCHOTELS), and 6 quant-Holds downgraded to Sell where valuation had "
        f"detached from the business (POWERINDIA, ASIANPAINT, POONAWALLA, BHEL, TATATECH, ANANDRATHI)."
    )
    insert_chart(doc, charts["c1"], 1,
                 "Book value by verdict — most value is in clean Holds; the escalated sleeve is where Principal attention buys the most.",
                 "PORTFOLIO_QUAL_SUMMARY.csv", AS_OF)
    insert_chart(doc, charts["c2"], 2,
                 "Top 15 positions by value, colored by verdict. No Sell sits inside the top 15; the largest Sell (POWERINDIA, Rs 4.4 L) ranks #19 by value.",
                 "PORTFOLIO_QUAL_SUMMARY.csv", AS_OF)

    # ---- 2. The 11 Sells ------------------------------------------------------
    add_heading(doc, "2. The 11 Sell recommendations", 1)
    rows = []
    reasons = {
        "TATAPOWER": "Rising leverage plus an unprovisioned ~$490mn arbitration award; close call, escalated.",
        "POWERINDIA": "~150x trailing PE fails reverse-DCF; FY26 order momentum is one HVDC mega-order base effect.",
        "JIOFIN": "1.1x book prices a ~14% normalized ROE the business (~1-2% actual) is nowhere near.",
        "DEEPAKNTR": "~40x on a depressed, volatile base while Haldia enters phenol/acetone at comparable scale.",
        "ASIANPAINT": "Structural tier-2/3 share loss to Birla Opus/JSW Dulux; ~59x multiple never de-rated; CCI probe open.",
        "POONAWALLA": "3.7x book prices ~21-26% steady-state ROE, above management's own Jun-2028 target.",
        "BHEL": "~60x corrected PE / 5.8x book on a best-ever 6% ROE after a +105% re-rating; margins already backsliding.",
        "COCHINSHIP": "Core-ops PBT a fraction of reported (subsidy accrual); richest shipyard multiple on the lowest ROE.",
        "HINDCOPPER": "FY26 growth is copper-price pass-through, not execution; expansion has verified equipment delays.",
        "TATATECH": "57x for a two-year organic stall; JLR concentration crashed a quarter's profit 96% this fiscal.",
        "ANANDRATHI": "Pristine franchise, but ~85-90x demands a decade of flawless 25-30% compounding management itself no longer guides.",
    }
    for _, r_ in sell_df.iterrows():
        rows.append([r_.symbol, f"{r_.value_inr / 1e5:,.2f}", r_.quant_rec,
                     f"{r_.growth_3y_pct}", reasons.get(r_.symbol, first_sentences(r_.summary, 1))])
    add_source_table(doc, 3, "Sell recommendations, by holding value.",
                     ["Symbol", "Value (Rs lakh)", "Quant said", "Est. 3y growth (%/yr)", "Core reason"],
                     rows, "results/pf_qual_<SYMBOL>.json (11 files)", AS_OF,
                     right_align_cols={1, 3}, col_widths_in=[0.9, 0.9, 0.8, 0.9, 3.6])

    # ---- 3. Full book ----------------------------------------------------------
    add_heading(doc, "3. All 59 holdings", 1)
    rows = []
    for _, r_ in df.iterrows():
        rows.append([r_.symbol, f"{r_.value_inr / 1e5:,.2f}", f"{r_.pct_of_book:.1f}",
                     r_.quant_rec, r_.analyst_rec, f"{r_.growth_3y_pct}",
                     "YES" if r_.escalation else ""])
    add_source_table(doc, 4, "Full book, by holding value. Analyst rec is the final verdict; growth is the analyst's own 3-year estimate ([OPINION]-class).",
                     ["Symbol", "Value (Rs lakh)", "% book", "Quant", "Analyst (FINAL)", "3y gr (%/yr)", "Escalated"],
                     rows, "PORTFOLIO_QUAL_SUMMARY.csv (59 rows)", AS_OF,
                     right_align_cols={1, 2, 5}, col_widths_in=[1.1, 1.0, 0.7, 0.8, 1.1, 0.9, 0.9])

    # ---- 4. Escalations ---------------------------------------------------------
    add_heading(doc, "4. The 32 escalations — for Principal adjudication", 1)
    doc.add_paragraph(
        "Condensed to their core tension below; the desk has deliberately NOT resolved these. "
        "Full verbatim escalation texts (and each stock's complete research file with sources) are in the "
        "Escalations sheet of PORTFOLIO_RECOMMENDATIONS.xlsx and results/ESCALATIONS_FOR_PRINCIPAL.md."
    )
    themes = [
        ("4.1 Quant-invisible corporate actions", ["SUNPHARMA", "PERSISTENT", "TMCV"]),
        ("4.2 Sell-side judgment calls (all 10 escalated Sells)", [s for s in sell_df.symbol if df.set_index("symbol").loc[s, "escalation"]]),
        ("4.3 Hold — genuine coin-flips and watch-items", None),  # remainder
        ("4.4 Methodology gaps (fix before the 750-universe rollout)", ["SIEMENS", "SUZLON", "M&M"]),
    ]
    used = set()
    esc_df = df[df.escalation].set_index("symbol")
    for title, syms in themes:
        add_heading(doc, title, 2)
        if syms is None:
            syms = [s for s in esc_df.index if s not in used]
        for s in syms:
            if s in used or s not in esc_df.index:
                continue
            used.add(s)
            r_ = esc_df.loc[s]
            p = doc.add_paragraph()
            run = p.add_run(f"{s} ({r_.analyst_rec}; quant {r_.quant_rec}; Rs {r_.value_inr / 1e5:,.1f} L) — ")
            run.bold = True
            p.add_run(first_sentences(r_.escalation_reason, 3))

    # ---- 5. Near-term calendar ----------------------------------------------------
    add_heading(doc, "5. Knife-edge names with Q1 FY27 prints inside two weeks", 1)
    cal = [["BANDHANBNK", "21-Jul-2026", "Hold priced for guided ROE recovery; capital-raise overhang"],
           ["IDFCFIRSTB", "25-Jul-2026", "Management mid-teens ROE promise vs Street 9-12%"],
           ["SUMICHEM", "27-Jul-2026", "~55x, zero cushion, kharif sowing -20.8% YoY"],
           ["BEL", "27-Jul-2026", "~50x vs 2x own 10-yr median PE; margin print is the test"],
           ["BAJAJHFL", "29/30-Jul-2026", "3.2x book for 12.1% ROE; coin-flip Hold"],
           ["MARUTI", "31-Jul-2026", "Two-year PBT margin compression is the open question"],
           ["ITC", "~end-Jul-2026", "Tax-shock volume hit; print could flip the Hold"],
           ["VBL", "~end-Jul-2026", "Peak-summer quarter vs monsoon downgrade + Campa"],
           ["TMPV", "Q2 (JLR print)", "JLR margin guidance already cut; next print could tip to Sell"]]
    add_source_table(doc, 5, "If any escalated call is to be re-checked, doing it after these prints is the efficient order.",
                     ["Symbol", "Print date", "What the print decides"],
                     cal, "per-stock research files (results/pf_qual_<SYMBOL>.json)", AS_OF,
                     col_widths_in=[1.1, 1.2, 4.2])

    # ---- 6. Next steps --------------------------------------------------------------
    add_heading(doc, "6. Decisions requested / next steps", 1)
    for t in [
        "Adjudicate the 32 escalations (Sec. 4) — in particular the six quant-Hold->Sell downgrades and the three pending-M&A balance-sheet shifts (SUNPHARMA/Organon, PERSISTENT/Nagarro, TMCV/Iveco).",
        "Rule on the quant-methodology open gaps (authoritative 7-item list in FROZEN_METHODOLOGY.md — headlined by demerger PE-blend, DTA-inflated PAT, captive-NBFC ROCE) before the full 750-universe scorecard build trusts those fields; plus the vintage question of recomputing quant penalty columns for the v4 analyst-growth<10% flag (12 names affected, no final call changes).",
        "No live action is taken from this review; any execution of Sells in the real NDPMS account is the Principal's own step outside this system.",
    ]:
        doc.add_paragraph(t, style="List Bullet")

    doc.save(docx_path)
    print("Docx written:", docx_path)


if __name__ == "__main__":
    df = load_merged()
    print(f"Merged {len(df)} holdings | Hold={ (df.analyst_rec=='Hold').sum() } Sell={ (df.analyst_rec=='Sell').sum() } Esc={ df.escalation.sum() }")
    xlsx = RESULTS_DIR / "PORTFOLIO_RECOMMENDATIONS.xlsx"
    build_excel(df, xlsx)
    charts = build_charts(df, RESULTS_DIR)
    docx_out = REPORTS_DIR / f"PORTFOLIO_HOLDINGS_REVIEW_{AS_OF}.docx"
    build_docx(df, charts, docx_out)
    print("DONE")
