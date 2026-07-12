"""
docx_style_kit.py — reusable python-docx + matplotlib helper module implementing
STYLE_GUIDE.md (Shreyas_Ionic_AMC de-AI-ification style system, WS-2).
Owner: Tanvi Desai (Product, E-026). Feeds every 09_PRODUCT/scripts/build_*.py builder.

Import this from any report builder instead of hand-setting Calibri/default colors:

    from docx_style_kit import (
        apply_firm_styles, add_title_page, add_heading, add_exhibit_caption,
        insert_chart, add_source_table, style_chart_axes, FIRM_PALETTE_MPL,
        FIRM_INK, FIRM_NAVY, FIRM_GOLD, FIRM_TEAL, FIRM_RUST, FIRM_STONE,
    )

Design source: STYLE_GUIDE.md §(b) Document design, §(c) Charts, §(d) Tables.
Fonts confirmed present on the build machine (C:\\Windows\\Fonts): georgia.ttf/-b/-i/-z,
bahnschrift.ttf — no install dependency.

NOTE on the existing dataviz palette (BLUE/AQUA/YELLOW/GREEN/VIOLET/RED used in
build_principal_report.py etc.): that is the generic dataviz-skill placeholder palette,
which the dataviz skill itself instructs teams to "swap for your own brand." The palette
below IS that swap for Principal-facing product deliverables (Investor Letter, strategy
packs, execution-sheet docs). [OPINION/house — Tanvi] New Principal-facing builders should
import this module; older chart code is not required to be retrofitted retroactively, but
any NEW build should use FIRM_PALETTE_MPL, not the old six.
"""
from __future__ import annotations

import datetime as dt
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# PALETTE — 6 hex colors, STYLE_GUIDE.md §(c). Stone matches the "5F5E57" gray
# already in use (build_principal_report.py:133) for continuity.
# ---------------------------------------------------------------------------
FIRM_INK   = "1C1C1A"   # primary text / axis / ink
FIRM_NAVY  = "1F3A5D"   # primary brand / primary series / headings
FIRM_GOLD  = "B08D57"   # accent / highlight / conviction-weighted
FIRM_TEAL  = "2E6E62"   # secondary series / positive-signed numbers
FIRM_RUST  = "A34A28"   # negative-signed / kills / losses / warnings
FIRM_STONE = "5F5E57"   # neutral / gridlines / captions / baseline

# matplotlib wants "#RRGGBB"
FIRM_PALETTE_MPL = [f"#{c}" for c in
                    (FIRM_NAVY, FIRM_GOLD, FIRM_TEAL, FIRM_RUST, FIRM_STONE, FIRM_INK)]

BANNED_OFFICE_BLUE = "4472C4"  # Word default Accent-1 -- never use this

BODY_FONT = "Georgia"
HEAD_FONT = "Bahnschrift"


def get_palette_color(i: int) -> str:
    """Cyclic accessor into the 6-color firm palette (mpl hex, with '#')."""
    return FIRM_PALETTE_MPL[i % len(FIRM_PALETTE_MPL)]


# ---------------------------------------------------------------------------
# DOCX: base styles
# ---------------------------------------------------------------------------
def apply_firm_styles(doc: Document) -> None:
    """Override Normal + Heading 1/2/3 to the firm typography pairing.
    Call this ONCE right after Document() / Document(template) -- before adding content."""
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(FIRM_INK)

    heading_sizes = {"Heading 1": (16, True), "Heading 2": (13, True), "Heading 3": (11, True)}
    for style_name, (size, bold) in heading_sizes.items():
        if style_name in doc.styles:
            st = doc.styles[style_name]
            st.font.name = HEAD_FONT
            st.font.size = Pt(size)
            st.font.bold = bold
            st.font.color.rgb = RGBColor.from_string(FIRM_NAVY)
            # kill any inherited theme color the built-in style carries
            rpr = st.element.get_or_add_rPr()
            for tag in ("color",):
                for el in rpr.findall(qn(f"w:{tag}")):
                    rpr.remove(el)


def add_heading(doc: Document, text: str, level: int = 1):
    """Wrapper around doc.add_heading that guarantees direct-formatted firm styling
    even if the underlying style object didn't take (some Word templates ignore
    style-level color on Heading styles)."""
    h = doc.add_heading("", level=level)
    run = h.add_run(text)
    run.font.name = HEAD_FONT
    run.font.color.rgb = RGBColor.from_string(FIRM_NAVY)
    run.font.bold = True
    return h


def add_title_page(doc: Document, title: str, subtitle: str, date_str: str | None = None,
                    classification: str = "Internal — Shreyas_Ionic_AMC",
                    author: str | None = None) -> None:
    """Title-page furniture per STYLE_GUIDE.md §(b): direct-formatted (never the
    theme-linked default Word 'Title' style, which carries the default blue)."""
    date_str = date_str or dt.date.today().isoformat()

    tpar = doc.add_paragraph()
    tpar.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tpar.add_run(title)
    tr.font.name = HEAD_FONT
    tr.font.size = Pt(26)
    tr.bold = True
    tr.font.color.rgb = RGBColor.from_string(FIRM_NAVY)

    spar = doc.add_paragraph()
    spar.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = spar.add_run(subtitle)
    sr.italic = True
    sr.font.name = BODY_FONT
    sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor.from_string(FIRM_STONE)

    meta_bits = [date_str, classification]
    if author:
        meta_bits.append(author)
    mpar = doc.add_paragraph()
    mpar.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = mpar.add_run(" · ".join(meta_bits))
    mr.font.name = BODY_FONT
    mr.font.size = Pt(9.5)
    mr.font.color.rgb = RGBColor.from_string(FIRM_STONE)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# DOCX: numbered exhibits with footnoted source + as-of date
# ---------------------------------------------------------------------------
def _add_left_rule(paragraph, color_hex: str = FIRM_NAVY, sz: int = 12) -> None:
    """Add a thin left border to a paragraph (the exhibit-caption rule)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(sz))
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)
    paragraph.paragraph_format.left_indent = Pt(8)


def add_exhibit_caption(doc: Document, number, text: str, source: str, as_of: str):
    """Exhibit N. <caption> + Source/as-of line, per STYLE_GUIDE.md §(b)/(c).
    number: int or str ('3' or '3a'). source: file path (+ row count if relevant).
    as_of: 'YYYY-MM-DD'."""
    cap = doc.add_paragraph()
    _add_left_rule(cap)
    r1 = cap.add_run(f"Exhibit {number}. ")
    r1.bold = True
    r1.font.name = BODY_FONT
    r1.font.size = Pt(9.5)
    r1.font.color.rgb = RGBColor.from_string(FIRM_INK)
    r2 = cap.add_run(text)
    r2.font.name = BODY_FONT
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = RGBColor.from_string(FIRM_INK)

    src = doc.add_paragraph()
    _add_left_rule(src)
    sr = src.add_run(f"Source: {source} · as of {as_of}")
    sr.italic = True
    sr.font.name = BODY_FONT
    sr.font.size = Pt(8)
    sr.font.color.rgb = RGBColor.from_string(FIRM_STONE)


def insert_chart(doc: Document, png_path, number, caption: str, source: str, as_of: str,
                  width_inches: float = 6.3):
    """Insert a chart image + its numbered exhibit caption in one call."""
    doc.add_picture(str(png_path), width=Inches(width_inches))
    last_par = doc.paragraphs[-1]
    last_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_exhibit_caption(doc, number, caption, source, as_of)


# ---------------------------------------------------------------------------
# DOCX: three-line tables (no vertical rules, units in header)
# ---------------------------------------------------------------------------
def _set_cell_border(cell, edge: str, sz: int = 8, color: str = FIRM_INK, val: str = "single"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    tag = f"w:{edge}"
    el = tcBorders.find(qn(tag))
    if el is None:
        el = OxmlElement(tag)
        tcBorders.append(el)
    el.set(qn("w:val"), val)
    el.set(qn("w:sz"), str(sz))
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), color)


def _clear_cell_border(cell, edge: str):
    _set_cell_border(cell, edge, sz=0, val="nil")


def add_source_table(doc: Document, number, caption: str, headers, rows,
                      source: str, as_of: str, right_align_cols=None,
                      col_widths_in=None):
    """Three-line table per STYLE_GUIDE.md §(d): rule above header, rule below
    header, rule at table bottom, NO vertical rules, NO inside grid.
    headers already carry units, e.g. 'CAGR (%)'. right_align_cols: set of column
    indices (0-based) to right-align (numeric columns)."""
    right_align_cols = right_align_cols or set()
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    # leave default (unnamed) table style -- it carries no visible borders; we add only
    # the specific top/header/bottom rules below (three-line table, no vertical rules)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    if col_widths_in:
        for j, w in enumerate(col_widths_in):
            for row in t.rows:
                row.cells[j].width = Inches(w)

    # header row: top rule + bottom rule; body: no rules except last-row bottom
    for j, htxt in enumerate(headers):
        cell = t.rows[0].cells[j]
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT if j in right_align_cols else WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(htxt)
        run.bold = True
        run.font.name = BODY_FONT
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(FIRM_INK)
        _set_cell_border(cell, "top", sz=10, color=FIRM_INK)
        _set_cell_border(cell, "bottom", sz=6, color=FIRM_INK)
        _clear_cell_border(cell, "left")
        _clear_cell_border(cell, "right")

    for i, row in enumerate(rows, 1):
        for j, v in enumerate(row):
            cell = t.rows[i].cells[j]
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT if j in right_align_cols else WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(str(v))
            run.font.name = BODY_FONT
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(FIRM_INK)
            _clear_cell_border(cell, "left")
            _clear_cell_border(cell, "right")
            is_last = (i == len(rows))
            _set_cell_border(cell, "bottom", sz=10 if is_last else 0,
                              color=FIRM_INK, val="single" if is_last else "nil")

    add_exhibit_caption(doc, number, caption, source, as_of)


# ---------------------------------------------------------------------------
# MATPLOTLIB: chart axes styling — no default-matplotlib look
# ---------------------------------------------------------------------------
def style_chart_axes(ax, y_grid: bool = True):
    """Apply STYLE_GUIDE.md §(c): spines off top/right, faint left/bottom,
    faint dotted y-gridlines only, firm ink/stone text colors, no default cycle."""
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_alpha(0.4)
    ax.spines["left"].set_color(f"#{FIRM_STONE}")
    ax.spines["bottom"].set_alpha(0.4)
    ax.spines["bottom"].set_color(f"#{FIRM_STONE}")
    ax.tick_params(colors=f"#{FIRM_STONE}", labelsize=8)
    ax.title.set_color(f"#{FIRM_INK}")
    ax.xaxis.label.set_color(f"#{FIRM_STONE}")
    ax.yaxis.label.set_color(f"#{FIRM_STONE}")
    if y_grid:
        ax.grid(axis="y", color=f"#{FIRM_STONE}", alpha=0.18, linewidth=0.6, linestyle=":")
        ax.grid(axis="x", visible=False)
    else:
        ax.grid(visible=False)
    ax.set_axisbelow(True)


def source_caption_mpl(fig, source: str, as_of: str, y: float = -0.02):
    """Small source/as-of caption directly on the figure (used when the PNG is
    embedded standalone, e.g. in HTML/PPT rather than a docx with add_exhibit_caption)."""
    fig.text(0.01, y, f"Source: {source} · as of {as_of}", fontsize=7.5,
              color=f"#{FIRM_STONE}", style="italic", ha="left")


# ---------------------------------------------------------------------------
# SELF-TEST: build a 2-page before/after sample docx
# ---------------------------------------------------------------------------
def _build_sample():
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    root = Path(__file__).resolve().parents[3]
    out_dir = root / "Shreyas_Ionic_AMC" / "09_PRODUCT" / "reports"
    img_dir = out_dir / "_img_style_kit_sample"
    img_dir.mkdir(parents=True, exist_ok=True)
    today = dt.date.today().isoformat()

    # ---- sample chart, firm palette + style_chart_axes ----
    fig, ax = plt.subplots(figsize=(6.3, 3.0), dpi=200)
    months = ["Mar", "Apr", "May", "Jun", "Jul"]
    strangle = [0.14, 0.19, 0.11, 0.22, 0.22]
    hurdle = [0.10] * 5
    ax.plot(months, strangle, color=get_palette_color(0), lw=1.8, marker="o", label="S-04 strangle")
    ax.plot(months, hurdle, color=get_palette_color(3), lw=1.2, ls=(0, (4, 2)), label="Cost hurdle")
    ax.annotate(f"{strangle[-1]:.2f}%/spot", (4, strangle[-1]), textcoords="offset points",
                xytext=(6, 4), fontsize=8, color=get_palette_color(0))
    ax.set_ylabel("% of spot, managed exit")
    ax.set_title("S-04 strangle: managed return vs 2×-cost hurdle")
    style_chart_axes(ax)
    ax.legend(frameon=False, fontsize=8, loc="lower right")
    fig.tight_layout()
    chart_path = img_dir / "sample_chart.png"
    fig.savefig(chart_path, bbox_inches="tight")
    plt.close(fig)

    doc = Document()
    apply_firm_styles(doc)

    add_title_page(
        doc,
        title="STYLE KIT — Before / After Sample",
        subtitle="Two pages: prior house style (Calibri, default theme, unfiltered prose) vs the "
                 "new firm style (Georgia/Bahnschrift, firm palette, style-lint-clean prose)",
        date_str=today,
        author="Tanvi Desai, Head of Product",
    )

    # ---------------- PAGE 1: BEFORE ----------------
    add_heading(doc, "1. BEFORE — prior house style", level=1)
    before = doc.add_paragraph()
    before.add_run(
        "In today's fast-paced landscape, it's important to note that our robust and "
        "comprehensive strangle sleeve continues to leverage a seamless, holistic approach. "
        "Let's dive in: this game-changer boasts significant, remarkable performance that "
        "underscores a watershed moment for the book. Studies show the results are truly "
        "impactful — it's not just a number, it's a testament to our meticulous process. "
        "The future looks bright, and this thriving, cutting-edge strategy will surely "
        "empower the team to harness, streamline, and elevate outcomes going forward."
    ).font.name = "Calibri"
    for run in before.runs:
        run.font.size = Pt(10)
    p2 = doc.add_paragraph()
    p2.add_run("Key highlights:").bold = True
    for item in ["Innovation", "Growth", "Scale", "Opportunity", "Momentum"]:
        b = doc.add_paragraph(style=None)
        b.add_run(f"- {item}").font.name = "Calibri"
    doc.add_paragraph()

    # a default-styled table (Word built-in grid look) to show the "before" table tell
    tbefore = doc.add_table(rows=3, cols=3)
    tbefore.style = "Light Grid Accent 1"  # the default-blue-grid Office look we are banning
    hdr = tbefore.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Month", "Return", "Hurdle"
    data_rows = [("Jun", "0.22%", "0.10%"), ("Jul", "0.22%", "0.10%")]
    for i, row in enumerate(data_rows, 1):
        for j, v in enumerate(row):
            tbefore.rows[i].cells[j].text = v
    doc.add_paragraph()

    doc.add_page_break()

    # ---------------- PAGE 2: AFTER ----------------
    add_heading(doc, "2. AFTER — firm style (STYLE_GUIDE.md-compliant)", level=1)
    after = doc.add_paragraph()
    after.add_run(
        "S-04 (the short strangle) is the one survivor of the four original option sleeves. "
        "It certified 12/12 cells at 2× modeled costs and moved to paper-watch on 2026-07-04. "
        "The managed-exit return held at +0.22% of spot in June and July, against a cost hurdle "
        "of 0.10% of spot — a real edge, not a rounding artifact."
    )
    after2 = doc.add_paragraph()
    after2.add_run(
        "I am not yet ready to call this durable. Two months is not a track record. Neel Basu's "
        "attribution has not run on it yet, and the paper-watch phase exists precisely to find out "
        "whether managed-exit fills are as clean live as they were in the backtest."
    )

    insert_chart(
        doc, chart_path, number=1,
        caption="S-04 strangle managed return vs 2×-cost hurdle, Mar-Jul 2026.",
        source="06_TRADING_DESK/STRATEGY_REGISTER.md (S-04 row)",
        as_of=today,
    )

    add_source_table(
        doc, number=2,
        caption="S-04 monthly managed-exit return vs cost hurdle.",
        headers=["Month", "Return (% spot)", "Hurdle (% spot)"],
        rows=[["Jun 2026", "0.22", "0.10"], ["Jul 2026", "0.22", "0.10"]],
        source="06_TRADING_DESK/STRATEGY_REGISTER.md (S-04 row)",
        as_of=today,
        right_align_cols={1, 2},
    )

    out_path = out_dir / "_style_sample.docx"
    doc.save(str(out_path))
    print(f"Sample written: {out_path}")
    return out_path


if __name__ == "__main__":
    _build_sample()
