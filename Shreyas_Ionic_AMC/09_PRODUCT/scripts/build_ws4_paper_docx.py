"""Render SYSTEM_VS_LLM_PAPER_DRAFT.md -> firm-styled docx with the 3 WS-4 charts
embedded at their anchor points. Reuses the firm's markdown table/heading parser
pattern (as in build_firm_blueprint.py) layered with docx_style_kit typography."""
import re, sys
from pathlib import Path

sys.path.insert(0, r"c:\Users\Shreyas.1Gupta\OneDrive - Angel Broking Limited\Desktop\Backup\NIFTY 500\Shreyas_Ionic_AMC\09_PRODUCT\scripts")
from docx_style_kit import (apply_firm_styles, add_title_page, add_heading, insert_chart,
                             FIRM_INK, FIRM_STONE)
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(r"c:\Users\Shreyas.1Gupta\OneDrive - Angel Broking Limited\Desktop\Backup\NIFTY 500")
REP = ROOT / "Shreyas_Ionic_AMC/09_PRODUCT/reports"
IMG = REP / "_img_ws4"
AS_OF = "2026-07-15"
md = (REP / "SYSTEM_VS_LLM_PAPER_DRAFT.md").read_text(encoding="utf-8")

# anchor markers: insert chart N right after this exact table's LAST row line
CHART_AFTER = {
    "| C (firm pipeline) | 1.80 |": ("chart3_primary_study_arms.png", "3",
        "Pre-registered study: defects found by arm (A/B/C/C2), Opus 4.8 base.",
        "ws4_battery/results/opus_arms_grade/OPUS_ARMS_RESULT.txt"),
    "| Haiku 4.5 | 9/16 | 1/4 | $0.025": ("chart1_cost_vs_accuracy.png", "1",
        "Cost vs. accuracy across four Claude tiers on the same battery.",
        "MODEL_GRID/COST_ESTIMATE.txt; ws4_battery/results/xmodel_grade/BATTERY_RESULT.txt"),
    "| Opus 4.8 | 9.25 | 9.75 | **+0.50**": ("chart2_judge_self_preference.png", "2",
        "Judge self-preference: same answers, two judges.",
        "MODEL_GRID/GRID_QUALITY_CORRECTED.txt"),
}

doc = Document()
apply_firm_styles(doc)
doc.styles["Normal"].font.size = Pt(10)

add_title_page(
    doc,
    title="Does Process Beat the Model?",
    subtitle="A Pre-Registered, Blind Test of Governed Multi-Agent Review Against a Single "
             "Frontier LLM on Backtest Defect Detection",
    date_str=AS_OF,
    author="Firm S — internal research note",
)
doc.add_page_break()


def add_md_run(p, text):
    for tok in re.split(r"(\*\*.+?\*\*|`.+?`|\*[^*]+?\*)", text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9)
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 1:
            r = p.add_run(tok[1:-1]); r.italic = True
        else:
            p.add_run(tok)


lines = md.splitlines()
i = 0
table_num = 0
while i < len(lines):
    ln = lines[i]
    if ln.startswith("|") and i + 1 < len(lines) and set(lines[i + 1].replace("|", "").replace(" ", "")) <= set("-:"):
        rows = []
        raw_block_lines = []
        while i < len(lines) and lines[i].startswith("|"):
            raw_block_lines.append(lines[i])
            cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
            if not set("".join(cells)) <= set("-: "):
                rows.append(cells)
            i += 1
        raw_block = "\n".join(raw_block_lines)
        ncol = max(len(r) for r in rows)
        table_num += 1
        t = doc.add_table(rows=len(rows), cols=ncol)
        t.style = "Light Grid Accent 1"
        for ri, r in enumerate(rows):
            for ci in range(ncol):
                cell = t.cell(ri, ci)
                cell.text = ""
                p = cell.paragraphs[0]
                add_md_run(p, r[ci] if ci < len(r) else "")
                for run in p.runs:
                    run.font.size = Pt(8)
                    run.font.name = "Georgia"
                    if ri == 0:
                        run.bold = True
        doc.add_paragraph()
        # check anchor (against RAW markdown block text, which still has the | | formatting
        # the CHART_AFTER keys were copied from -- matching against parsed/joined cells was
        # the bug in the first build: 0 images embedded despite no errors, caught on readback)
        for key, (png, num, cap, src) in CHART_AFTER.items():
            if key in raw_block:
                insert_chart(doc, IMG / png, number=num, caption=cap, source=src, as_of=AS_OF)
                doc.add_paragraph()
                print("  chart anchored:", png)
        continue
    if ln.startswith("# "):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_heading(doc, re.sub(r"[*`]", "", ln[2:]), level=1)
    elif ln.startswith("## "):
        add_heading(doc, re.sub(r"[*`]", "", ln[3:]), level=1)
    elif ln.startswith("### "):
        add_heading(doc, re.sub(r"[*`]", "", ln[4:]), level=2)
    elif re.match(r"^\s*\d+\. ", ln):
        p = doc.add_paragraph(style="List Number"); add_md_run(p, re.sub(r"^\s*\d+\. ", "", ln))
    elif re.match(r"^\s*[-*] ", ln):
        p = doc.add_paragraph(style="List Bullet"); add_md_run(p, re.sub(r"^\s*[-*] ", "", ln))
    elif ln.strip() == "---":
        pass  # section rule -> skip (title page + headings already separate sections visually)
    elif ln.strip().startswith("```"):
        # code/exhibit block: collect until closing fence, render as monospace box
        i += 1
        block = []
        while i < len(lines) and not lines[i].strip().startswith("```"):
            block.append(lines[i]); i += 1
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(14)
        run = p.add_run("\n".join(block))
        run.font.name = "Consolas"; run.font.size = Pt(8); run.font.color.rgb.__class__
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor.from_string(FIRM_STONE)
    elif ln.strip():
        p = doc.add_paragraph()
        add_md_run(p, ln.strip())
        for run in p.runs:
            run.font.size = Pt(10)
    i += 1

out_path = REP / "FIRM_S_SYSTEM_VS_LLM_20260715.docx"
doc.save(str(out_path))
print("paper docx saved:", out_path)
