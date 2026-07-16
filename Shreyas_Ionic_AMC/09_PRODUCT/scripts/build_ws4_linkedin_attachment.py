"""Build the LinkedIn-attachment docx: the reader-facing companion to LINKEDIN_POST_DRAFT.md.
Shorter than the full paper, same 'lead with clean wins' emphasis, neutral alias only,
no internal-only editorial notes (pivot history, pending-audit markers, Principal routing)."""
import sys
from pathlib import Path
sys.path.insert(0, r"c:\Users\Shreyas.1Gupta\OneDrive - Angel Broking Limited\Desktop\Backup\NIFTY 500\Shreyas_Ionic_AMC\09_PRODUCT\scripts")
from docx_style_kit import (apply_firm_styles, add_title_page, add_heading, insert_chart,
                             add_source_table)
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(r"c:\Users\Shreyas.1Gupta\OneDrive - Angel Broking Limited\Desktop\Backup\NIFTY 500")
REP = ROOT / "Shreyas_Ionic_AMC/09_PRODUCT/reports"
IMG = REP / "_img_ws4"
AS_OF = "2026-07-15"

doc = Document()
apply_firm_styles(doc)

add_title_page(
    doc,
    title="Cheap, Accurate, and Honest About Its Own Bias",
    subtitle="A benchmark of four AI models on real research-review work, and what it took "
             "to catch a grading bias hiding in the result",
    date_str=AS_OF,
    author="Firm S — a personal research benchmark",
)

add_heading(doc, "Why this exists", level=1)
doc.add_paragraph(
    "I run a personal research process, call it Firm S, that reviews quantitative research "
    "the way a real desk should: every test pre-registered before it runs, every claim "
    "checked adversarially, every result graded blind. I used that process to build a 20-task "
    "exam and ran it across four AI models to see what you actually get for the money."
)

add_heading(doc, "The exam", level=1)
doc.add_paragraph(
    "Twenty review tasks, built from real data traps: lookahead bugs, timestamp errors, "
    "settlement quirks, statistical tricks that make randomness look like skill. Sixteen "
    "tasks each hide one verified defect, verified meaning a script proves the defect changes "
    "the answer. Four tasks are clean; the correct answer there is “no defect,” because a "
    "reviewer that invents problems to look thorough is exactly as dangerous as one that "
    "misses real ones. Grading was blind: answers were stripped of identifying details, "
    "shuffled, and scored against a fixed rubric by a grader that did not know which model "
    "wrote which answer."
)

add_heading(doc, "Finding 1: cost and accuracy do not move together", level=1)
doc.add_paragraph(
    "The mid-tier model (Sonnet 5) matched the flagship model (Fable 5) on defects found, "
    "fifteen out of sixteen either way, at roughly one-tenth the cost. The most expensive "
    "model in the lineup (Opus 4.8) was not the most accurate. The cheapest model (Haiku 4.5) "
    "found the fewest defects but also invented the fewest false ones; answer length tracked "
    "false alarms across the board, so the two most verbose models were also the two that "
    "flagged the most non-existent problems on the clean tasks."
)
insert_chart(doc, IMG / "chart1_cost_vs_accuracy.png", number=1,
             caption="Cost vs. accuracy across four Claude tiers on the same 20-task battery.",
             source="internal cost/accuracy benchmark, cost estimated at published per-token pricing",
             as_of=AS_OF)
add_source_table(
    doc, number="1a", caption="Full cost/accuracy table.",
    headers=["Model", "Defects found (of 16)", "False positives (of 4 clean)", "Cost, 20 tasks (USD est.)", "Cost per defect found"],
    rows=[
        ["Sonnet 5", "15/16", "3/4", "$0.148", "$0.0099"],
        ["Fable 5", "15/16", "2/4", "$1.492", "$0.0995"],
        ["Opus 4.8", "14/16", "4/4", "$2.110", "$0.1507"],
        ["Haiku 4.5", "9/16", "1/4", "$0.025", "$0.0028"],
    ],
    source="internal cost/accuracy benchmark", as_of=AS_OF, right_align_cols={1, 2, 3, 4},
)

add_heading(doc, "Finding 2: I measured my own grader’s bias, by accident", level=1)
doc.add_paragraph(
    "While grading a separate, smaller set of open-ended answers, one model graded a ranking "
    "that looked wrong: it placed a stronger model below a weaker one. A neutral second judge "
    "reversed the ranking. Comparing both judges’ scores for every model showed a clean "
    "pattern: each judge inflated its own model family by roughly half a point to a full point "
    "on a ten-point scale, while the two models neither judge shared a family with barely "
    "moved. This is a direct, measured instance of a bias that is usually discussed only in "
    "theory: an AI model grading another AI model tends to grade its own family generously."
)
insert_chart(doc, IMG / "chart2_judge_self_preference.png", number=2,
             caption="The same answers, graded twice, by two different judges.",
             source="internal grading-bias check (leave-one-out correction)", as_of=AS_OF)
doc.add_paragraph(
    "The practical takeaway is simple: if you are using one AI model to grade another, do not "
    "trust a single judge, and check for this specifically if the judge and any of the models "
    "being graded share a family. It is a cheap, avoidable bias, and it is easy to miss if you "
    "are not looking for it."
)

add_heading(doc, "What I am not claiming here", level=1)
doc.add_paragraph(
    "This note is about model selection and grading hygiene, not about whether adding more "
    "process or more agents around a model improves a result, which is a separate and more "
    "nuanced question I tested independently and am treating carefully rather than folding "
    "into a headline. It is also not a claim about any specific market, trade, or investment; "
    "this is research-process work, done on personal time, with no capital involved."
)

add_heading(doc, "Method, in one paragraph", level=1)
doc.add_paragraph(
    "Every task, the rubric, and the pass criteria were fixed and committed before any model "
    "saw them. Grading was blind: a scrub pass removed anything that could identify which "
    "model produced an answer, each answer got a random id, and the id-to-model mapping was "
    "sealed until every grade was filed. The method is meant to be copyable: freeze your test "
    "before you run it, grade blind, include clean controls that penalize false alarms, and "
    "check your grader for the same bias you are trying to measure in the thing being graded."
)

out_path = REP / "FIRM_S_LINKEDIN_ATTACHMENT_20260715.docx"
doc.save(str(out_path))
print("LinkedIn attachment docx saved:", out_path)
