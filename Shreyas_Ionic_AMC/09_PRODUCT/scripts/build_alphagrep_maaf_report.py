"""ALPHAGREP MAAF NFO ANALYSIS -- human-readable Word report for the Principal.
Owner: Neel Basu (Attribution). Prepared 2026-07-05 for the Principal's meeting; NFO opens 6-Jul-2026.
Reusable: re-run to regenerate. Charts follow the firm dataviz palette + one-axis/thin-mark method.
All source numbers are labeled [DECK pX] / [VERIFIED vs our data] / [UNVERIFIABLE].
Companion reproducible computation: verify_agmaaf_numbers.py (same dir).
"""
import datetime as dt
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as mtick
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[3]
OUT = ROOT / "Shreyas_Ionic_AMC" / "09_PRODUCT" / "reports"
OUT.mkdir(parents=True, exist_ok=True)
STAMP = "2026-07-05"
IMG = OUT / "_img_agmaaf"
IMG.mkdir(exist_ok=True)

BLUE, AQUA, YELLOW, GREEN, VIOLET, RED = "#2a78d6", "#1baf7a", "#eda100", "#008300", "#4a3aa7", "#e34948"
GRID = "#e6e5e0"
INK, INK2 = "#1a1a19", "#5f5e57"
plt.rcParams.update({
    "figure.facecolor": "white", "axes.facecolor": "white",
    "axes.edgecolor": GRID, "axes.grid": True, "grid.color": GRID, "grid.linewidth": 0.6,
    "axes.spines.top": False, "axes.spines.right": False,
    "font.size": 9, "axes.titlesize": 11, "axes.titleweight": "bold",
    "text.color": INK, "axes.labelcolor": INK2, "xtick.color": INK2, "ytick.color": INK2,
})

# ---------------------------------------------------------------- CHART 1: claimed metrics (BACKTEST)
# p13 table: NIFTY / Static MA / Active MA / AGMAAF
labels = ["NIFTY\n(p13)", "Static\nMulti-Asset", "Active\nMulti-Asset", "AGMAAF"]
cagr = [11.04, 9.56, 10.42, 13.90]
sharpe = [0.55, 0.77, 1.44, 1.90]
maxdd = [-59.50, -39.64, -13.67, -12.79]
cols = [YELLOW, AQUA, VIOLET, BLUE]
fig, axes = plt.subplots(1, 3, figsize=(7.4, 3.0), dpi=200)
for ax, vals, ttl, fmt in [(axes[0], cagr, "CAGR %", "{:.1f}"),
                            (axes[1], sharpe, "Sharpe", "{:.2f}"),
                            (axes[2], maxdd, "Max drawdown %", "{:.1f}")]:
    bars = ax.bar(range(4), vals, color=cols, width=0.72)
    for b, v in zip(bars, vals):
        ax.annotate(fmt.format(v), (b.get_x() + b.get_width() / 2, v),
                    ha="center", va="bottom" if v >= 0 else "top", fontsize=7.5, color=INK2)
    ax.set_title(ttl, fontsize=10)
    ax.set_xticks(range(4)); ax.set_xticklabels(labels, fontsize=6.6)
    ax.axhline(0, color=INK2, lw=0.6)
fig.suptitle("What the deck claims (p13) -- INTERNAL BACKTEST, NO LIVE RECORD", fontsize=10.5, fontweight="bold")
fig.tight_layout(rect=[0, 0, 1, 0.95])
fig.savefig(IMG / "claimed.png", bbox_inches="tight"); plt.close(fig)

# ---------------------------------------------------------------- CHART 2: attribution waterfall (signature)
# 13.90% headline decomposed at AVG net exposures (p11): eq 28.5 / FI 45.75 / commodity 25.75
# beta contributions from OUR verified asset returns (equity=approx NIFTY50 TRI, FI=liquid+0.5, gold=GOLDBEES)
steps = [("Equity\nbeta", 3.72, BLUE), ("Fixed-income\nbeta", 3.42, AQUA),
         ("Commodity\nbeta", 3.71, YELLOW), ("Residual:\ntiming+selection", 3.05, RED)]
fig, ax = plt.subplots(figsize=(7.4, 3.4), dpi=200)
cum = 0.0
for i, (lab, val, c) in enumerate(steps):
    ax.bar(i, val, bottom=cum, color=c, width=0.66, edgecolor="white")
    ax.annotate("+%.2f" % val, (i, cum + val / 2), ha="center", va="center", fontsize=8.5,
                color="white", fontweight="bold")
    cum += val
ax.bar(len(steps), cum, color=INK, width=0.66)
ax.annotate("%.2f%%" % cum, (len(steps), cum / 2), ha="center", va="center", fontsize=9,
            color="white", fontweight="bold")
# 78% beta bracket
ax.annotate("diversified BETA = 10.85pp (78%)\nbuyable from any static multi-asset mix",
            (1.0, 11.9), fontsize=7.8, color=INK2, ha="center")
ax.annotate("SKILL claim (22%)\nin-sample, spliced,\nex-post regime labels",
            (3.0, 8.0), fontsize=7.6, color=RED, ha="center")
ax.set_xticks(range(len(steps) + 1))
ax.set_xticklabels([s[0] for s in steps] + ["AGMAAF\nclaimed CAGR"], fontsize=7.2)
ax.set_ylabel("Contribution to CAGR (pp)")
ax.set_title("Where the 13.90% comes from -- our attribution (equity beta = NIFTY50 approx-TRI)")
ax.set_ylim(0, 15)
fig.tight_layout(); fig.savefig(IMG / "attrib.png", bbox_inches="tight"); plt.close(fig)

# ---------------------------------------------------------------- CHART 3: verification (2 panels)
fig, (axL, axR) = plt.subplots(1, 2, figsize=(7.4, 3.2), dpi=200)
# LEFT: NIFTY50 -- deck 'TRI' vs our data
grp = ["CAGR %", "Volatility %", "Max DD % (abs)"]
deck = [11.04, 20.66, 59.50]
ourp = [11.76, 20.98, 59.86]       # our PRICE index
ourt = [13.06, 20.98, 59.86]       # our approx TRI (only CAGR differs)
x = range(3); w = 0.26
axL.bar([xi - w for xi in x], deck, w, color=YELLOW, label="Deck 'NIFTY TRI' (p13)")
axL.bar(list(x), ourp, w, color=BLUE, label="Our NIFTY50 PRICE")
axL.bar([xi + w for xi in x], ourt, w, color=VIOLET, label="Our NIFTY50 approx-TRI")
for xi, (d, pp, tt) in enumerate(zip(deck, ourp, ourt)):
    for off, v, cc in [(-w, d, YELLOW), (0, pp, BLUE), (w, tt, VIOLET)]:
        axL.annotate("%.1f" % v, (xi + off, v), ha="center", va="bottom", fontsize=6.2, color=INK2)
axL.set_xticks(list(x)); axL.set_xticklabels(grp, fontsize=7.2)
axL.legend(frameon=False, fontsize=6.6, loc="upper right")
axL.set_title("NIFTY 50: our data reproduces vol & maxDD;\ntheir 'TRI' behaves like PRICE", fontsize=8.8)
# RIGHT: gold super-cycle (performance-chasing) -- load GOLDBEES if available
try:
    gb = pd.read_parquet(ROOT / "datasets" / "etf_gold_silver" / "goldbees_daily.parquet")
    gb["date"] = pd.to_datetime(gb["timestamp"]).dt.tz_convert("Asia/Kolkata").dt.tz_localize(None)
    gb = gb.sort_values("date").set_index("date")["close"].astype(float)
    sc = gb[(gb.index >= "2024-07-01") & (gb.index <= "2026-06-30")]
    reb = sc / sc.iloc[0] * 100.0
    axR.plot(reb.index, reb.values, color=YELLOW, lw=1.8)
    axR.fill_between(reb.index, 100, reb.values, color=YELLOW, alpha=0.12)
    axR.annotate("+%.0f%%" % (reb.iloc[-1] - 100), (reb.index[-1], reb.iloc[-1]),
                 textcoords="offset points", xytext=(-4, 2), fontsize=9, color="#a06a00",
                 fontweight="bold", ha="right")
    axR.axhline(100, color=INK2, lw=0.6)
    axR.set_ylabel("GOLDBEES (rebased 100)")
    axR.set_title("Gold super-cycle: +112% Jul-24 -> Jun-26\n(NFO launches a 25%-commodity fund right after)", fontsize=8.8)
except Exception as e:
    axR.text(0.5, 0.5, "GOLDBEES load skipped:\n%s" % e, ha="center", va="center", fontsize=7)
fig.tight_layout(); fig.savefig(IMG / "verify.png", bbox_inches="tight"); plt.close(fig)

# ================================================================ DOCX
doc = Document()
st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10)

def h(text, lvl=1): doc.add_heading(text, level=lvl)

def p(text, bold=False, size=10, color=None, italic=False):
    para = doc.add_paragraph(); r = para.add_run(text); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor.from_string(color)
    return para

def shade(cell, hexcol):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexcol)
    tcPr.append(sh)

def table(headers, rows, shade_col=None, shade_map=None, fs=8.5):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers)); t.style = "Light Grid Accent 1"
    for j, htxt in enumerate(headers):
        c = t.rows[0].cells[j].paragraphs[0].add_run(htxt); c.bold = True; c.font.size = Pt(fs)
    for i, row in enumerate(rows, 1):
        for j, v in enumerate(row):
            r = t.rows[i].cells[j].paragraphs[0].add_run(str(v)); r.font.size = Pt(fs)
            if shade_col is not None and j == shade_col and shade_map:
                key = str(v).strip().upper()
                if key in shade_map:
                    shade(t.rows[i].cells[j], shade_map[key]); r.font.color.rgb = RGBColor.from_string("FFFFFF"); r.bold = True
    doc.add_paragraph()
    return t

RAG = {"RED": "c0392b", "AMBER": "d68910", "GREEN": "1e8449"}

# ---- title
title = doc.add_paragraph(); tr = title.add_run("AlphaGrep Multi Asset Allocation Fund (AGMAAF) -- NFO Due-Diligence")
tr.bold = True; tr.font.size = Pt(17); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sr = sub.add_run("Independent analysis for the Principal  |  NFO opens 6-Jul-2026  |  prepared %s by Neel Basu, "
                 "Performance Attribution, Shreyas Ionic AMC" % STAMP)
sr.font.size = Pt(9.5); sr.font.color.rgb = RGBColor.from_string("5F5E57"); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
lbl = doc.add_paragraph()
lr = lbl.add_run("Every figure is tagged: [DECK pX] = their claim  |  [VERIFIED] = checked against our data  |  "
                 "[UNVERIFIABLE] = we lack the source series.  This is due-diligence, not investment advice.")
lr.font.size = Pt(8.5); lr.italic = True; lr.font.color.rgb = RGBColor.from_string("5F5E57")

# ================================================================ 1. EXECUTIVE VERDICT
h("1. Executive verdict")
p("AGMAAF is a competently-structured, tax-efficient, open-ended multi-asset fund from a genuinely credible global "
  "quant sponsor. But the PITCH is a fully backtested, in-sample-optimized story on spliced history, launched with "
  "classic performance-chasing timing. Our decomposition finds that about 78% of the headline 13.90% CAGR is ordinary "
  "diversified BETA available from any static equity/debt/gold mix; only ~3pp is a timing-plus-selection residual, and "
  "that residual is measured in-sample, on data where the equity engine did not exist before 2011 and the commodity "
  "engine did not exist before 2017. There is no live, verifiable track record anywhere in the deck -- even though the "
  "sponsor runs live AIFs, a PMS and an offshore fund whose records exist at SEBI. Verdict: a diversified multi-asset "
  "BETA product with an UNPROVEN timing overlay -- not the 13.9%/1.90-Sharpe skill machine the deck implies.", size=10.5)
p("This mirrors our own hardest lesson (Jul-2026, IC-1): 71% of a flagship 'edge' turned out to be unconditional beta. "
  "Here it is ~78%. Headlines decompose.", size=9.5, italic=True, color="5F5E57")

h("Scorecard", 2)
table(["Dimension", "Rating", "One-line finding"], [
    ["Live track record", "RED", "100% backtest; ZERO live NAVs shown despite sponsor running live AIFs/PMS/offshore [DECK p3,p12]"],
    ["History integrity", "RED", "Equity sleeve = NIFTY TRI pre-2011; commodity = GOLD BEES pre-2017 (spliced) [DECK p13]"],
    ["Headline Sharpe 1.90", "RED", "It IS the optimization objective (Layer-1 maximizes Sharpe), reported in-sample [DECK p7,p10,p13]"],
    ["Product timing", "RED", "Launches 25%-avg-commodity fund right after gold +112% (Jul24-Jun26) [VERIFIED]"],
    ["Number discipline", "AMBER", "3 different CAGRs (13.90/14.17/16.32); risk stats identical across different windows"],
    ["Fee transparency", "AMBER", "TER not disclosed in deck; 'net of fees/slippage' claimed but unquantified"],
    ["Sponsor pedigree", "AMBER", "Real ~$1bn global quant firm, 200+ researchers -- but HFT/latency edge does NOT transfer to a weekly MF"],
    ["Structure / terms / tax", "GREEN", "Equity taxation (12.5% LTCG@24m) via >35% gross equity; 1%/15d exit load; Rs500 min -- investor-friendly [DECK p7,p17]"],
], shade_col=1, shade_map=RAG)
p("Overall: 4 RED / 3 AMBER / 1 GREEN. The GREEN (structure/tax) is real and to the investor's benefit; the REDs are "
  "all about EVIDENCE -- the fund may well be fine, but nothing in this deck lets you know that yet.", size=9.5)
doc.add_picture(str(IMG / "claimed.png"), width=Inches(6.9))
cap = doc.add_paragraph(); cr = cap.add_run("Figure 1. The deck's own claimed metrics (p13). Every bar is an internal "
    "backtest with no live component. Read the levels, then read Sections 4-5 for why they do not mean what they appear to.")
cr.font.size = Pt(8); cr.italic = True; cr.font.color.rgb = RGBColor.from_string("5F5E57")

# ================================================================ 2. WHAT THE FUND IS
doc.add_page_break()
h("2. What the fund actually is (plain language)")
p("An open-ended, actively-but-systematically managed multi-asset fund that rebalances WEEKLY across three sleeves "
  "using two proprietary model layers: Layer-1 sets asset-class weights (risk-parity + Sharpe maximisation + a macro "
  "regime overlay); Layer-2 picks the equity names (multi-factor: value/momentum/quality/size + 'alpha' signals) and "
  "the commodity mix (gold/silver/copper/crude) [DECK p7-p10].", size=10)
table(["Term", "Detail", "Read-through"], [
    ["Structure", "Open-ended multi-asset [p17]", "Buy/sell any day at NAV -- no lock-in beyond exit load"],
    ["Asset allocation (SID)", "Equity 35-60%, Debt 10-60%, Commodities 10-40%, InvITs 0-10% [p17]", "SID GROSS-equity band"],
    ["Asset allocation (deck)", "Net equity 10-60%, FI 10-60%, commodity 10-40% [p7]", "NET (directional) band -- different measure, see below"],
    ["Avg backtest exposure", "Net equity 28.5%, FI 45.75%, commodity 25.75% [p11]", "~Half the book is debt on average -- this is a CONSERVATIVE mix"],
    ["Tax", "Maintains >35% GROSS equity via arbitrage + REITs [p7]", "Equity taxation: 12.5% LTCG after 24m (vs slab on debt funds) -- investor-friendly"],
    ["Exit load", "1% if <15 days; nil after [p17]", "Cheap to exit -- genuinely investor-friendly"],
    ["Minimums", "Rs500 lump / Rs500 SIP, 6 min installments [p17]", "Low barrier"],
    ["Benchmark", "35% NIFTY 200 TRI + 45% NIFTY Composite Debt + 20% MCX iCOMDEX [p17]", "A real, fair multi-asset benchmark"],
    ["Fund Manager", "Mr. Ravneet Singh [p17]", "No individual attributable track record shown -- FLAG"],
    ["TER", "Not in deck", "Ask -- it directly reduces every number shown"],
], fs=8.2)
p("The net-vs-gross equity distinction is not sloppiness -- it is the tax engine. The fund can hold, say, ~10% "
  "directional equity + ~25% market-neutral arbitrage/REITs = >35% gross equity (qualifying for 12.5% LTCG) while "
  "running low NET equity risk. That is a legitimate, investor-friendly design. The deck just never explains it on one "
  "page, so the '10-60%' (p7) and '35-60%' (p17) bands look contradictory when they are actually measuring different "
  "things (net vs gross).", size=9.5)

# ================================================================ 3. CLAIMS vs FINE PRINT
h("3. Claims vs fine print -- the splice")
p("The single most important slide in the deck is the fine print on p13. The 20-year 'track record' is not one strategy "
  "run for 20 years; it is today's model stitched onto passive indices for the years the model did not exist.", size=10)
table(["Marketing claim", "Fine print (p13, p8, p14)", "What it means"], [
    ["'20-year track record', Rs1L -> Rs12.47L [p12]", "Equity = NIFTY TRI before 2011; stock-selection engine only from 2011",
     "For 2006-2011 there was NO equity alpha -- it was the index. ~5 of 20 years are passive-spliced."],
    ["'Multi-commodity' diversification [p8]", "Commodity = GOLD BEES only before 2017; multi-commodity after",
     "For 2006-2017 the 'commodity engine' was just gold. ~11 of 20 years are gold-only."],
    ["'Smooth highway through the GFC' [p14]", "Max drawdown -12.79% is IDENTICAL on p13/p14/p18",
     "That drawdown occurs in the 2020-26 window (COVID) -- so 2008 was MILDER, because in 2008 the book was spliced "
     "passive NIFTY + heavy debt/gold. The 'GFC survival' is a splice + asset-mix artifact, not model skill."],
    ["Gold spot pre-2007", "MCX Gold Spot used before March 2007 [p13]", "Earliest data is a different instrument again."],
    ["'Alpha at both layers' [enabler p2]", "Allocation ranges 'based on backtested portfolios 2007-26' [p8]",
     "Even the sleeve ranges are fitted to the same backtest."],
], fs=8.2)
p("Net: the honest 'model-as-it-exists-today, all-engines-live' history is 2017-2026 (~9 years), not 20. Ask for the "
  "stats on that sub-period alone (Meeting Q2).", size=9.5)

# ================================================================ 4. INDEPENDENT VERIFICATION
doc.add_page_break()
h("4. Our independent verification (checked against our data)")
p("We hold NSE official index NAVs 2005-2026 (factor_navs_principal.parquet, D-009 verified), 174 official index series "
  "2016-2026 (nse_official_all_indices.parquet, triple-verified), and D-009-passed GOLDBEES daily. Here is what we could "
  "check and what we could not.", size=10)
table(["Deck claim", "Our data says", "Verdict"], [
    ["NIFTY 'TRI' 20.66% vol, -59.50% maxDD [p13]", "NIFTY 50 PRICE: 20.98% vol, -59.86% maxDD (2006-2026)",
     "VERIFIED -- our data reproduces their NIFTY series (vol & DD match to <0.4pp)"],
    ["NIFTY 'TRI' CAGR 11.04% [p13]", "NIFTY 50 PRICE 11.76%; true TRI ~13.06% (price + ~1.3% div)",
     "MISLABEL -- their 'TRI' behaves like the PRICE index; a real TRI is ~1.3-1.5pp higher"],
    ["Benchmark leg: NIFTY 200 TRI (2020-26) [p18]", "NIFTY 200 PRICE 13.18%, TRI ~14.48%, vol 18.1%, maxDD -38.2%",
     "VERIFIED plausible -- blended 11.48% benchmark reconciles with debt+iCOMDEX legs"],
    ["'Gold super-cycle Jul-24->May-26, gold +50%' [p5]", "GOLDBEES +112.5% total over that window (~50% ANNUALISED)",
     "VERIFIED -- and it underlines the performance-chasing launch timing"],
    ["Static Multi-Asset comparator = 9.56% [p13]", "Our 60/25/15 reconstruction = ~11.1% (price eq) to ~11.9% (TRI eq)",
     "QUESTION -- our fair static baseline is ~1.5-2.3pp HIGHER, which would compress AGMAAF's apparent edge"],
    ["NIFTY Composite DEBT index leg [p17]", "We hold only Nifty Composite G-SEC (govt-only), not the corporate composite",
     "UNVERIFIABLE -- cannot check the 45% debt leg precisely"],
    ["MCX iCOMDEX composite leg [p17]", "We hold Nifty Commodities (an EQUITY index), not the MCX futures index",
     "UNVERIFIABLE -- cannot check the 20% commodity-futures leg"],
    ["AGMAAF model returns / Sharpe 1.90 [p12-13]", "Proprietary model -- no external series exists",
     "UNVERIFIABLE by construction -- only a LIVE record could confirm it"],
], fs=8.0)
doc.add_picture(str(IMG / "verify.png"), width=Inches(6.9))
cap = doc.add_paragraph(); cr = cap.add_run("Figure 2. Left: our NIFTY 50 data reproduces the deck's volatility and max "
    "drawdown almost exactly (so both sides are using the same market history) -- but the deck's '11.04% NIFTY TRI' lines "
    "up with the PRICE index, not the ~13% total-return index. Right: gold (GOLDBEES) more than doubled in the two years "
    "before this NFO -- the exact backdrop for a performance-chasing commodity launch.")
cr.font.size = Pt(8); cr.italic = True; cr.font.color.rgb = RGBColor.from_string("5F5E57")

# ================================================================ 5. QUANT CRITIQUE + ATTRIBUTION
doc.add_page_break()
h("5. The quant critique -- attribution, in-sample optimisation, deflated Sharpe, regime hindsight")

h("5a. Attribution: 78% of the headline is beta", 2)
p("We decomposed the claimed 13.90% CAGR using the fund's OWN average net exposures (p11: equity 28.5% / FI 45.75% / "
  "commodity 25.75%) and OUR verified asset-class returns over the same era (equity = NIFTY 50 approx-TRI 13.06%; "
  "FI = HDFC Liquid proxy 6.97% + 0.5pp = ~7.47%; commodity = GOLDBEES 14.40%).", size=10)
table(["Component", "Contribution to CAGR", "Share of 13.90%", "Method"], [
    ["Equity beta", "+3.72 pp", "27%", "0.285 x 13.06% (NIFTY50 approx-TRI) [VERIFIED input]"],
    ["Fixed-income beta", "+3.42 pp", "25%", "0.4575 x 7.47% (liquid-fund proxy + 0.5) [VERIFIED-ish]"],
    ["Commodity beta", "+3.71 pp", "27%", "0.2575 x 14.40% (GOLDBEES full era) [VERIFIED input]"],
    ["= Diversified beta", "+10.85 pp", "78%", "Available from ANY static multi-asset mix -- no skill required"],
    ["Residual: timing + selection", "+3.05 pp", "22%", "13.90% - beta; the ONLY part that could be skill"],
], fs=8.3)
doc.add_picture(str(IMG / "attrib.png"), width=Inches(6.6))
p("The residual is +3.0pp (using total-return equity) to +3.4pp (using price equity) -- call it ~3pp. Now shrink it "
  "further with what is actually attainable: the equity selection engine existed only ~15 of 20 years (post-2011) and "
  "harvesting factor premia costs 3.5-10.7pp/yr at high rebalance frequency (our own KB finding); the commodity engine "
  "existed only ~9 of 20 years (post-2017) and multi-commodity most likely UNDERPERFORMED pure gold over 2017-26 "
  "(gold was the star). A generous selection contribution is +0.5 to +0.8pp -- leaving ~2.2-2.5pp of pure market-TIMING "
  "alpha over 20 years. Sustained TAA timing alpha of that size is something almost no real manager achieves, and here "
  "it is measured entirely in-sample with ex-post regime labels. That single ~2.3pp is the least credible number in the "
  "pitch.", size=9.5)

h("5b. The Sharpe 1.90 is the optimisation objective, not an estimate", 2)
p("Layer-1 explicitly 'maximises Sharpe' (p7, p10). The deck then reports the maximised Sharpe (1.90) on the SAME "
  "history. That is not a test statistic -- it is a fitted value. We ran a deflated-Sharpe illustration to be precise "
  "about what that does and does not prove:", size=10)
table(["Trials evaluated (N)", "Noise-only Sharpe ceiling (ann.)", "Deflated-Sharpe (prob true SR>0)", "Live months to confirm 1.90 (MinTRL)"], [
    ["1 (no selection)", "0.00", "~1.00", "~11"],
    ["50", "0.50", "~1.00", "~20"],
    ["200", "0.61", "~1.00", "~24"],
    ["1000", "0.72", "~1.00", "~28"],
], fs=8.3)
p("Honest read (this cuts BOTH ways): a Sharpe of 1.90 is high enough that even best-of-1000 discrete trial-selection "
  "does NOT explain it away -- IF it were a genuine realised sample, it would be significant (SE +-0.24; naive 95% CI "
  "[1.43, 2.37]). So the problem is NOT sampling noise, and we should not pretend deflation drags 1.90 to zero. The "
  "problem is that 1.90 is the in-sample OPTIMISATION TARGET on partly-synthetic (spliced) data -- a regime that "
  "deflated-Sharpe math cannot repair, because it assumes real, unoptimised returns. The clean resolution is "
  "out-of-sample: the table's right column shows just ~1-2 years of LIVE data would confirm a real 1.90 -- and the "
  "sponsor, who runs live vehicles, shows exactly zero. That silence is the tell.", size=9.5)

h("5c. Regime table is labelled with hindsight", 2)
p("p5 tags a single 'winning asset class' per regime AFTER the fact. That is how you'd grade the past, not how a model "
  "trades the future. The live question -- unproven in the deck -- is the DETECTION LAG: on the COVID crash (Feb-Mar "
  "2020) and the 2022 rate-hike, how many days/weeks after the regime turned did the weekly model actually reallocate? "
  "In a strict point-in-time backtest, that lag typically erases much of the 'right asset at the right time' benefit. "
  "The p5 returns themselves are largely fine (we verified the gold super-cycle); it is the ex-post LABELLING and the "
  "implied real-time agility that are unproven.", size=9.5)

h("5d. Number discipline", 2)
table(["Inconsistency", "Explanation", "Verdict"], [
    ["CAGR 13.90% [p13] vs 14.17% [enabler] vs 16.32% [p18]", "Windows: Jan06-May26 / Jan07-Apr26 / Jan20-Apr26",
     "Mostly window artifacts (each reconciles to its terminal value); 16.32% is the cherry-picked hot window"],
    ["Rs12.47L [p12] vs Rs12.9L [enabler]", "Same as above -- different windows",
     "One discrepancy, not two: 12.47L->13.89% CAGR, 12.9L->14.17% CAGR"],
    ["Vol/Sharpe/maxDD identical (7.44/1.90/-12.79) across different windows", "Risk stats not recomputed per window",
     "Sloppiness -- undermines confidence in the bookkeeping"],
    ["-12.79% maxDD on p13/p14/p18", "The worst DD sits inside 2020-26 (COVID)",
     "p14's 'GFC survived at -12.79%' mislabels the COVID drawdown as the 2008 one"],
], fs=8.2)

# ================================================================ 6. PEDIGREE
doc.add_page_break()
h("6. Pedigree and capability transferability")
p("Be fair: AlphaGrep is a real, substantial global quant/HFT firm -- offices in 8 countries, ~$1bn across AM + prop, "
  "500+ professionals incl. 200+ researchers, trades 30+ exchanges [DECK p3]. The systematic discipline, data "
  "infrastructure and research depth are genuine assets, and the AMC already runs live CAT-III AIFs, a long-only PMS "
  "and a GIFT-City offshore fund.", size=10)
table(["Capability", "Transfers to a weekly T+1 mutual fund?", "Note"], [
    ["Colocation / latency / HFT execution", "NO", "A weekly-rebalanced MF cannot use microsecond edges -- irrelevant here"],
    ["Microstructure / order-book 'alpha' (p8)", "MOSTLY NO", "Order-book signals decay in minutes; useless at weekly horizon"],
    ["Systematic research discipline & data infra", "YES", "Real and valuable -- but common to many quant AMCs now"],
    ["Multi-factor equity modelling", "YES (commoditised)", "Value/momentum/quality are public premia; edge is in cost-control & capacity"],
    ["Running money vs trading flow", "UNPROVEN", "Prop/HFT skill != asset-gathering MF skill; different game, no MF live record"],
], fs=8.3)
p("So the brand does count for something -- but the specific thing AlphaGrep is famous for (latency) is exactly the "
  "thing this product cannot use. The relevant, transferable skills (factor research, execution cost control) are real "
  "but no longer rare, and their pay-off here is entirely unproven live. Do not let the HFT halo stand in for a track "
  "record (Meeting Q13).", size=9.5)

# ================================================================ 7. MEETING QUESTIONS
doc.add_page_break()
h("7. Questions for the meeting (ordered by importance)")
p("This is the section that matters. Each question is designed to separate a confident, honest quant shop from a "
  "marketing exercise. For each: what a GOOD answer sounds like, and what a BAD/evasive answer sounds like.", size=10)

def q(n, question, good, bad):
    para = doc.add_paragraph()
    r = para.add_run("Q%d. %s" % (n, question)); r.bold = True; r.font.size = Pt(9.8)
    g = doc.add_paragraph(); gr = g.add_run("   GOOD answer: " + good); gr.font.size = Pt(9); gr.font.color.rgb = RGBColor.from_string("1E8449")
    b = doc.add_paragraph(); br = b.add_run("   BAD answer: " + bad); br.font.size = Pt(9); br.font.color.rgb = RGBColor.from_string("C0392B")

q(1, "You run live CAT-III AIFs, a PMS and an offshore fund. Why is there NOT ONE live NAV in this deck -- only a backtest? Show the live track record of your existing systematic vehicles.",
  "Produces live NAVs of comparable systematic strategies with a realistic (not 1.90) realised Sharpe; explains AGMAAF is new but the engines run live elsewhere.",
  "'The backtest is representative', 'different strategy', or deflects. If a live-vehicle sponsor won't show live numbers, assume they don't help the story.")
q(2, "The equity engine existed only from 2011 and multi-commodity from 2017 (p13). What are the CAGR/Sharpe/maxDD for 2017-2026 -- the model AS IT EXISTS TODAY, all engines live, no splicing?",
  "Provides the ~9-year all-engines-live sub-period stats and concedes the earlier years are passive-spliced.",
  "Insists the full 20-year number 'is' the model, or can't produce the un-spliced sub-period.")
q(3, "Layer-1 maximises Sharpe and you report the maximised Sharpe (1.90) on the same data. How many model variants/configs were evaluated, and what is the WALK-FORWARD / out-of-sample Sharpe (purged CV)?",
  "Gives a trial count and an OOS/walk-forward Sharpe that is materially below 1.90; describes purging/embargo.",
  "'Proprietary', or no in-sample-vs-OOS distinction. A shop that optimises Sharpe must know its OOS Sharpe.")
q(4, "Our attribution says ~78% of the 13.90% is asset-class beta at your average weights and only ~3pp is timing+selection. Break the residual into timing vs selection, net of honest weekly-rebalance costs.",
  "Shares their own attribution, separates timing from selection, quantifies turnover and cost drag.",
  "Claims the whole 13.9% is 'the model', or has no cost breakdown for weekly rebalancing.")
q(5, "The regime table (p5) labels winners ex-post. What is the real-time DETECTION LAG -- on COVID (Feb-Mar 2020) and the 2022 hikes, how many days after the turn did the weekly model reallocate, in a strict point-in-time backtest?",
  "Shows PIT signals with a realistic lag and that the edge survives the lag.",
  "No PIT test; regime returns computed with perfect-hindsight labels.")
q(6, "You launch a 25%-avg-commodity fund right after gold returned ~112% (we verified Jul24-Jun26). What is the current live model commodity weight today, and does the backtest edge survive if you EXCLUDE the 2020-26 gold run?",
  "Edge is stable ex-gold-run; current commodity weight is moderate, not a chase.",
  "Very high current commodity weight; edge collapses without 2020-26 gold.")
q(7, "Your p13 'NIFTY TRI' at 11.04% matches the PRICE index in our data, not the ~13% total-return index. Is the benchmark price or total-return -- and is AGMAAF's OWN equity sleeve measured on total-return?",
  "Both on a consistent TRI basis, or a clean correction.",
  "Benchmark is price but the fund is TRI (asymmetric, flatters outperformance), or 'we'll check'.")
q(8, "What is the all-in TER (Regular and Direct)? Does the 13.90%/1.90 backtest net the ACTUAL TER plus realistic commodity-ETF roll, arbitrage and weekly-rebalance costs?",
  "Discloses TER; confirms the backtest nets the full live TER and realistic costs.",
  "TER undisclosed, or the backtest used a lower cost assumption than the live fund will bear.")
q(9, "Net equity averages 28.5% but you keep >35% gross via arbitrage+REITs for 12.5% LTCG. Is the 24-month equity-taxation treatment assured, and what is the drag from running the arbitrage book?",
  "Confirms the tax mechanics and quantifies a small arbitrage drag.",
  "Vague on tax assurance, or ignores the cost of the arbitrage sleeve.")
q(10, "Mr. Ravneet Singh -- what is his individual, attributable live record running systematic multi-asset money? Who controls the model parameters, and what is the key-person risk?",
  "A concrete attributable record + a documented model-governance process.",
  "No individual record; 'team-based'; parameters controlled by an unnamed group.")
q(11, "Your -12.79% maxDD is identical on the full-period, the 2008 slide and the 2020-26 slide -- implying the worst DD was COVID, not 2008 (when the equity sleeve was passive). What is the maxDD of TODAY'S model through a real 2008-scale equity shock?",
  "Acknowledges the splice; provides a stress-tested drawdown of the current all-engines model.",
  "Repeats 'smooth highway through the GFC'.")
q(12, "What AUM does this strategy's capacity top out at? Weekly rebalancing across single-stock equity + commodity ETFs/derivatives + arbitrage has real impact limits -- how does the edge decay with size?",
  "An honest capacity curve and a stated soft-close level.",
  "'No capacity limit' for a weekly-rebalanced multi-asset book.")
q(13, "AlphaGrep's global edge is colocation/latency HFT, which a weekly T+1 MF cannot use. Which SPECIFIC, transferable capabilities -- not latency -- drive THIS fund's edge?",
  "Names data infrastructure, factor research and execution cost-control -- not latency or the brand.",
  "Leans on the HFT reputation as if microsecond edges transfer to a weekly fund.")
q(14, "Our reconstruction of your 60/25/15 static comparator earns ~11%, not the 9.56% you show. How is your static baseline built -- and does a fair static baseline shrink AGMAAF's apparent edge?",
  "Shares the construction and reconciles the gap.",
  "Cannot reproduce its own comparator, or the baseline is quietly understated.")

# ================================================================ 8. BOTTOM LINE
doc.add_page_break()
h("8. Bottom line -- for whom, when, how much")
p("Framed two ways per our D-032 discipline: the AMC (firm) lens and the Principal-personal lens. These are different "
  "decisions and should not be blurred.", size=10)

h("AMC / firm lens", 2)
p("Treat AGMAAF as a COMPETITOR and a STUDY OBJECT, not an allocation. It is a near-perfect live specimen of the exact "
  "traps our process is built to catch -- in-sample Sharpe optimisation, spliced history, ex-post regime labels, "
  "performance-chasing launch timing. Recommended firm actions: (1) PAPER-TRACK it from day one -- log the live AMFI NAV "
  "weekly against its own benchmark and against a cheap DIY 35/45/20 static, exactly as we paper-track our own sleeves; "
  "(2) use it as a validation foil for our multi-asset / regime research; (3) do NOT spend firm research capital trying "
  "to reverse-engineer a 1.90 that is almost certainly an in-sample artifact. Zero firm capital; high learning value.", size=9.5)

h("Principal-personal lens", 2)
p("If the Principal wants genuine multi-asset diversification with equity-like taxation, the STRUCTURE is sound and "
  "investor-friendly (12.5% LTCG, cheap exit, low minimum). But there is no reason to enter at the NFO:", size=9.5)
table(["Question", "Answer"], [
    ["For whom?", "An investor who wants a conservative (~half-debt), tax-efficient, hands-off multi-asset holding -- NOT someone seeking the 'equity-like 13.9%' the deck implies (that's ~78% beta + an unproven overlay)."],
    ["When?", "NOT at NFO. NFOs carry no pricing advantage; wait 12-24 months for LIVE NAVs and let the timing overlay meet real markets (MinTRL says ~1-2 years settles the 1.90 claim)."],
    ["How much?", "If at all after live data: small and as a DIVERSIFIER, not a core equity substitute. A DIY 35/45/20 (index equity + a debt fund + a gold ETF) captures ~78% of the return at near-zero manager risk and lower cost."],
    ["Kill criteria to watch post-NFO", "Live Sharpe far below 1.90; commodity weight stays extreme after gold cools; live return tracks the static baseline (i.e., the overlay adds nothing); FM departure."],
], fs=8.6)

h("Monitoring sources to enable post-NFO (require approval -- NOT fetched now)", 2)
p("Per firm hard rule (no auto-fetch of new external sources), we did NOT pull any live data. Flagging two sources for "
  "the Data Officer's D-009 gate WHEN approved: (1) AMFI daily NAV file (public) -- to paper-track AGMAAF's live NAV vs "
  "benchmark from launch; (2) SEBI AIF / PMS disclosures -- to locate AlphaGrep's EXISTING live systematic records, which "
  "is the single best evidence the deck omits. Both are post-NFO monitoring inputs, not needed for today's meeting.", size=9.5)

pfin = doc.add_paragraph()
rr = pfin.add_run("Prepared by Neel Basu, Performance Attribution -- Shreyas Ionic AMC -- %s. Verification numbers "
    "reproducible via 09_PRODUCT/scripts/verify_agmaaf_numbers.py against factor_navs_principal.parquet, "
    "nse_official_all_indices.parquet and goldbees_daily.parquet. This is due-diligence for the Principal, not "
    "investment advice; the firm holds no position and proposes none." % STAMP)
rr.font.size = Pt(8); rr.font.color.rgb = RGBColor.from_string("5F5E57")

out = OUT / ("ALPHAGREP_MAAF_ANALYSIS_%s.docx" % STAMP)
doc.save(str(out))
print("SAVED:", out)
