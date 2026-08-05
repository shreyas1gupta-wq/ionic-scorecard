# -*- coding: utf-8 -*-
"""build_pac_qa_brief.py — PAC / CEO question-and-answer brief for the NDPMS Portfolio
Review engine (2026-08-04).

A prep document for whoever presents the product-approval deck: the questions a Product
Approval Committee is actually likely to ask, and a defensible answer to each. INTERNAL.

Design rules for the answers, deliberately:
  - Never claim evidence we do not have. Where the honest answer is "we cannot show that
    yet", it says so and says what we CAN show instead.
  - Lead with the number or the mechanism, not with reassurance.
  - Where an answer concedes a weakness, it also states the control that bounds it.

Output: 09_PRODUCT/reports/PAC_QA_BRIEF.docx
Usage:  python build_pac_qa_brief.py
"""
import os
import sys

os.environ.setdefault("PYTHONIOENCODING", "utf-8")
HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)

from docx import Document                      # noqa: E402
from docx.shared import Pt, Inches, RGBColor   # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402

import docx_style_kit as K                     # noqa: E402

OUT = os.path.abspath(os.path.join(HERE, "..", "reports", "PAC_QA_BRIEF.docx"))
AS_OF = "2026-08-04"

NAVY = RGBColor(0x1E, 0x3A, 0x8A)
INK = RGBColor(0x11, 0x18, 0x27)
SLATE = RGBColor(0x5B, 0x64, 0x748 & 0xFF)
RED = RGBColor(0xC0, 0x2B, 0x1D)

# (section, [(question, answer, optional_flag)])
# flag: None | "hard" (a question we should expect to be pressed on) | "concede"
QA = [
 ("A. Methodology and provenance", [
  ("Where did the pillar weights come from? Were they optimised on data?",
   "They were set by judgment, not optimised, and that is the honest answer. The design intent is "
   "documented: the 3-year score is fundamentals-tilted and the 1-year score is behaviour-tilted, "
   "which the shipped weights deliver at 58/42 and 48/52 respectively (fundamentals block being "
   "Quality, Growth and Value). No weight was fitted to maximise a backtested return. That is a "
   "deliberate choice — with the length of clean point-in-time history we have (8 quarterly "
   "rebalances), any weight optimisation would be curve-fitting, and we would rather present "
   "un-optimised weights we can defend from first principles than tuned weights we cannot.", "hard"),
  ("Then how do you know the weights are right?",
   "We do not, and we do not claim to. What we claim is narrower and testable: the weights encode a "
   "stated investment view (over three years fundamentals should dominate; over one year market "
   "behaviour should), they are applied identically to every name, and they are frozen and versioned "
   "so any change is visible and attributable. The weights are a hypothesis the forward record will "
   "judge. Anyone who tells you their factor weights are provably optimal on Indian equity with this "
   "much clean history is overclaiming.", "hard"),
  ("Why two horizons instead of one number?",
   "Because averaging them destroys the information. A good business in a poor tape and a weak "
   "business in a strong tape can produce the same blended score for opposite reasons, and the "
   "action a reviewer should take differs. Keeping them apart also lets the Sell rule be "
   "conservative: a name fails if EITHER horizon fails, so a fundamental problem cannot be masked "
   "by momentum, or vice versa. The client still sees one number (the Ionic Score); the two legs "
   "exist for the desk.", None),
  ("Why is Value only 18% of the 3-year score? Shouldn't valuation matter more for a Sell decision?",
   "This is the most reasonable challenge on the sheet and it is live internally. The case for "
   "raising it: this product reviews EXISTING holdings, and over three years the dominant way an "
   "existing holding hurts a client is de-rating from having overpaid, which is precisely what "
   "Value measures. At 18% it is also the smallest of the three fundamental pillars, which is odd "
   "for an exit discipline. The case for caution: 60% of the Value pillar is P/E by construction, "
   "so moving Value from 18% to 25% takes P/E alone from 10.8% to 15.0% of the whole 3-year score, "
   "and a low-P/E tilt in India loads heavily on PSUs, cyclicals and value traps. Valuation is also "
   "already expressed twice more downstream — in the analyst's reverse-DCF judgment and in the "
   "growth-led forward adjustment — so raising the pillar risks triple-counting it. Our position: "
   "we agree with the direction, we will not change it by assertion, and we have a point-in-time "
   "harness to test it. See the standing item at the end of this brief.", "hard"),
  ("Why did you exclude discounted cash flow as a pillar?",
   "Because the mechanical version produced implausible output and we would rather drop a pillar "
   "than ship a number we do not believe. A two-stage terminal-value formula on this universe "
   "implied 85 to 121 per cent growth to justify observed prices; even after correcting to a "
   "single-stage perpetuity it read 8 to 13 per cent, which is a wide enough swing on a small "
   "modelling choice that the output was not decision-grade. Valuation now enters as the analyst's "
   "own reverse-DCF judgment, written down per name. It is a candidate to return as a weighted "
   "pillar once that judgment has a track record.", None),
  ("Why is Quality ranked within sector but Growth ranked across the whole universe?",
   "Because the two questions differ. Return on capital is only meaningful against businesses with "
   "comparable capital intensity, so a bank, a cement plant and a software firm must be compared to "
   "their own kind. Growth is a claim on scarce capital across the whole market, and a 20 per cent "
   "grower is interesting whether or not its sector is growing. Value is ranked within sector AND "
   "size tier for the same reason as Quality, since multiples are not comparable across either.", None),
  ("What is the regime tilt, and isn't a discretionary overlay a back door?",
   "It is a small, disclosed adjustment reflecting the house macro call, currently Value or Cyclical "
   "Rotation as primary. It moves the Sector and Macro pillar by roughly plus 4.6 for cyclicals and "
   "minus 2.6 for non-cyclicals, which on an 11 per cent pillar weight is well under a point of "
   "composite score. It is deliberately mild, it is set by one named person on a documented cadence "
   "rather than per-stock, and it was NEUTRALISED entirely in our backtest so no historical claim "
   "depends on it. If the committee prefers it removed from the client-facing score, that is a "
   "clean change to make.", "hard"),
  ("Why do banks and NBFCs get an exemption from the leverage gate? That looks like special pleading.",
   "Because leverage is their product, not a symptom of distress: a bank with low debt to equity is "
   "a bank that is not lending. Applying the gate uniformly did not produce conservatism, it "
   "produced a false signal, and it was hitting roughly a sixth of one real client book by value. "
   "The exemption is narrow and it is not a free pass — those names are flagged as requiring a "
   "sector-specific read (capital adequacy, asset quality, credit costs, margin path) rather than "
   "being silently marked healthy. The pre-existing bug where the same names were penalised twice, "
   "once by the gate and again by the red-flag count, was found and fixed.", None),
  ("Why is the Sell line at 40 rather than 50?",
   "40 is the point below which the composite is signalling a genuine problem rather than mere "
   "mediocrity, and it produces a Sell rate of roughly a quarter of the universe, which matches "
   "what we would expect to be genuinely worth exiting in a broad Indian universe. A 50 line would "
   "flag close to half the book, which is not a review, it is a liquidation. The threshold is also "
   "not the only gate: a red balance-sheet or liquidity flag caps a score at 40 regardless of how "
   "the pillars scored, so the worst cases are pushed below the line by construction.", None),
  ("Are you double-counting valuation across the score, the analyst view and the forward adjustment?",
   "Partially yes, and we bound it deliberately. Valuation appears in the Value pillar, in the "
   "analyst's reverse-DCF, and indirectly in the growth-led forward adjustment. Two hard caps exist "
   "specifically to stop this compounding in the client's favour: a name growing below 10 per cent "
   "can be adjusted down but never up, and a name the analyst rates Sell can never be lifted by the "
   "adjustment. Both were added after observing the earlier version perversely raising the score of "
   "names we were telling clients to exit. It is the main reason we are cautious about raising the "
   "Value weight further.", "concede"),
 ]),
 ("B. Evidence and performance", [
  ("Does it work? Show me the backtest.",
   "Over a clean point-in-time window from December 2021 to September 2024, eight quarterly "
   "rebalances, the top decile by score returned 41.3 per cent CAGR against 32.4 for the bottom "
   "decile and 30.3 for the cap-weighted Nifty 500, and did so at a maximum drawdown of 2.3 per "
   "cent against 12.6 and 5.6. The drawdown gap is the one result we consider genuinely "
   "interesting, because it is consistent with what the score is built to do. We do not present "
   "this as proof of a return edge, for the reasons in the next three answers.", None),
  ("Your own placebo test puts the top decile at the 44th percentile of random baskets. Isn't the signal noise?",
   "On return selection over that window, yes, and we say so in our own report rather than leaving "
   "it for someone else to find. Two thousand random ten-name baskets placed our top decile at the "
   "44th percentile, and the long-short Sharpe was approximately zero. What survives that test is "
   "not the return ranking but the risk profile: the shallower drawdown is a property of screening "
   "out weak balance sheets and thin liquidity, and it does not depend on the return ranking being "
   "significant. We are comfortable defending a downside-protection claim and not comfortable "
   "defending an alpha claim.", "hard"),
  ("It also failed to beat equal-weighting the universe. Why ship it?",
   "Correct, it lagged an equal-weighted universe by 2.8 per cent a year over that window, which "
   "was dominated by a small-cap breadth rally in which almost any diversified basket beat a "
   "cap-weighted index. That is a real limitation of the test window, not a defence of the score. "
   "The product justification does not rest on beating equal weight: this is not a portfolio "
   "construction engine competing with an index, it is a review process that tells an existing "
   "holder which of the names they already own carry identifiable problems, with a written reason "
   "and an audit trail. Equal-weighting the universe is not an option available to a client who "
   "already owns 47 specific stocks.", "hard"),
  ("Why is the backtest only eight quarters?",
   "Because that is how much clean point-in-time fundamental history we have without lookahead. The "
   "annual PIT panel refreshes through FY2023 and prices begin mid-2021. We could produce a longer "
   "and much better-looking backtest by using current fundamentals over historical prices, and that "
   "would be lookahead and worthless. Extending the panel is an active data-office task, and it is "
   "the single change that would most improve the evidence base.", None),
  ("Then what exactly is the claimed edge?",
   "Three things, in descending order of confidence. First, consistency: every holding gets the "
   "same arithmetic, the same gates and the same disclosure, which removes the variance between "
   "reviewers rather than adding return. Second, downside screening: the balance-sheet and "
   "liquidity gates demonstrably kept the top decile out of the deepest drawdowns in-sample. Third, "
   "and least proven, the analyst overlay — the written human case that can rescue a mechanical "
   "Sell. That third one is the actual product and it is structurally untestable historically, "
   "because present-day judgment cannot be reconstructed point-in-time. It can only be proven "
   "forward.", "hard"),
  ("So the part you consider the real product is the part you cannot test?",
   "Yes, and that is an uncomfortable but accurate statement of where we are. It is also true of any "
   "research process that relies on human judgment. What we have done is make the judgment "
   "auditable rather than asking to be trusted: every override is written down, dated, attributed "
   "and versioned, so the forward record will attribute wins and losses to the human layer "
   "specifically rather than to the process as a whole. The forward paper test is the mechanism, "
   "and starting it does not require this approval.", "concede"),
  ("What is the forward test plan, concretely?",
   "Record today's full Ionic Score and call for the covered universe as a frozen baseline, then "
   "track realised outcomes against it with no revision to the historical record, reviewed at the "
   "existing monthly and quarterly cadences. The register pins the specification and the code "
   "version at entry, so a mid-test change to methodology starts a new version with a restarted "
   "clock rather than quietly improving the old result. That freeze discipline is already firm "
   "policy for strategies and applies here unchanged.", None),
 ]),
 ("C. The human layer", [
  ("If an analyst can override the score, is the score doing anything?",
   "The score does two things the analyst cannot. It originates every Sell, so nothing gets exited "
   "on a hunch; and it forces a written justification whenever a human disagrees with it, which is "
   "how the override register exists at all. The override is also one-directional, which is the key "
   "design decision — see the next question.", None),
  ("Why can an analyst turn a Sell into a Hold but not a Hold into a Sell?",
   "Because the two errors are not symmetric in cost or in likelihood. Talking yourself into keeping "
   "a name you like is the more common and more expensive failure, so we made optimism carry the "
   "burden of proof: rescuing a Sell requires a written case that goes on the record next to the "
   "name. Forcing a Sell that the engine did not originate is barred outright; the analyst escalates "
   "instead and a person with portfolio responsibility rules on it. The ingestion layer clamps any "
   "attempt to do it anyway and logs the attempt.", None),
  ("You have 126 escalations against 751 names. Isn't a one-in-six escalation rate a broken process?",
   "It is higher than we want and we treat it as a calibration problem rather than a badge of "
   "rigour. An escalation channel only functions if a human can read every item in it, and at one "
   "in six that is a real workload. Two things follow: the reviewer-facing workbook now exists so "
   "the queue is triageable by size and materiality rather than read end to end, and the analyst "
   "brief has been updated with the measured rate and explicit guidance to separate a genuine "
   "decision from recorded uncertainty. We expect that rate to fall on the next incremental cycle "
   "and we will report it either way.", "concede"),
  ("What stops analysts rubber-stamping the machine, or fabricating a source?",
   "Structurally: every load-bearing number requires a source that was actually opened and is listed "
   "per name, the recommendation must follow from the analyst's own written rationale (a bullish "
   "write-up under a Sell call is a named failure mode), and a self-review checklist runs per stock. "
   "Procedurally: the firm runs a periodic honesty probe that seeds a deliberately flawed claim to "
   "test whether dissent actually flows, and a red-team function whose job is to attack conclusions "
   "before they are relied upon. Neither eliminates the risk; both make it detectable.", None),
  ("The research is produced by language models. How is that defensible for client-facing work?",
   "By treating the model as a research assistant whose output is constrained, checked and "
   "attributed, not as an oracle. The constraints are explicit: a fixed output schema, a prohibition "
   "on inventing figures, a requirement to cite sources actually opened, a vocabulary restricted to "
   "Sell, Trim and Hold, and a one-directional override that prevents the model from originating an "
   "exit. The checks are mechanical and independent of the model. And nothing reaches a client "
   "without a named human countersignature. The committee should weigh this as a process question, "
   "and we would rather answer it directly than have it discovered.", "hard"),
 ]),
 ("D. The fund frameworks", [
  ("Why do you run two fund frameworks? Which one wins?",
   "They answer different questions and neither is allowed to win alone. The short-term framework "
   "ranks a fund on how much of its benchmark's rise it captures relative to how much of the fall, "
   "over a common six-month window across 181 funds. The long-term framework scores a curated 40 "
   "funds for multi-year holding. A fund is only recommended for sale when BOTH independently say "
   "sell; a positive signal on either side vetoes the sale and any disagreement defaults to Hold. "
   "The bar is deliberately high because switching a fund costs the client tax and time.", None),
  ("The long-term framework covers only 40 funds. What happens to everything else a client holds?",
   "It gets the short-term view only, or an explicit No View. We do not extend a curated ranking to "
   "funds it was not built for. In the sample client book, 24 of 37 schemes carried a real score and "
   "13 were reported as No View.", None),
  ("Hybrid, sectoral, index, liquid and debt funds have no framework at all. That is a large hole in a real client's book.",
   "Agreed, and it is disclosed on the page rather than papered over. Those categories sit outside "
   "both engines by design, and the deck prints No View with the reason instead of manufacturing a "
   "grade. In practice this is why the fund section of a review is honest but incomplete, and "
   "extending category coverage is on the roadmap. We would rather show a client a labelled gap "
   "than a fabricated verdict, but we accept the committee may want this closed before wider "
   "rollout.", "concede"),
  ("Why are funds reviewed twice a year when stocks are weekly?",
   "Because fund quality signals are slow and net-asset-value history only accrues monthly, so a "
   "weekly re-run would mostly re-express noise as change and generate switching recommendations "
   "that cost the client tax for no informational gain. The model runs at April and October ends, "
   "chosen from a study of formation-date pairs, with monthly data refresh in between so each run "
   "has a full window.", None),
  ("Is your fund benchmark correct?",
   "Not fully, and it is a known open item. The index sheet feeding the short-term framework is a "
   "price index rather than a total-return index, which flatters excess return by roughly 1.2 to "
   "1.5 per cent a year and therefore systematically UNDERSTATES sells. The direction of the error "
   "is conservative for a client, which is why it is not urgent, but it must be corrected before "
   "the next full model run and it is scheduled.", "concede"),
 ]),
 ("E. Client-facing and compliance", [
  ("Is this investment advice? What is the regulatory posture?",
   "It is a review of holdings a client already owns under a non-discretionary mandate. The "
   "vocabulary is Sell, Trim or Hold; there are no buy recommendations, no target prices and no "
   "implied return promises. Nothing is executed until the client authorises it, which every deck "
   "states with a signature line. Tax characterisations are labelled indicative with the client's "
   "own adviser named as the authority before dealing.", None),
  ("Clients will ask what to buy. What do we say?",
   "That this document does not answer that question, and why. Proceeds from recommended exits are "
   "shown parked in liquid or overnight instruments pending a separate conversation about goals and "
   "policy, and the deck says explicitly that no redeployment is assumed or recommended. That is a "
   "product boundary, not an oversight, and holding it is what keeps the review inside the "
   "non-discretionary framing.", None),
  ("Without target prices, is the client getting enough to act on?",
   "The client gets, per holding, a score, a call, a written case, the strongest counter-argument we "
   "could construct against our own view, a valuation judgment expressed as what the current "
   "multiple already assumes, and the condition that would change our mind. We consider that more "
   "actionable than a price target, which encodes a false precision about timing that we cannot "
   "support.", None),
  ("The tax numbers are estimates. What is our exposure if a client relies on them?",
   "Every tax figure is labelled indicative and preliminary on the page, with the instruction to "
   "confirm holding period, character and applicable rates with the client's tax adviser before "
   "dealing. Where the underlying data is missing we say so specifically rather than estimating "
   "silently: equity cost basis is not available from a consolidated account statement, so share-sale "
   "tax is disclosed as an estimate on approximately half of proceeds, while fund gains use the "
   "actual cost basis from the statement.", None),
  ("Client holdings are sitting in a code repository. What about confidentiality?",
   "The repository is private. Client data lives there because the deck is generated from it and "
   "reproducibility requires it, and every page is marked Private and Confidential. This is a real "
   "control question and the committee is right to ask it: the current answer is repository access "
   "control, and if the committee wants client data separated from code entirely, that is a change "
   "we should scope rather than defend against.", "hard"),
  ("Who signs off before a client sees this?",
   "The Principal or the CEO countersigns every deck, every time, with no standing exemption for a "
   "repeat client or a routine refresh. The scheduled pipeline is deliberately automated only as far "
   "as producing a marked draft. That gate is the control that stops an automated process ever "
   "speaking to a client on its own.", None),
 ]),
 ("F. Operations and risk", [
  ("What happens when the data pipeline breaks?",
   "It has, and the honest version of this answer is more useful than a reassuring one. Our option "
   "capture task stalled on 3 August: it logged a successful login and the full instrument universe, "
   "then wrote 2 files against 91 the previous session. It was caught by the daily end-of-day check "
   "and flagged, and the diagnosis exposed a weakness in the check itself — the task logs success "
   "BEFORE doing work, so a log-based health test reads green while nothing is captured. The check "
   "is being changed to assert on files written. That failure does not touch the review engine, "
   "which runs on client statements and the scored universe, but it is a fair illustration of the "
   "class of risk and of how we handle it.", "hard"),
  ("Is this dependent on one person?",
   "Partly, and reducing that is the reason the operating manual and the analyst brief exist as "
   "self-contained documents rather than as knowledge in someone's head. A new analyst receives one "
   "skill file and one workbook and can produce conforming output. The engine is scripted and "
   "rerunnable end to end. What remains person-dependent is judgment at the fund-manager and "
   "sign-off layers, which is appropriate, and the operational muscle memory for the data pipeline, "
   "which is genuinely thin.", "concede"),
  ("What does one review cost to produce, and how many clients can this serve?",
   "The expensive input is analyst research, and the design keeps that bounded rather than "
   "proportional to universe size: the weekly refresh routes names into full research only when "
   "earnings landed, a cheap delta look when only news moved, and no cost at all otherwise, so "
   "maintaining 751 names does not cost 751 research passes a week. Per client, deck assembly is "
   "scripted and effectively free once holdings are matched. The binding constraint on client count "
   "is sign-off and relationship-manager time, not compute.", None),
  ("If a client disputes a call a year from now, can we reconstruct why we said it?",
   "Yes. Each name carries a dated research file with the rationale, the sources opened and the "
   "forward growth estimate; the state file journals every change to a call with its trigger; and "
   "the deck is regenerated from that state, so the artefact and the reasoning cannot drift apart. "
   "Overrides record who moved what and why.", None),
  ("What is on the known-defects list right now?",
   "Fund expense ratios render from a placeholder rather than each scheme's real figure, which is "
   "visible on the scorecard pages and is first to close. Category coverage gaps in the fund "
   "frameworks. The price-versus-total-return benchmark issue. Equity cost basis unavailable for "
   "tax. Fund look-through built but not wired into every page. And roughly a thousand instances of "
   "internal field names appearing in analyst prose that can reach a client page, of which the "
   "exposed ones are fixed and the remainder is a scoped decision. All of these are in the deck.", None),
  ("Have you actually found real defects in your own work, or is the QA theatre?",
   "We have, and the deck devotes a page to them because a product review that shows only successes "
   "is not evidence. Concrete examples: a page asserted a concentration-limit breach that did not "
   "exist in the book while a correctly-computed page in the same deck said otherwise; fund "
   "name-similarity matching paired a mid-cap fund with a multi-cap fund and a liquid fund with an "
   "equity fund; a holdings row silently absorbed another holding's data during statement parsing; "
   "an unscored holding displayed as a score of zero, which is a real and very bad score; and a "
   "table row was covered by the panel beneath it while passing both geometry checks. Each was "
   "found by our own process before any client saw it, and each is closed in code rather than in a "
   "checklist.", None),
 ]),
 ("G. The uncomfortable ones", [
  ("This was largely built with AI agents. Who is accountable for a wrong call?",
   "The firm is, through the named human who signs the deck. That is why the signature requirement "
   "has no exemptions. The agents are constrained producers of research and layout whose output is "
   "schema-checked, source-required, machine-verified and human-countersigned; they hold no "
   "decision rights, and the one place they could originate a client action — a Sell — is reserved "
   "to a mechanical score with a written human confirmation on top.", "hard"),
  ("What would make you kill or pause this product?",
   "A forward record showing the analyst overlay destroys value relative to the mechanical score "
   "would remove the product's main justification. A confidentiality incident involving client "
   "holdings would pause it immediately. A regulatory reading that a holdings review of this "
   "specificity constitutes advice we are not licensed to give would stop it outright, which is why "
   "the vocabulary and the no-target-price rule are treated as permanent constraints and not "
   "stylistic preferences.", None),
  ("What are you asking us to approve, precisely?",
   "Adoption of this engine as the standard NDPMS review deliverable for onboarding and periodic "
   "reviews, on the frozen methodology, the four-layer quality gate and the countersignature "
   "requirement as described. Confirmation of the Sell, Trim and Hold vocabulary and the "
   "no-buy, no-target-price posture as permanent product constraints. Confirmation of the "
   "one-directional override rule and the escalation channel. And a direction on which of the "
   "known-defect items to close first.", None),
 ]),
]

STANDING = [
    ("Value weight in the 3-year score",
     "Proposal on the table is to raise Value from 18 to 25 per cent. Direction agreed; magnitude "
     "and funding source not settled. Because 60 per cent of the Value pillar is P/E, 25 per cent "
     "would put P/E alone at 15.0 per cent of the composite (from 10.8). The 7 points should come "
     "from Stage and Technical rather than Quality or Growth, since momentum has the weakest "
     "three-year justification. To be settled by a sensitivity sweep on the existing point-in-time "
     "harness judged on call-churn and drawdown stability, NOT on which weight produced the best "
     "return over 8 quarters, which would be curve-fitting. Requires quant-head and red-team "
     "sign-off, and it will move every score in the covered universe and break comparability with "
     "the shipped client books, so it should be a versioned amendment with a restarted forward clock."),
    ("Documentation correction already made",
     "The methodology previously described the horizon splits as 63/37 and 40/60. Against the "
     "shipped weights the fundamentals-versus-market split is 58/42 for the 3-year and 48/52 for "
     "the 1-year score. The old 63/37 reconciles to the pre-DCF-removal spec and was never updated "
     "when DCF was dropped; the 40/60 does not reconcile to either version. Corrected in the "
     "methodology with the weights left untouched. Quote 58/42 and 48/52."),
    ("Internal field names in client-visible prose",
     "Roughly a thousand instances across about 40 tokens. The instances actually exposed by a "
     "shipped deck are fixed; the rest is a decision between a render-layer translation, which is "
     "safer and protects every future client, and a supervised rewrite of the research corpus."),
]


def _p(doc, text, size=10, bold=False, color=INK, space_after=6, italic=False, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    return p


def build():
    doc = Document()
    K.apply_firm_styles(doc)
    K.add_title_page(
        doc,
        "NDPMS Portfolio Review Engine",
        "Product Approval Committee — anticipated questions and prepared rationale",
        AS_OF,
    )

    _p(doc, "How to use this brief", 13, bold=True, color=NAVY, space_after=4)
    _p(doc, "One question per block, with an answer written to be said out loud. Questions marked "
            "PRESS are the ones most likely to be pushed on a second time; questions marked CONCEDE "
            "are ones where the honest answer admits a real weakness, and the answer states the "
            "control that bounds it rather than arguing. Nothing here claims evidence we do not "
            "have. Where the truthful answer is that we cannot show something yet, it says so and "
            "says what we can show instead.", 10, space_after=14)

    n = 0
    for section, items in QA:
        K.add_heading(doc, section, level=1)
        for q, a, flag in items:
            n += 1
            tag = {"hard": "  [PRESS]", "concede": "  [CONCEDE]"}.get(flag, "")
            qp = doc.add_paragraph()
            rq = qp.add_run(f"Q{n}. {q}")
            rq.font.size = Pt(10.5)
            rq.font.bold = True
            rq.font.color.rgb = NAVY
            if tag:
                rt = qp.add_run(tag)
                rt.font.size = Pt(8)
                rt.font.bold = True
                rt.font.color.rgb = RED
            qp.paragraph_format.space_after = Pt(3)
            qp.paragraph_format.keep_with_next = True
            _p(doc, a, 10, space_after=12)

    doc.add_page_break()
    K.add_heading(doc, "Standing items the committee will be asked to direct", level=1)
    for title, body in STANDING:
        _p(doc, title, 10.5, bold=True, color=NAVY, space_after=3)
        _p(doc, body, 10, space_after=12)

    _p(doc, f"Internal committee preparation document. {n} questions. Prepared {AS_OF}. "
            "Figures quoted are as of that date and traceable to the methodology, the "
            "point-in-time backtest report, and the coverage counts in the product deck.",
       8.5, italic=True, color=SLATE, space_after=0)

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    path = OUT
    for attempt in range(3):
        try:
            doc.save(path)
            break
        except PermissionError:
            path = OUT.replace(".docx", f"_v{attempt + 2}.docx")
    print(f"Saved {n} Q&A across {len(QA)} sections -> {path}")
    return path


if __name__ == "__main__":
    build()
