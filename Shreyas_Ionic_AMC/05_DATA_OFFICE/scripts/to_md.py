"""to_md.py — convert docx/xlsx/csv/parquet(/pdf if pypdf present) to lean Markdown.
TOKEN SAVER: agents read the .md digest instead of binary-parsing or dumping raw files.
Usage: python to_md.py <input> [output.md] [--rows N (tables, default 40)] [--sheet NAME]
Proxy-proof: uses only libs already on this machine (python-docx, openpyxl, pandas).
"""
import sys
from pathlib import Path

import pandas as pd


def docx_to_md(p: Path) -> str:
    from docx import Document
    doc = Document(str(p))
    out = [f"# {p.stem} (converted from docx)"]
    for para in doc.paragraphs:
        t = para.text.strip()
        if not t:
            continue
        style = (para.style.name or "").lower()
        if "heading" in style:
            lvl = "".join(c for c in style if c.isdigit()) or "2"
            out.append("#" * min(int(lvl) + 1, 6) + " " + t)
        else:
            out.append(t)
    for i, tbl in enumerate(doc.tables):
        out.append(f"\n**Table {i+1}:**\n")
        rows = [[c.text.strip().replace("\n", " ") for c in r.cells] for r in tbl.rows]
        if rows:
            out.append("| " + " | ".join(rows[0]) + " |")
            out.append("|" + "---|" * len(rows[0]))
            out += ["| " + " | ".join(r) + " |" for r in rows[1:]]
    return "\n".join(out)


def table_to_md(p: Path, rows: int, sheet=None) -> str:
    if p.suffix.lower() == ".csv":
        df = pd.read_csv(p)
        parts = {None: df}
    elif p.suffix.lower() == ".parquet":
        parts = {None: pd.read_parquet(p)}
    else:  # xlsx
        xls = pd.read_excel(p, sheet_name=sheet)  # None -> all sheets (dict)
        parts = xls if isinstance(xls, dict) else {sheet: xls}
    out = [f"# {p.name} (converted; showing head {rows} rows per sheet)"]
    for name, df in parts.items():
        if name:
            out.append(f"\n## Sheet: {name}")
        out.append(f"shape: {df.shape[0]} rows x {df.shape[1]} cols | columns: {list(df.columns)}")
        out.append(df.head(rows).to_markdown(index=False))
        num = df.select_dtypes("number")
        if len(num.columns):
            out.append("\n**Numeric summary:**\n" + num.describe().round(3).to_markdown())
    return "\n".join(out)


def pdf_to_md(p: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        return f"# {p.name}\n(pypdf not installed — cannot convert PDF; read directly or install pypdf)"
    r = PdfReader(str(p))
    return f"# {p.stem} (from pdf, {len(r.pages)} pages)\n\n" + "\n\n".join(
        f"## Page {i+1}\n{pg.extract_text() or '(no text)'}" for i, pg in enumerate(r.pages))


if __name__ == "__main__":
    args = [a for a in sys.argv[1:] if not a.startswith("--")]
    rows = 40
    sheet = None
    for i, a in enumerate(sys.argv):
        if a == "--rows":
            rows = int(sys.argv[i + 1])
        if a == "--sheet":
            sheet = sys.argv[i + 1]
    src = Path(args[0])
    dst = Path(args[1]) if len(args) > 1 else src.with_suffix(".md")
    ext = src.suffix.lower()
    if ext == ".docx":
        md = docx_to_md(src)
    elif ext in (".xlsx", ".xls", ".csv", ".parquet"):
        md = table_to_md(src, rows, sheet)
    elif ext == ".pdf":
        md = pdf_to_md(src)
    else:
        sys.exit(f"unsupported: {ext}")
    dst.write_text(md, encoding="utf-8")
    print(f"wrote {dst}  ({len(md)//1000}k chars from {src.stat().st_size//1024}KB source)")
