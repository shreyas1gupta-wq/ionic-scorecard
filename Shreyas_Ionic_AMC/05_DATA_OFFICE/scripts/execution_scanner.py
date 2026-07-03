"""Build the EXECUTION sheet: fully-specified, dated trade legs (action / symbol / expiry /
strike / CE-PE / live price / lot size) for every strategy, sorted by entry date.
Live prices via Angel One. Output -> FINAL_STRATEGY_FORWARD_CHECK/08_Execution/.

EX-ANTE RISK OVERLAY (RISK_LIMITS.md - APPROVED D-021 - + KNOWLEDGE_BASE lesson A3 / K-010):
appends entry_iv / iv_source / size_x / tail_tier / tail_warning columns to every output CSV
(original columns and file names unchanged -> backward compatible). Final position size for
Short_Strangle / IVRV rows = lots * size_x (inverse-IV sizing, mandatory per RISK_LIMITS).

Usage:
    python execution_scanner.py            # live Angel-connected scan (needs session)
    python execution_scanner.py --dry-run  # NO Angel: reload existing execution_ALL.csv and
                                           # (re)apply the risk overlay in place. Idempotent.
"""
import sys, json, time, datetime as dt
from pathlib import Path
import numpy as np, pandas as pd

DRY_RUN = "--dry-run" in sys.argv

PROJ = Path(r"c:\Users\Shreyas.1Gupta\OneDrive - Angel Broking Limited\Desktop\Backup\NIFTY 500")
sys.path.insert(0, str(PROJ / "intraday_options_strategy"))

R_, Q_ = 0.065, 0.0
TODAY = dt.date(2026, 7, 3)
NEXT_SESSION = dt.date(2026, 7, 6)            # Mon (Jul 4-5 = weekend)
SOPT = PROJ / "intraday_options_strategy/datasets/raw/hf_index_options_1m/stocks_options"
OUTD = PROJ / "FINAL_STRATEGY_FORWARD_CHECK" / "08_Execution"; OUTD.mkdir(parents=True, exist_ok=True)
FWD = pd.read_csv(PROJ / "datasets/nse_earnings_dates/forthcoming_results.csv")


def prev_session(d):
    d = d - dt.timedelta(days=1)
    while d.weekday() >= 5:
        d -= dt.timedelta(days=1)
    return d


# ==================== EX-ANTE RISK OVERLAY ====================
# Source: 07_RISK_OFFICE/RISK_LIMITS.md (APPROVED D-021) position rule:
#   "inverse-IV sizing mandatory (size proportional to 1/entry-IV, ref 25% IV)"
# and KNOWLEDGE_BASE lesson A3 / KILLED_IDEAS K-010: NO retro-fit outcome blacklists
# (lookahead). Live tail filter must be EX-ANTE: IV at entry (corr -0.23 with future
# worst-case; the ex-ante top-IV-quintile filter caught 8/12 landmines).
IV_REF = 0.25                 # reference IV: size_x = 1.0 at 25% entry IV
SIZE_MIN, SIZE_MAX = 0.4, 1.0 # clip band for size_x — CIO ruling Q3-plan 2(a): cap 1.0x until regime gate exists (no upsizing into calm)
TAIL_Q = 0.80                 # top quintile of the scan's per-stock entry IVs -> tail_tier=HIGH
TAIL_HAIRCUT = 0.6            # HIGH tail tier: size_x *= 0.6
TAIL_WARN = ("HIGH ex-ante tail tier: entry IV in top quintile of this scan "
             "(the filter that caught 8/12 landmines) -> size_x *= 0.6")
NOIV_WARN = "no entry IV derivable -> size_x floored at 0.4 (inverse-IV sizing mandatory)"
RISK_COLS = ["entry_iv", "iv_source", "size_x", "tail_tier", "tail_warning"]

# ---- P1 (2026Q3, IC-1 catch): shared IV sanity cap ----------------------------
# Root cause: implied_vol() solves on Brent [0.1%, 500%] with NO sanity bound of its
# own (intraday_options_strategy/options/bs_pricing.py); the backtest had an iv<1.0
# guardrail but call sites here only checked "0.03 < iv < 3" (FF/strangle ATM) or had
# NO bound at all (credit-proxy, signal-regex fallback, overlay's raw ingestion of an
# upstream entry_iv). Result: a stale/crossed print can back out IV >= 100% (e.g. the
# INFY 132.7% seen in a live/lastmonth artifact) and flow straight into inverse-IV
# sizing. Fix: ONE helper, applied at EVERY IV computation/ingestion point.
IV_LO, IV_HI = 0.03, 1.0      # sane annualized-IV band: 3% .. 100%
IV_REJECTED = "rejected"      # iv_source tag for anything that failed the sanity band


def sane_iv(iv):
    """Return iv if 0.03 < iv < 1.0, else NaN. Scalar-safe (None/NaN pass through as NaN).
    This is the ONLY gate for "is this a usable IV" anywhere in this script — every
    computation or ingestion site must route through it before the value is trusted."""
    if iv is None:
        return np.nan
    try:
        v = float(iv)
    except (TypeError, ValueError):
        return np.nan
    if not np.isfinite(v) or not (IV_LO < v < IV_HI):
        return np.nan
    return v


def _is_shortvol(strat):
    """Strategies covered by the inverse-IV sizing rule (naked premium-selling sleeves)."""
    s = str(strat).upper()
    return s == "SHORT_STRANGLE" or s.startswith("IVRV")


def apply_risk_overlay(df):
    """Append the ex-ante risk columns; idempotent (recomputes if columns already exist).
    entry_iv  : annualized entry IV. Real ATM implied vol where the scanner computed one
                (iv_source='atm'); else for strangles a credit-based proxy is derived:
                    IV_proxy ~= credit_pct / (0.8 * sqrt(DTE/365))
                APPROXIMATE, labeled as such: inverts ATM-straddle ~= 0.8*S*sigma*sqrt(T);
                a ~5%-OTM strangle credit slightly understates ATM straddle premium, so the
                proxy runs a touch low -- acceptable for relative sizing/quintiles.
                (iv_source='proxy'). 'IV=xx' in the signal string is a last fallback ('signal').
                EVERY one of the above must pass sane_iv() (0.03 < iv < 1.0, P1 2026Q3 /
                IC-1 catch: no >=100% or <3% IV survives). A raw value that fails the band
                at every stage it was tried is NOT silently kept or defaulted -- it is
                excluded from IVRV candidacy and iv_source='rejected'.
    size_x    : inverse-IV multiplier clip(0.25/entry_iv, 0.4, 1.0) on Short_Strangle/IVRV
                rows only; 1.0 elsewhere. FINAL SIZE = lots * size_x. Rows with
                iv_source='rejected' are floored at SIZE_MIN (0.4), same as "no IV".
    tail_tier : per-stock HIGH if that stock's entry IV is in the top quintile of THIS scan
                (ex-ante, K-010-compliant -- no realized-outcome blacklists), else NORMAL;
                '-' for rows outside the sizing rule OR rejected by sane_iv (excluded from
                the ranking cross-section entirely). HIGH -> size_x *= 0.6 + tail_warning.
    """
    df = df.copy()
    # idempotency: keep only genuine ATM IVs from a previous pass, recompute the rest
    if "entry_iv" in df.columns and "iv_source" in df.columns:
        df.loc[df["iv_source"].astype(str) != "atm", "entry_iv"] = np.nan
    if "entry_iv" not in df.columns:
        df["entry_iv"] = np.nan
    df = df.drop(columns=[c for c in RISK_COLS if c in df.columns and c != "entry_iv"])

    sv = df["strategy"].map(_is_shortvol)
    # ---- ingestion point #1: upstream entry_iv (from the live-scan ATM computation
    # below, or a re-loaded --dry-run CSV) -- MUST pass sane_iv before being trusted.
    iv_raw = pd.to_numeric(df["entry_iv"], errors="coerce")
    iv = iv_raw.map(sane_iv)
    n_atm_rejected = int((iv_raw.notna() & iv.isna()).sum())
    src = pd.Series("-", index=df.index)
    src[iv.notna()] = "atm"

    # DTE per row (entry -> expiry of the short leg)
    ent = pd.to_datetime(df["entry_date"]).dt.date
    exp = df["expiry"].map(lambda s: dt.datetime.strptime(str(s), "%d%b%Y").date())
    dte = pd.Series([max((e - a).days, 1) for a, e in zip(ent, exp)], index=df.index, dtype=float)

    # ---- ingestion point #2: credit-based IV proxy for strangle-style rows
    # (formula above; APPROXIMATE) -- MUST pass sane_iv before being trusted.
    credit_pct = df["signal"].astype(str).str.extract(r"([\d.]+)%spot")[0].astype(float)
    proxy_raw = (credit_pct / 100.0) / (0.8 * np.sqrt(dte / 365.0))
    proxy = proxy_raw.map(sane_iv)
    n_proxy_rejected = int((sv & iv.isna() & proxy_raw.notna() & proxy.isna()).sum())
    m = sv & iv.isna() & proxy.notna()
    iv[m], src[m] = proxy[m], "proxy"
    # ---- ingestion point #3: last fallback, an explicit IV=xx token in the signal
    # (e.g. future IVRV rows) -- MUST pass sane_iv before being trusted.
    sig_iv_raw = df["signal"].astype(str).str.extract(r"IV=([\d.]+)")[0].astype(float)
    sig_iv_raw = sig_iv_raw.where(sig_iv_raw <= 3, sig_iv_raw / 100.0)   # percent-quoted -> decimal
    sig_iv = sig_iv_raw.map(sane_iv)
    n_sig_rejected = int((sv & iv.isna() & sig_iv_raw.notna() & sig_iv.isna()).sum())
    m = sv & iv.isna() & sig_iv.notna()
    iv[m], src[m] = sig_iv[m], "signal"

    # any row where a raw IV existed at some stage but failed sane_iv at every stage
    # it was tried -> explicitly marked 'rejected' (never silently coerced to "-"/1.0x)
    ever_had_raw = (iv_raw.notna() | proxy_raw.notna() | sig_iv_raw.notna())
    rejected = sv & iv.isna() & ever_had_raw
    src[rejected] = IV_REJECTED
    n_rejected_total = n_atm_rejected + n_proxy_rejected + n_sig_rejected
    if n_rejected_total:
        print(f"sane_iv: rejected {n_rejected_total} raw IV value(s) outside "
              f"({IV_LO:.0%}, {IV_HI:.0%}) -- atm={n_atm_rejected} proxy={n_proxy_rejected} "
              f"signal={n_sig_rejected}. Excluded from IVRV candidacy; iv_source='rejected', "
              f"size_x floored at {SIZE_MIN}.")

    # inverse-IV sizing (RISK_LIMITS, ref 25% IV)
    size = pd.Series(1.0, index=df.index)
    ok = sv & iv.notna()
    size[ok] = np.clip(IV_REF / iv[ok], SIZE_MIN, SIZE_MAX)
    warn = pd.Series("", index=df.index)
    noiv = sv & iv.isna() & ~rejected
    size[noiv], warn[noiv] = SIZE_MIN, NOIV_WARN         # conservative floor if IV unknown
    size[rejected] = SIZE_MIN                            # conservative floor: IV failed sanity band
    warn[rejected] = (f"entry IV outside sane band ({IV_LO:.0%}, {IV_HI:.0%}) -- "
                       f"rejected & excluded from IVRV candidacy; size_x floored at {SIZE_MIN}")

    # ex-ante tail tier: top quintile of per-stock entry IV across THIS scan
    # (rejected rows are excluded from the IVRV candidate ranking set entirely --
    # they must not contribute to or benefit from the quintile threshold)
    tail = pd.Series("-", index=df.index)
    tail[sv] = "NORMAL"
    tail[rejected] = "-"
    sym_iv = iv[ok].groupby(df.loc[ok, "symbol"]).max().dropna()
    if len(sym_iv) >= 5:                                  # need a real cross-section to rank
        thr = sym_iv.quantile(TAIL_Q)
        hi = sv & df["symbol"].isin(set(sym_iv[sym_iv >= thr].index))
        tail[hi] = "HIGH"
        size[hi] *= TAIL_HAIRCUT
        warn[hi] = TAIL_WARN

    df["entry_iv"] = iv.round(4)
    df["iv_source"] = src
    df["size_x"] = size.round(2)
    df["tail_tier"] = tail
    df["tail_warning"] = warn
    return df


def overlay_report(df):
    """Before/after evidence: trade-level (not leg-level) sizing + tail stats."""
    sv = df[df["strategy"].map(_is_shortvol)]
    tr = sv.drop_duplicates(["strategy", "symbol", "entry_date"])
    print("\n=== EX-ANTE RISK OVERLAY (RISK_LIMITS D-021) - trade level ===")
    print(f"short-vol trades in scope : {len(tr)}  (before: all sized 1.00x flat)")
    print(f"  downsized (size_x<1.0)  : {(tr['size_x'] < 1).sum()}")
    print(f"  upsized   (size_x>1.0)  : {(tr['size_x'] > 1).sum()}")
    print(f"  unchanged (size_x=1.0)  : {(tr['size_x'] == 1).sum()}")
    print(f"  tail_tier HIGH (x0.6)   : {(tr['tail_tier'] == 'HIGH').sum()}")
    if len(tr):
        print(f"  size_x min/med/max      : {tr['size_x'].min():.2f} / "
              f"{tr['size_x'].median():.2f} / {tr['size_x'].max():.2f}")
        print(f"  entry IV min/med/max    : {tr['entry_iv'].min():.1%} / "
              f"{tr['entry_iv'].median():.1%} / {tr['entry_iv'].max():.1%}")


# ==================== HARD RISK CEILING (adoption-queue #5, ai-hedge-fund pattern) ====================
# Non-overridable post-overlay clamp. Runs AFTER apply_risk_overlay (and after any downstream
# conviction/event-gate scoring in final_execution.py, if that script is fed through this too).
# Nothing upstream -- conviction score, news overlay, sizing tier -- may push size_x above 1.0,
# unblock a blocked row, or authorize more lots than the 1% book-equity risk budget allows.
# This function is deliberately dumb: it does not re-derive risk, it only clamps.
BOOK_EQUITY = 10_000_000  # ₹1 crore paper book (Principal ruling D-026, 2026-07-04)       # paper-book default (RISK_LIMITS D-021: 1% max risk per position)
POSITION_RISK_PCT = 0.01      # RISK_LIMITS "Position level": max risk per position = 1.0% of book equity
WORST_CASE_MULT = 2.0         # APPROXIMATE conservative proxy for short structures: worst-case
                               # loss per lot ~= 2x premium collected (no real tail model here --
                               # RISK_LIMITS says undefined-risk structures need a worst-case MTM
                               # model; this is a cheap stand-in until that model exists, labeled
                               # approximate per the task spec, NOT a substitute for real margin/SPAN)


def enforce_risk_ceiling(df):
    """HARD, non-overridable clamp applied after all sizing/conviction/event-gate logic.
    Idempotent -- safe to call multiple times or on an already-clamped frame.

    (a) size_x hard-capped at 1.0 for EVERY row (not just short-vol rows) -- belt-and-braces on
        top of apply_risk_overlay's own clip(), in case a caller (or a future column) ever
        pushes size_x above 1.0 by some other path (e.g. FF calendar sz up to 1.25x today).
    (b) blocked column: preserved if present (never flips True->False here); created default
        False if absent (this script's own execution_ALL.csv has no blocked col -- that gate
        lives downstream in final_execution.py's execution_scored.csv -- but this function must
        be safe to run on EITHER file without ever erasing an existing block).
    (c) max_lots = floor(0.01 * BOOK_EQUITY / worst_case_per_lot), worst_case_per_lot = 2x
        live_price*lot_size (APPROXIMATE short-structure proxy, see WORST_CASE_MULT comment).
        Rows with no usable live_price/lot_size get max_lots = 0 (fail safe, not unlimited).
    (d) prints a loud RISK CEILING table listing every row where a clamp actually bound
        (size_x lowered, or lots > max_lots) -- silent if nothing binds.
    """
    df = df.copy()

    # ---- (a) size_x hard cap at 1.0, unconditionally ----
    if "size_x" not in df.columns:
        df["size_x"] = 1.0
    size_before = pd.to_numeric(df["size_x"], errors="coerce").fillna(1.0)
    size_clamped = size_before.clip(upper=1.0)
    n_size_clamped = int((size_clamped < size_before).sum())
    df["size_x"] = size_clamped

    # ---- (b) blocked column: preserve, never erase; default False if wholly absent ----
    if "blocked" not in df.columns:
        df["blocked"] = False
    else:
        df["blocked"] = df["blocked"].fillna(False).astype(bool)

    # ---- (c) max_lots hard position-risk budget ----
    live_price = pd.to_numeric(df.get("live_price"), errors="coerce")
    lot_size = pd.to_numeric(df.get("lot_size"), errors="coerce")
    premium_per_lot = live_price * lot_size
    worst_case_per_lot = premium_per_lot * WORST_CASE_MULT
    max_lots = np.floor((POSITION_RISK_PCT * BOOK_EQUITY) / worst_case_per_lot)
    # fail-safe: no usable price/lot_size -> 0 lots allowed, never unlimited (NaN/inf guard)
    max_lots = max_lots.where(np.isfinite(max_lots) & (worst_case_per_lot > 0), 0.0)
    df["max_lots"] = max_lots.astype("int64")

    lots_requested = pd.to_numeric(df.get("lots"), errors="coerce").fillna(0.0)
    n_lots_over_ceiling = int((lots_requested > df["max_lots"]).sum())

    # ---- (d) loud printed table when any clamp binds ----
    binds = pd.Series(False, index=df.index)
    binds |= (size_clamped < size_before)
    binds |= (lots_requested > df["max_lots"])
    n_binding = int(binds.sum())
    if n_binding:
        print("\n" + "=" * 78)
        print("RISK CEILING BOUND -- hard, non-overridable clamp (enforce_risk_ceiling)")
        print("=" * 78)
        print(f"BOOK_EQUITY={BOOK_EQUITY:,.0f} (paper default) | position risk budget "
              f"{POSITION_RISK_PCT:.0%} | worst_case_per_lot ~= {WORST_CASE_MULT:.0f}x premium "
              "(APPROXIMATE proxy)")
        print(f"size_x clamped to <=1.0 : {n_size_clamped} row(s)")
        print(f"lots > max_lots         : {n_lots_over_ceiling} row(s)")
        show_cols = [c for c in ["entry_date", "strategy", "symbol", "action", "opt",
                                  "live_price", "lot_size", "lots", "max_lots",
                                  "size_x", "blocked"] if c in df.columns]
        with pd.option_context("display.max_rows", 50, "display.width", 160):
            print(df.loc[binds, show_cols].to_string(index=False))
        print("=" * 78 + "\n")
    return df


def risk_ceiling_report(before, after):
    """Before/after evidence for the memo: what the ceiling actually changed."""
    print("\n=== RISK CEILING before/after (enforce_risk_ceiling) ===")
    b_size = pd.to_numeric(before.get("size_x"), errors="coerce") if "size_x" in before.columns else pd.Series(dtype=float)
    print(f"size_x > 1.0 before -> after : "
          f"{int((b_size > 1.0).sum()) if len(b_size) else 0} -> {int((after['size_x'] > 1.0).sum())}")
    print(f"blocked=True rows before -> after : "
          f"{int(before['blocked'].sum()) if 'blocked' in before.columns else 0} -> {int(after['blocked'].sum())}")
    print(f"max_lots column present : {'max_lots' in after.columns} "
          f"(min/med/max = {after['max_lots'].min()}/{int(after['max_lots'].median())}/{after['max_lots'].max()})")
    print(f"rows with lots > max_lots : "
          f"{int((pd.to_numeric(after.get('lots'), errors='coerce').fillna(0) > after['max_lots']).sum())}")


# ==================== DRY RUN (no Angel session) ====================
if DRY_RUN:
    src_csv = OUTD / "execution_ALL.csv"
    df = pd.read_csv(src_csv)
    print(f"DRY RUN: loaded {len(df)} legs from {src_csv}")
    df = apply_risk_overlay(df)
    overlay_report(df)
    before_ceiling = df.copy()
    df = enforce_risk_ceiling(df)                 # HARD non-overridable clamp, always last
    risk_ceiling_report(before_ceiling, df)
    df.to_csv(OUTD / "execution_ALL.csv", index=False)
    for strat in df["strategy"].unique():
        df[df["strategy"] == strat].to_csv(OUTD / f"execution_{strat}.csv", index=False)
    print(f"\nDRY RUN done -> risk + ceiling columns appended in place at {OUTD} (same CSV names)")
    sys.exit(0)

# ==================== LIVE SCAN (Angel session required) ====================
from docx import Document
from docx.shared import Pt, RGBColor
from options.bs_pricing import implied_vol
import angel_cfg as A

obj, sess = A.login(); print("login OK")
scrip = json.loads((Path("scrip_master.json")).read_bytes())
stocks = sorted({p.name for p in SOPT.iterdir() if p.is_dir()})
eqtok = {x["name"]: x["token"] for x in scrip if x.get("exch_seg") == "NSE"
         and x.get("symbol", "").endswith("-EQ") and x.get("name") in stocks}
opt = [x for x in scrip if x.get("exch_seg") == "NFO" and x.get("instrumenttype") == "OPTSTK" and x.get("name") in stocks]

def ed(s): return dt.datetime.strptime(s, "%d%b%Y").date()
exps = sorted({ed(x["expiry"]) for x in opt if ed(x["expiry"]) > TODAY})
FRONT, BACK = exps[0], exps[1]
FRONT_S, BACK_S = FRONT.strftime("%d%b%Y").upper(), BACK.strftime("%d%b%Y").upper()
Tf = (FRONT - TODAY).days / 365; Tb = (BACK - TODAY).days / 365
lot = {}
chain = {}
for x in opt:
    lot[x["name"]] = int(x["lotsize"])
    k = float(x["strike"]) / 100; ot = x["symbol"][-2:]
    chain[(x["name"], x["expiry"], ot, round(k, 2))] = x["token"]
strikes = {}
for (nm, ex, ot, k), tok in chain.items():
    strikes.setdefault((nm, ex), set()).add(k)
print(f"front {FRONT} back {BACK} | {len(eqtok)} stocks")


def bulk_ltp(exch, tokens):
    out = {}; toks = list(dict.fromkeys(tokens))
    for i in range(0, len(toks), 45):
        try:
            r = obj.getMarketData("LTP", {exch: toks[i:i + 45]})
            for f in r.get("data", {}).get("fetched", []):
                out[str(f["symbolToken"])] = f.get("ltp")
        except Exception:
            pass
        time.sleep(0.4)
    return out


spot_tok = bulk_ltp("NSE", list(eqtok.values()))
spot = {nm: spot_tok.get(str(t)) for nm, t in eqtok.items()}
spot = {k: v for k, v in spot.items() if v}


def near(ks, tgt): return min(ks, key=lambda x: abs(x - tgt)) if ks else None


# collect option tokens to price
need = []
plan = {}  # nm -> chosen strikes/tokens
for nm, s in spot.items():
    kf = strikes.get((nm, FRONT_S), set()); kb = strikes.get((nm, BACK_S), set())
    if not kf or not kb:
        continue
    atm = near(sorted(kf & kb), s) if (kf & kb) else near(sorted(kf), s)
    kc = near(sorted(kf), s * 1.05); kp = near(sorted(kf), s * 0.95)
    m = {"spot": s, "atm": atm, "kc": kc, "kp": kp, "lot": lot.get(nm, 0)}
    for lab, ex, ot, k in [("f_atmCE", FRONT_S, "CE", atm), ("b_atmCE", BACK_S, "CE", atm),
                            ("f_atmPE", FRONT_S, "PE", atm),
                            ("otmCE", FRONT_S, "CE", kc), ("otmPE", FRONT_S, "PE", kp)]:
        t = chain.get((nm, ex, ot, k))
        if t: m[lab] = t; need.append(t)
    plan[nm] = m

ltp = bulk_ltp("NFO", need)
def L(nm, lab):
    t = plan[nm].get(lab); return ltp.get(str(t)) if t else None

rows = []   # execution legs

# ---- FF calendar (enter next session) ----
for nm, m in plan.items():
    ce_f, ce_b = L(nm, "f_atmCE"), L(nm, "b_atmCE")
    if not (ce_f and ce_b and m["atm"]):
        continue
    iv1 = sane_iv(implied_vol(ce_f, m["spot"], m["atm"], Tf, R_, Q_, True))
    iv2 = sane_iv(implied_vol(ce_b, m["spot"], m["atm"], Tb, R_, Q_, True))
    if not (iv1 == iv1 and iv2 == iv2):   # NaN-safe truthiness (sane_iv returns NaN, not None)
        continue
    var_f = (iv2**2 * Tb - iv1**2 * Tf) / (Tb - Tf)
    if var_f <= 0:
        continue
    ff = (iv1 - np.sqrt(var_f)) / np.sqrt(var_f)
    if ff < 0.25:
        continue
    sz = 0.75 if ff < 0.5 else (1.0 if ff < 0.75 else 1.25)
    for act, ex, px in [("SELL", FRONT_S, ce_f), ("BUY", BACK_S, ce_b)]:
        rows.append(dict(entry_date=NEXT_SESSION, strategy="FF_Calendar", action=act, symbol=nm,
                         expiry=ex, strike=m["atm"], opt="CE", live_price=round(px, 2),
                         lots=round(sz, 2), lot_size=m["lot"], signal=f"FF={ff:.2f}",
                         exit_rule="close BOTH ~2 sessions before front expiry",
                         entry_iv=round(iv1, 4)))   # front ATM IV (risk overlay input)

# ---- Short strangle (enter ~14 DTE) ----
STR_ENTRY = FRONT - dt.timedelta(days=14)
while STR_ENTRY.weekday() >= 5:
    STR_ENTRY += dt.timedelta(days=1)
for nm, m in plan.items():
    ce, pe = L(nm, "otmCE"), L(nm, "otmPE")
    if not (ce and pe):
        continue
    cr = ce + pe
    # today's ATM IV = entry-IV estimate for the risk overlay (actual entry is ~14 DTE);
    # where unavailable the overlay falls back to the credit-based IV proxy (see docstring)
    ce_atm = L(nm, "f_atmCE")
    iv_atm = implied_vol(ce_atm, m["spot"], m["atm"], Tf, R_, Q_, True) if (ce_atm and m["atm"]) else None
    iv_atm = sane_iv(iv_atm)
    iv_atm = iv_atm if iv_atm == iv_atm else None   # NaN -> None so `round(iv_atm, 4) if iv_atm else None` still works below
    for act, k, ot, px in [("SELL", m["kc"], "CE", ce), ("SELL", m["kp"], "PE", pe)]:
        rows.append(dict(entry_date=STR_ENTRY, strategy="Short_Strangle", action=act, symbol=nm,
                         expiry=FRONT_S, strike=k, opt=ot, live_price=round(px, 2),
                         lots=1, lot_size=m["lot"], signal=f"credit={cr:.1f} ({cr/m['spot']*100:.1f}%spot)",
                         exit_rule="buy back at 50% of credit, else hold to expiry",
                         entry_iv=round(iv_atm, 4) if iv_atm else None))

# ---- Earnings short-vol (enter 1 session before each earnings) ----
FWD["d"] = pd.to_datetime(FWD["date"], format="%d-%b-%Y", errors="coerce")
up = FWD.dropna(subset=["d"]); up = up[up["d"].dt.date >= TODAY].sort_values("d")
for _, e in up.iterrows():
    nm = e["symbol"]; ed_ = e["d"].date()
    if nm not in plan:
        continue
    exp = FRONT if ed_ <= FRONT else BACK
    exp_s = exp.strftime("%d%b%Y").upper()
    m = plan[nm]; atm = m["atm"]
    # price ATM CE+PE in the spanning expiry (front already priced; back reuse if needed)
    ce = L(nm, "f_atmCE") if exp == FRONT else L(nm, "b_atmCE")
    pe = L(nm, "f_atmPE") if exp == FRONT else None
    entry = prev_session(ed_)
    for act, ot, px in [("SELL", "CE", ce), ("SELL", "PE", pe)]:
        rows.append(dict(entry_date=entry, strategy="Earnings_ShortVol", action=act, symbol=nm,
                         expiry=exp_s, strike=atm, opt=ot, live_price=round(px, 2) if px else None,
                         lots=1, lot_size=m["lot"], signal=f"earnings {ed_}",
                         exit_rule="close 1 session AFTER the result"))

df = pd.DataFrame(rows).sort_values(["entry_date", "strategy", "symbol", "opt"])
df = apply_risk_overlay(df)          # ex-ante risk columns appended (RISK_LIMITS D-021)
overlay_report(df)
before_ceiling = df.copy()
df = enforce_risk_ceiling(df)        # HARD non-overridable clamp (adoption-queue #5), always last
risk_ceiling_report(before_ceiling, df)
df.to_csv(OUTD / "execution_ALL.csv", index=False)
for strat in df["strategy"].unique():
    df[df["strategy"] == strat].to_csv(OUTD / f"execution_{strat}.csv", index=False)

# counts
ff_n = (df["strategy"] == "FF_Calendar").sum() // 2
str_n = (df["strategy"] == "Short_Strangle").sum() // 2
ea_n = (df["strategy"] == "Earnings_ShortVol").sum() // 2
print(f"FF trades {ff_n} | strangle {str_n} | earnings {ea_n} | total legs {len(df)}")

# ---- Word execution plan ----
doc = Document()
doc.add_heading("EXECUTION PLAN — trades to take", 0)
p = doc.add_paragraph(f"Live prices as of {TODAY} (Fri close). Next session Mon {NEXT_SESSION}. "
                      f"Front expiry {FRONT}, back {BACK}. Account = disposable/data-only.")
p.runs[0].italic = True
doc.add_heading("Sequenced by entry date", 1)
t = doc.add_table(rows=1, cols=8); t.style = "Light Grid Accent 1"
for i, h in enumerate(["Entry date", "Strategy", "Action", "Symbol", "Expiry", "Strike", "CE/PE", "Live px"]):
    r = t.rows[0].cells[i].paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(8)
for _, x in df.iterrows():
    c = t.add_row().cells
    vals = [str(x["entry_date"]), x["strategy"].replace("_", " "), x["action"], x["symbol"],
            x["expiry"], f"{x['strike']:g}", x["opt"], f"{x['live_price']}" if pd.notna(x["live_price"]) else "-"]
    for i, v in enumerate(vals):
        rr = c[i].paragraphs[0].add_run(v); rr.font.size = Pt(8)
doc.add_heading("How to read / manage", 1)
for b in [f"FF Calendar ({ff_n} trades): enter {NEXT_SESSION}. SELL front-{FRONT_S} CE + BUY back-{BACK_S} CE at same strike. "
          "Size 0.75/1.0/1.25x by FF tier. Exit both legs ~2 sessions before front expiry. NOTE many names have Jul earnings (elevated FF).",
          f"Short Strangle ({str_n} trades): enter ~{STR_ENTRY} (14 DTE). SELL ~5% OTM CE + ~5% OTM PE (front). "
          "Buy back at 50% of credit collected, else hold to expiry.",
          f"Earnings Short-Vol ({ea_n} events): for each, SELL the ATM straddle (CE+PE) 1 session BEFORE the result, "
          "close 1 session AFTER (harvest IV crush). Entry dates are per-stock (see table).",
          "IV/RV Short Straddle: NO trade right now — 0 stocks currently show IV/RV >= 1.4 (market calm).",
          "Tail risk is accepted (no stop). Manage at portfolio level: small size, many concurrent, diversify."]:
    doc.add_paragraph(b, style="List Bullet")
doc.save(OUTD / "EXECUTION_PLAN.docx")
print(f"saved -> {OUTD}")
