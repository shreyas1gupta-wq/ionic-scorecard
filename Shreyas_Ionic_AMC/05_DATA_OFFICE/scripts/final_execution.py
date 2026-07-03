"""FINAL news-adjusted execution sheet + Word plan. Base conviction (edge+signal+earnings risk)
adjusted by a NEWS/sectoral overlay (IT pack detailed from research; other sectors flagged by
earnings-in-window + sector cluster). Produces execution_scored.csv, conviction_summary.csv,
and EXECUTION_PLAN.docx with conviction, sector, earnings flag, news note + a sectoral note.

EX-ANTE RISK GATING (RISK_LIMITS.md - APPROVED D-021): carries through the scanner's
entry_iv / iv_source / size_x / tail_tier / tail_warning columns and adds a hard EVENT GATE:
any Short_Strangle / IVRV trade with earnings inside its holding window (entry -> expiry) gets
blocked=True and conviction capped at 35 ("no naked short-vol through a name's known binary").
Previously this was only a conviction deduction. Earnings_ShortVol is exempt: trading the
binary IS that sleeve's design. New columns are APPENDED; existing names/columns unchanged.
Run execution_scanner.py --dry-run first if execution_ALL.csv lacks the risk columns.

P1 (2026Q3, IC-1 catch) NOTE: this script does NOT compute IV itself -- it only reads
entry_iv/iv_source/size_x through from execution_ALL.csv. The IV sanity cap (sane_iv(),
0.03 < iv < 1.0, iv_source='rejected' excluded from IVRV candidacy) lives entirely in
execution_scanner.py's apply_risk_overlay(). As long as execution_ALL.csv was produced
(or --dry-run re-passed) by the patched scanner, this script inherits already-clamped
values with no further guard needed here; it does not need its own copy of sane_iv().
"""
import datetime as dt, re
from pathlib import Path
import numpy as np, pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor

PROJ = Path(r"c:\Users\Shreyas.1Gupta\OneDrive - Angel Broking Limited\Desktop\Backup\NIFTY 500")
EXD = PROJ / "FINAL_STRATEGY_FORWARD_CHECK" / "08_Execution"
FWD = pd.read_csv(PROJ / "datasets/nse_earnings_dates/forthcoming_results.csv")
FWD["d"] = pd.to_datetime(FWD["date"], format="%d-%b-%Y", errors="coerce")
earn_dates = {}
for _, r in FWD.dropna(subset=["d"]).iterrows():
    earn_dates.setdefault(r["symbol"], []).append(r["d"].date())
FRONT_EXP = dt.date(2026, 7, 28)

import importlib.util
spec = importlib.util.spec_from_file_location("cs", str(Path(__file__).resolve().parent / "conviction_scorer.py"))
cs = importlib.util.module_from_spec(spec); spec.loader.exec_module(cs)   # reuse SECTOR + score_trade

# --- NEWS overlay (from 6-sector research sweep). (flag, note, conviction_adj) ---
NEWS = {
    # IT (results week 9-27 Jul + US H-1B $100k visa overhang)
    "TCS": ("HIGH", "Q1 9-Jul: IT season-opener; margin guidance sets sector tone", -11),
    "INFY": ("HIGH", "Q1 23-Jul: FY27 guidance reset = biggest single-stock IT gap catalyst", -12),
    "COFORGE": ("HIGH", "Q1 27-Jul: mid-cap, FIRST post-Encora consolidation quarter", -14),
    "HCLTECH": ("ELEVATED", "Q1 13-Jul + full-year guidance", -5),
    "TECHM": ("ELEVATED", "Q1 16-Jul; turnaround stock, higher realized vol", -6),
    # Banks / financials (private banks all 18-Jul)
    "HDFCBANK": ("HIGH", "Q1 18-Jul + CEO-reappoint/RBI-approval/governance-probe all cluster 15-18 Jul", -14),
    "ICICIBANK": ("ELEVATED", "Q1 18-Jul (clean bank, no binary overhang)", -5),
    "AXISBANK": ("ELEVATED", "Q1 18-Jul; history of asset-quality gap surprises", -7),
    "HDFCLIFE": ("ELEVATED", "Q1 15-Jul; VNB/APE print + HDFC-group sentiment", -5),
    "ANGELONE": ("HIGH", "Q1 15-16 Jul + LIVE SEBI weekly-expiry/F&O-tenure consultation-paper risk", -13),
    "BANKINDIA": ("NORMAL", "PSU bank; Q1 likely ~late-Jul (after 28-Jul expiry); healthy business update", 0),
    # Adani group (report 21-22 Jul, group-headline correlated)
    "ADANIGREEN": ("ELEVATED", "Q1 22-Jul; Adani-group headline-correlated", -6),
    "ADANIENSOL": ("ELEVATED", "Q1 21-Jul; Adani-group", -6),
    "ADANIPOWER": ("ELEVATED", "Q1 22-Jul; Adani-group", -6),
    # Pharma (US Sec-232 tariff overhang; generics exempt for now)
    "DRREDDY": ("HIGH", "Q1 22-Jul + fresh 7-obs USFDA 483 (biologics site) + GLP-1 headlines", -12),
    "CIPLA": ("ELEVATED", "Q1 23-Jul + unresolved Lanreotide supplier OAI", -6),
    "APOLLOHOSP": ("ELEVATED", "active demerger (HealthCo spin-off); Q1 likely mid-Aug -> verify date", -5),
    # FMCG / cement (all report late-Jul)
    "ASIANPAINT": ("ELEVATED", "Q1 ~28-29 Jul; Birla-Opus paint-war / market-share pressure", -6),
    "NESTLEIND": ("ELEVATED", "Q1 22-Jul; special-dividend decision 3-Jul", -5),
    "COLPAL": ("ELEVATED", "Q1 late-Jul; soft-volume + competition backdrop", -6),
    "ULTRACEMCO": ("ELEVATED", "Q1 20-Jul; weak July cement pricing + premium valuation", -6),
    # Auto / defence / metals
    "BAJAJ-AUTO": ("ELEVATED", "Q1 21-Jul; strong June sales already priced", -5),
    "BHARATFORG": ("ELEVATED", "lumpy defence order-flow headlines; Q1 mid-Jul (verify)", -6),
    "BDL": ("HIGH", "imminent Q1 + hot defence sector + India-Pak geopolitics + stretched valuation", -13),
    "JSWSTEEL": ("HIGH", "Q1 17-Jul + EU safeguard/CBAM step-change (1-Jul) + open duty decisions", -12),
    "COALINDIA": ("NORMAL", "PSU energy; no notable near-term binary found", 0),
}
IT_SET = {s for s, v in cs.SECTOR.items() if v == "IT"}
IT_NOTE = "IT sector: results week 9-27 Jul + H-1B $100k fee / visa overhang -> elevated sector IV"


def news_for(sym, strat, entry, expiry):
    if sym in NEWS:
        return NEWS[sym]
    # earnings-in-window (from announced calendar) -> ELEVATED
    for d in earn_dates.get(sym, []):
        if entry <= d <= expiry:
            return ("ELEVATED", f"Q1 earnings {d} in window", -5)
    if sym in IT_SET:
        return ("ELEVATED", "IT sector visa/earnings-week overhang", -4)
    return ("NORMAL", "no notable idiosyncratic news (earnings/sector only)", 0)


d = pd.read_csv(EXD / "execution_ALL.csv")
d["entry_date"] = pd.to_datetime(d["entry_date"]).dt.date

# --- ex-ante risk columns from execution_scanner (RISK_LIMITS D-021 inverse-IV sizing) ---
RISK_DEFAULTS = {"entry_iv": np.nan, "iv_source": "-", "size_x": 1.0, "tail_tier": "-", "tail_warning": ""}
_missing = [c for c in RISK_DEFAULTS if c not in d.columns]
for c in _missing:
    d[c] = RISK_DEFAULTS[c]
if _missing:
    print(f"WARN: execution_ALL.csv missing risk columns {_missing} -> defaults applied. "
          "Run execution_scanner.py --dry-run first for real inverse-IV sizing.")

# --- EVENT GATE (RISK_LIMITS: "No naked short-vol through a name's known binary event") ---
BLOCK_CONV_CAP = 35   # hard block: conviction forced to <=35 (was only a -15/-18 deduction)


def _gated_strat(strat):
    """Sleeves under the hard event gate. Earnings_ShortVol is exempt (binary IS its edge)."""
    s = str(strat).upper()
    return s == "SHORT_STRANGLE" or s.startswith("IVRV")


rows = []
cache = {}
blocked_log = []   # (strategy, symbol, earnings_date, conv_before, conv_after) for evidence
for _, r in d.iterrows():
    exp = FRONT_EXP if "JUL" in str(r["expiry"]) else dt.date(2026, 8, 25)
    key = (r["strategy"], r["symbol"])
    if key not in cache:
        base, flags = cs.score_trade(r["strategy"], r["symbol"], r["signal"], r["entry_date"], exp)
        nflag, nnote, nadj = news_for(r["symbol"], r["strategy"], r["entry_date"], exp)
        conv = int(np.clip(base + nadj, 0, 100))
        blocked = False
        if _gated_strat(r["strategy"]):
            ged = cs.has_earnings_before_expiry(r["symbol"], r["entry_date"], exp)
            if ged:   # binary inside holding window -> HARD BLOCK, not a deduction
                blocked = True
                conv_pre = conv
                conv = min(conv, BLOCK_CONV_CAP)
                flags = list(flags) + [f"EVENT-GATE BLOCK: earnings {ged} inside holding window "
                                       "-> DO NOT ENTER naked short-vol (RISK_LIMITS)"]
                blocked_log.append((r["strategy"], r["symbol"], ged, conv_pre, conv))
        cache[key] = (conv, flags, nflag, nnote, blocked)
    conv, flags, nflag, nnote, blocked = cache[key]
    rr = r.to_dict()
    rr["sector"] = cs.sector(r["symbol"]); rr["conviction"] = conv
    rr["news_risk"] = nflag; rr["news_note"] = nnote
    rr["risk_flags"] = "; ".join(flags) if flags else "-"
    rr["blocked"] = blocked
    rows.append(rr)

out = pd.DataFrame(rows).sort_values(["entry_date", "conviction", "strategy", "symbol", "opt"],
                                     ascending=[True, False, True, True, True])
cols = ["entry_date", "strategy", "action", "symbol", "sector", "expiry", "strike", "opt",
        "live_price", "lots", "lot_size", "conviction", "news_risk", "signal", "risk_flags", "news_note", "exit_rule",
        # ex-ante risk columns APPENDED (backward compatible) - RISK_LIMITS D-021
        "entry_iv", "iv_source", "size_x", "tail_tier", "tail_warning", "blocked"]
out[cols].to_csv(EXD / "execution_scored.csv", index=False)
trades = out.drop_duplicates(["strategy", "symbol", "entry_date"])
trades[["entry_date", "strategy", "symbol", "sector", "signal", "conviction", "news_risk", "risk_flags", "news_note",
        "entry_iv", "size_x", "tail_tier", "blocked"]].to_csv(
    EXD / "execution_conviction_summary.csv", index=False)

# --- Word doc ---
doc = Document()
doc.add_heading("EXECUTION PLAN + CONVICTION", 0)
p = doc.add_paragraph(f"Live prices {dt.date(2026,7,3)} (Fri close); next session Mon 6-Jul. Front expiry {FRONT_EXP}. "
                      "Conviction /100 = statistical edge + signal strength - earnings/event risk +/- news. "
                      "Account = disposable/data-only."); p.runs[0].italic = True

doc.add_heading("Conviction scoring method", 1)
for b in ["Base by strategy: strangle 70 (89% fwd-hit), IV/RV 72, earnings 62 (+39% but 60% hit), FF 60 (71% hit).",
          "Signal: FF magnitude (0.5-1.5 best; >1.5 usually earnings-driven -> discount); strangle credit%; ",
          "Risk deductions: earnings inside the front expiry (calendars/strangles = gap risk) -15 to -18.",
          "News overlay: HIGH RISK -12/-14, ELEVATED -4/-6, NORMAL 0 (IT pack researched; others by earnings/sector).",
          "RISK_LIMITS (D-021): SzX = inverse-IV size multiplier clip(0.25/entry_IV, 0.4, 1.5) on strangle/IVRV "
          "(final size = lots x SzX); top-quintile entry-IV names = HIGH tail tier, extra x0.6.",
          "EVENT GATE (hard): strangle/IVRV with earnings inside the holding window = BLOCKED, conviction capped "
          "at 35 - no naked short-vol through a binary. Blk column: X = do not enter."]:
    doc.add_paragraph(b, style="List Bullet")

def tbl(df, title):
    doc.add_heading(title, 1)
    t = doc.add_table(rows=1, cols=11); t.style = "Light Grid Accent 1"
    hs = ["Entry", "Action", "Symbol", "Sector", "Expiry", "Strike", "CE/PE", "Px", "Conv", "SzX", "Blk"]
    for i, h in enumerate(hs):
        rr = t.rows[0].cells[i].paragraphs[0].add_run(h); rr.bold = True; rr.font.size = Pt(8)
    for _, x in df.iterrows():
        c = t.add_row().cells
        for i, v in enumerate([str(x["entry_date"]), x["action"], x["symbol"], x["sector"], x["expiry"],
                               f"{x['strike']:g}", x["opt"], f"{x['live_price']}" if pd.notna(x['live_price']) else "-",
                               f"{x['conviction']}", f"{x['size_x']:.2f}", "X" if x["blocked"] else ""]):
            rr = c[i].paragraphs[0].add_run(v); rr.font.size = Pt(8)

for strat, title in [("FF_Calendar", "FF Calendars (enter Mon 6-Jul) - by conviction"),
                     ("Earnings_ShortVol", "Earnings short-vol (enter 1 session before each result)"),
                     ("Short_Strangle", "Short strangles (enter ~14-Jul) - top 25 by conviction")]:
    sub = out[out["strategy"] == strat].sort_values("conviction", ascending=False)
    if strat == "Short_Strangle":
        top = sub.drop_duplicates(["symbol"]).head(50)   # 25 trades = 50 legs
        sub = sub[sub["symbol"].isin(top["symbol"])]
    tbl(sub, title)

doc.add_heading("EVENT-GATE BLOCKED (RISK_LIMITS: no naked short-vol through a binary)", 1)
blk = trades[trades["blocked"]]
if len(blk):
    doc.add_paragraph(f"{len(blk)} strangle/IVRV trades BLOCKED (earnings inside holding window) - do NOT enter:")
    for _, x in blk.sort_values(["strategy", "symbol"]).iterrows():
        doc.add_paragraph(f"{x['strategy']} {x['symbol']} ({x['sector']}) - conviction capped at {x['conviction']}",
                          style="List Bullet")
else:
    doc.add_paragraph("No trades blocked by the event gate in this scan.")

doc.add_heading("Per-stock news & risk flags", 1)
tr = trades.sort_values("conviction")
t = doc.add_table(rows=1, cols=5); t.style = "Light Grid Accent 1"
for i, h in enumerate(["Symbol", "Sector", "Conv", "News", "Note / risk"]):
    rr = t.rows[0].cells[i].paragraphs[0].add_run(h); rr.bold = True; rr.font.size = Pt(8)
for _, x in tr.drop_duplicates("symbol").iterrows():
    note = x["news_note"] if x["news_note"] != "no notable idiosyncratic news (earnings/sector only)" else (x["risk_flags"] if x["risk_flags"] != "-" else "-")
    c = t.add_row().cells
    for i, v in enumerate([x["symbol"], x["sector"], f"{x['conviction']}", x["news_risk"], note[:90]]):
        rr = c[i].paragraphs[0].add_run(str(v)); rr.font.size = Pt(8)

doc.add_heading("Sectoral / macro note (3 Jul - 10 Aug 2026)", 1)
for b in [IT_NOTE + ".",
          "Bank results week ~17-18 Jul (HDFCBANK/ICICIBANK/AXISBANK) - private-bank IV elevated mid-July.",
          "Adani-group (ADANIGREEN/ENSOL/POWER) report 21-22 Jul - group-correlated, headline-prone.",
          "Pharma (CIPLA/DRREDDY) 22-23 Jul; FMCG (NESTLEIND/COLPAL/DABUR) 22-29 Jul; Auto (BAJAJ-AUTO/M&M) 21-30 Jul.",
          "EARNINGS-FLAG CAVEAT: only 27 stocks have ANNOUNCED board dates so far. By the 14-Jul strangle entry many more late-July earnings will be confirmed - RE-RUN the earnings refresh right before entry and skip/downsize strangles on any name reporting 14-28 Jul.",
          "Macro to watch for broad vol spikes: RBI MPC (early Aug), US Fed (late Jul), monthly F&O expiry 28-Jul."]:
    doc.add_paragraph(b, style="List Bullet")
doc.save(EXD / "EXECUTION_PLAN.docx")

print(f"scored {len(out)} legs / {len(trades)} trades -> execution_scored.csv + EXECUTION_PLAN.docx")

# --- ex-ante risk gate evidence (before/after) ---
sv_tr = trades[trades["strategy"].map(_gated_strat)]
print("\n=== EVENT GATE + INVERSE-IV SIZING (RISK_LIMITS D-021) - trade level ===")
print(f"short-vol trades in scope   : {len(sv_tr)}")
print(f"  downsized (size_x < 1.0)  : {(sv_tr['size_x'] < 1).sum()}   (before: all 1.00x)")
print(f"  upsized   (size_x > 1.0)  : {(sv_tr['size_x'] > 1).sum()}")
print(f"  tail_tier HIGH (x0.6)     : {(sv_tr['tail_tier'] == 'HIGH').sum()}")
print(f"  BLOCKED by event gate     : {len(blocked_log)}   (before: deduction only, no block)")
if blocked_log:
    print("  blocked trades (conviction before -> after cap 35):")
    for st, sym, ged, cb, ca in sorted(blocked_log):
        print(f"    {st:16s} {sym:12s} earnings {ged}  conv {cb} -> {ca}")

print("\n=== TOP 10 by conviction ===")
print(trades.sort_values("conviction", ascending=False).head(10)[
    ["entry_date", "strategy", "symbol", "sector", "conviction", "news_risk"]].to_string(index=False))
print("\n=== LOWEST 10 (avoid / downsize) ===")
print(trades.sort_values("conviction").head(10)[
    ["entry_date", "strategy", "symbol", "sector", "conviction", "news_risk", "news_note"]].to_string(index=False))
