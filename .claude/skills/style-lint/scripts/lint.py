#!/usr/bin/env python
"""
style-lint / lint.py -- mechanical AI-tell + house-rule checker.
Owner: Tanvi Desai (Product) / Lakshmi Narayanan (Librarian), per WS-2 (SYSTEM_SCIENCE_PROGRAM).
Model tier: haiku-class (deterministic script; no judgment calls required to RUN it).

Checks a draft (.md/.txt/.docx) against:
  1. avoid-ai-writing's verified 53-category taxonomy (data/taxonomy.json), WebFetch-sourced
     2026-07-12 from https://github.com/conorbronsdon/avoid-ai-writing -- see STYLE_GUIDE.md.
  2. House structural rules (em-dash cap, rule-of-three, negation pivot, bare-noun bullets).
  3. House positive-rule advisories (claim-bearing lines without a file-path/date citation,
     uniform sentence-length cadence).

Usage:
    python lint.py <path-to-draft> [--out report.md]

No network calls. No third-party deps beyond python-docx (only needed for .docx input;
.md/.txt need nothing).
"""
import sys
import os
import re
import json
import statistics
from collections import defaultdict

HERE = os.path.dirname(os.path.abspath(__file__))
TAXONOMY_PATH = os.path.join(HERE, "..", "data", "taxonomy.json")


def load_taxonomy():
    with open(TAXONOMY_PATH, "r", encoding="utf-8") as f:
        return json.load(f)


def read_lines(path):
    """Return list of (line_no, text) pseudo-lines. .docx -> one entry per paragraph
    (plus table cells appended at the end, numbered as pseudo-lines continuing the count)."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        try:
            import docx
        except ImportError:
            print("ERROR: python-docx not installed; cannot read .docx. "
                  "pip install python-docx", file=sys.stderr)
            sys.exit(2)
        d = docx.Document(path)
        lines = []
        n = 0
        for para in d.paragraphs:
            n += 1
            if para.text.strip():
                lines.append((n, para.text))
        for ti, table in enumerate(d.tables):
            for ri, row in enumerate(table.rows):
                for cell in row.cells:
                    n += 1
                    if cell.text.strip():
                        lines.append((n, f"[table {ti} row {ri}] {cell.text}"))
        return lines
    else:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            raw = f.readlines()
        return [(i + 1, line.rstrip("\n")) for i, line in enumerate(raw)]


def word_count(text):
    return len(re.findall(r"[A-Za-z']+", text))


def find_all(pattern, text, flags=re.IGNORECASE):
    return list(re.finditer(pattern, text, flags))


def word_regex(phrase):
    """Build a case-insensitive, word-boundary-safe regex for a literal phrase/word."""
    escaped = re.escape(phrase)
    # allow the phrase to be a multi-word literal; boundary at start/end only
    return r"(?<![A-Za-z])" + escaped + r"(?![A-Za-z])"


def lint(path, tax):
    lines = read_lines(path)
    full_text = "\n".join(t for _, t in lines)
    total_words = max(word_count(full_text), 1)

    findings = []  # each: dict(line, category, severity, matched, suggestion)

    # --- Tier 1: always replace ---
    for phrase, repl in tax["tier1_always_replace"].items():
        rx = word_regex(phrase)
        for lineno, text in lines:
            for m in find_all(rx, text):
                findings.append(dict(line=lineno, category="tier1-word-replace",
                                      severity="P1", matched=m.group(0),
                                      suggestion=f"-> {repl}"))

    # --- Tier 1 copula avoidance (subset of tier1, called out separately) ---
    for phrase in tax["tier1_copula_avoidance"]:
        rx = word_regex(phrase)
        for lineno, text in lines:
            for m in find_all(rx, text):
                findings.append(dict(line=lineno, category="copula-avoidance",
                                      severity="P1", matched=m.group(0),
                                      suggestion="-> is / has (plain copula)"))

    # --- Tier 2: flag when 2+ occurrences in the same line/paragraph ---
    for lineno, text in lines:
        hits_here = []
        for phrase, repl in tax["tier2_flag_2plus_per_paragraph"].items():
            rx = word_regex(phrase)
            for m in find_all(rx, text):
                hits_here.append((m.group(0), repl))
        if len(hits_here) >= 2:
            for matched, repl in hits_here:
                findings.append(dict(line=lineno, category="tier2-cluster",
                                      severity="P1", matched=matched,
                                      suggestion=f"-> {repl} (2+ tier-2 words in this paragraph)"))

    # --- Tier 3: document-level density ---
    tier3_hits = []
    for word in tax["tier3_high_density_words"]:
        rx = word_regex(word)
        for lineno, text in lines:
            for m in find_all(rx, text):
                tier3_hits.append((lineno, m.group(0)))
    density_pct = 100.0 * len(tier3_hits) / total_words
    if density_pct >= tax["tier3_density_threshold_pct"]:
        sample_lines = sorted(set(l for l, _ in tier3_hits))[:10]
        findings.append(dict(line=",".join(map(str, sample_lines)) or "-",
                              category="tier3-density", severity="P1",
                              matched=f"{len(tier3_hits)} tier-3 words / {total_words} words "
                                      f"({density_pct:.1f}%)",
                              suggestion="cut or replace tier-3 intensity words; density >= "
                                         f"{tax['tier3_density_threshold_pct']}% threshold"))

    # --- Phrase families ---
    for family, spec in tax["phrase_families"].items():
        hits = []
        for item in spec["items"]:
            rx = word_regex(item) if " " not in item.strip() else re.escape(item)
            for lineno, text in lines:
                for m in re.finditer(rx, text, re.IGNORECASE):
                    hits.append((lineno, m.group(0)))
        if len(hits) >= spec.get("flag_at", 1):
            for lineno, matched in hits:
                findings.append(dict(line=lineno, category=f"phrase:{family}",
                                      severity=spec["severity"], matched=matched,
                                      suggestion="cut / rewrite plainly"))

    # --- Structural: em dash ---
    em_dash_count = full_text.count("—") + full_text.count(" -- ")
    rate_per_1000 = 1000.0 * em_dash_count / total_words
    cap = tax["structural_checks"]["em_dash_hard_cap_per_1000_words"]
    if em_dash_count > 0:
        sev = "P1" if rate_per_1000 > cap else "P2"
        findings.append(dict(line="doc", category="structural:em-dash", severity=sev,
                              matched=f"{em_dash_count} em-dashes, {rate_per_1000:.2f}/1000 words",
                              suggestion=f"target zero; hard cap {cap}/1000 words -- replace with "
                                         "commas/periods/parens or split the sentence"))

    # --- Structural: negation pivot ---
    npat = tax["structural_checks"]["negation_pivot_regex"]
    for lineno, text in lines:
        if re.search(npat, text, re.IGNORECASE):
            findings.append(dict(line=lineno, category="structural:negation-pivot",
                                  severity="P1", matched=text.strip()[:80],
                                  suggestion="rewrite 'it's not X -- it's Y' as a direct positive "
                                             "statement"))

    # --- Structural: rule of three (heuristic) ---
    rx3 = tax["structural_checks"]["rule_of_three_regex"]
    r3_hits = 0
    for lineno, text in lines:
        r3_hits += len(re.findall(rx3, text))
    if r3_hits >= 4:
        findings.append(dict(line="doc", category="structural:rule-of-three", severity="P2",
                              matched=f"{r3_hits} comma/'and' triplet groupings",
                              suggestion="compulsive rule-of-three -- vary list lengths, "
                                         "some pairs, some singles, some fours"))

    # --- Structural: bare-noun bullet runs ---
    min_items = tax["structural_checks"]["bullet_bare_noun_min_items"]
    run = []
    bullet_rx = re.compile(r"^\s*[-*]\s+(.*)$")
    for lineno, text in lines:
        m = bullet_rx.match(text)
        if m:
            item_text = m.group(1).strip()
            has_digit = bool(re.search(r"\d", item_text))
            has_verb_like = bool(re.search(r"\b(is|are|was|were|has|have|shows|means|drives|"
                                            r"kills|beats|misses)\b", item_text, re.IGNORECASE))
            word_ct = len(item_text.split())
            if not has_digit and not has_verb_like and word_ct <= 6:
                run.append(lineno)
            else:
                if len(run) >= min_items:
                    findings.append(dict(line=f"{run[0]}-{run[-1]}",
                                          category="structural:bullet-itis", severity="P1",
                                          matched=f"{len(run)} consecutive bare-noun bullets",
                                          suggestion="bullet-point-itis -- give each item a verb, "
                                                     "a number, or a source, or convert to prose"))
                run = []
        else:
            if len(run) >= min_items:
                findings.append(dict(line=f"{run[0]}-{run[-1]}",
                                      category="structural:bullet-itis", severity="P1",
                                      matched=f"{len(run)} consecutive bare-noun bullets",
                                      suggestion="bullet-point-itis -- give each item a verb, "
                                                 "a number, or a source, or convert to prose"))
            run = []
    if len(run) >= min_items:
        findings.append(dict(line=f"{run[0]}-{run[-1]}", category="structural:bullet-itis",
                              severity="P1", matched=f"{len(run)} consecutive bare-noun bullets",
                              suggestion="bullet-point-itis -- give each item a verb, a number, "
                                         "or a source, or convert to prose"))

    # --- House: claim-bearing lines without citation ---
    # "nearby" = same line, OR the immediately preceding/following scanned paragraph
    # (covers the common docx pattern: claim paragraph, then a separate "Source: ..." caption).
    cite_rx = tax["house_positive_checks"]["citation_regex"]
    for idx, (lineno, text) in enumerate(lines):
        if re.search(r"\d", text) and len(text.split()) >= 6:
            window = [text]
            if idx > 0:
                window.append(lines[idx - 1][1])
            if idx < len(lines) - 1:
                window.append(lines[idx + 1][1])
            if not any(re.search(cite_rx, w, re.IGNORECASE) for w in window):
                findings.append(dict(line=lineno, category="house:missing-citation",
                                      severity="house-P2", matched=text.strip()[:80],
                                      suggestion="[house] claim-bearing line has no file-path / "
                                                 "'as of YYYY-MM-DD' / 'source:' citation in this "
                                                 "or an adjacent line"))

    # --- House: cadence (sentence-length uniformity) ---
    sentences = [s.strip() for s in re.split(r"(?<=[.!?])\s+", full_text) if s.strip()]
    lens = [len(s.split()) for s in sentences if len(s.split()) > 0]
    if len(lens) >= 6:
        stdev = statistics.pstdev(lens)
        mean = statistics.mean(lens)
        if stdev < 3.0 and mean < 20:
            findings.append(dict(line="doc", category="house:cadence", severity="house-P2",
                                  matched=f"mean {mean:.1f} words/sentence, stdev {stdev:.1f}",
                                  suggestion="[house] uniform sentence rhythm -- alternate short "
                                             "(5-12w) and long (25w+) sentences"))

    return findings, total_words, len(lines)


SEV_WEIGHT = {"P0": 3, "P1": 2, "P2": 1, "house-P2": 1}


def render_report(path, findings, total_words, n_lines):
    out = []
    out.append(f"# style-lint report -- `{path}`\n")
    out.append(f"Words: {total_words} | Lines/paragraphs scanned: {n_lines} | "
               f"Findings: {len(findings)}\n")

    by_sev = defaultdict(int)
    for f in findings:
        by_sev[f["severity"]] += 1
    out.append("## Severity summary\n")
    for sev in ["P0", "P1", "P2", "house-P2"]:
        if by_sev[sev]:
            out.append(f"- {sev}: {by_sev[sev]}")
    score = sum(SEV_WEIGHT.get(f["severity"], 1) for f in findings)
    score_per_1000 = 1000.0 * score / max(total_words, 1)
    out.append(f"\n**[house] AI-tell weighted score: {score} "
               f"({score_per_1000:.2f} per 1,000 words)** -- our own metric "
               f"(P0=3, P1=2, P2/house-P2=1); not a port of the source repo's JS scorer.\n")

    if not findings:
        out.append("\nNo violations found against the loaded taxonomy. Still check by eye for "
                    "anything the mechanical pass can't catch (voice, honesty of the numbers, "
                    "whether kills/artifacts sit next to the headline).\n")
        return "\n".join(out)

    out.append("\n## Findings (line, category, severity, matched, suggestion)\n")
    out.append("| Line | Category | Severity | Matched | Suggestion |")
    out.append("|---|---|---|---|---|")
    for f in sorted(findings, key=lambda x: (str(x["line"]), x["category"])):
        matched = f["matched"].replace("|", "/").replace("\n", " ")[:100]
        suggestion = f["suggestion"].replace("|", "/")
        out.append(f"| {f['line']} | {f['category']} | {f['severity']} | {matched} | {suggestion} |")
    return "\n".join(out)


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)
    path = sys.argv[1]
    out_path = None
    if "--out" in sys.argv:
        out_path = sys.argv[sys.argv.index("--out") + 1]
    if not os.path.exists(path):
        print(f"ERROR: file not found: {path}", file=sys.stderr)
        sys.exit(2)
    tax = load_taxonomy()
    findings, total_words, n_lines = lint(path, tax)
    report = render_report(path, findings, total_words, n_lines)
    if out_path:
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(report)
        print(f"Report written: {out_path} ({len(findings)} findings)")
    else:
        print(report)


if __name__ == "__main__":
    main()
