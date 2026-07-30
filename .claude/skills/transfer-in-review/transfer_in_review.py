# -*- coding: utf-8 -*-
"""transfer_in_review.py — NDPMS transfer-in portfolio review.

Self-contained: no dependency on any specific firm's repo layout. Everything
this script needs (the NIFTY 500 reference list) ships alongside it in
reference/nifty500_current.json. Copy this whole `transfer-in-review` folder
wherever you like and it still works.

Checks an incoming client's existing holdings statement (CAS / Demat / CAMS /
Kfintech) against a transfer-in acceptance checklist and produces a Transfer
Report: an immediate pre-approved list (Stocks / MF & SIF / Bonds & InvIT)
plus an exceptions list with named routing.

CHECKLIST
  Stocks       -> beyond the NIFTY 500?                 exception -> "Yajash"
  Debt & InvIT -> listed AND credit rating > A+?         exception -> "Girdhar"
  MF & SIF     -> demat-mode (ISIN) holding, Direct plan? exception -> "RM"
  Every stock and every demat-mode MF/SIF row (pass or fail on the checks
  above) also gets an RM_purchase_history_confirmed column, left blank —
  this cannot be automated and must never be marked done by this script.

There is deliberately NO family-AUM / minimum-ticket-size threshold check in
this workflow — if your process needs one, add it explicitly; do not assume
one belongs here.

INPUT
  Preferred: a pre-extracted CSV/XLSX with columns (case-insensitive, extra
  columns ignored):
    type/category : EQ | MF | SIF | DEBT | INVIT   (optional — inferred if absent)
    isin           : strongly preferred, exact matching
    name           : security / scheme name as printed on the statement
    units          : quantity / units held
    value_inr      : current value in INR
    rating         : credit-rating string if visible on the statement (debt/InvIT)
    plan           : Direct | Regular, if visible (MF/SIF)
    family_member  : whose account this line came from (carried through, not scored)

  Also accepted: a raw CAS/CAMS/Kfintech PDF (requires `pip install pymupdf`).
  Extraction is HEURISTIC — line/ISIN scanning, not real table-structure
  parsing, because statement layouts vary by depository/RTA and this has not
  been validated against a real sample of every format. Every PDF-extracted
  row is tagged extraction_confidence=low and MUST be checked against the
  source statement before the report is treated as final. Prefer the CSV/XLSX
  path whenever you can get one.

USAGE
  python transfer_in_review.py --statement <file.pdf|csv|xlsx> --client "Name" --out <dir>

REQUIREMENTS
  pip install python-docx openpyxl          (always)
  pip install pymupdf                       (only if you'll feed it raw PDFs)
"""
from __future__ import annotations

import os
import re
import csv
import sys
import json
import argparse
import datetime as dt

HERE = os.path.dirname(os.path.abspath(__file__))
NIFTY500_JSON = os.path.join(HERE, "reference", "nifty500_current.json")

# ---------------------------------------------------------------------------
# Reference data / constants
# ---------------------------------------------------------------------------
# India-only: real ISINs here always start "IN" (NDPMS is an India workflow).
# Restricting the prefix (vs any 2-letter country) avoids false ISIN-shaped matches
# on stray codes/footer text when scanning raw PDF pages.
ISIN_RE = re.compile(r"\bIN[A-Z0-9]{9}[0-9]\b")

# long-term ratings strictly better than "A+" (AA- and above), plus the
# short-term top-tier equivalent (A1+) commonly seen on InvIT/NCD/CP paper.
# Anything NOT in this set — including A+, A, A-, BBB.., unrated, or an
# unrecognized string — does not auto-pass; conservative by design.
RATING_PASS_SET = {"AAA", "AA+", "AA", "AA-", "A1+"}

# ordered longest-token-first so alternation doesn't short-match e.g. "AA" inside "AA+".
# NOTE: plain \b word-boundaries do NOT work as the trailing delimiter here —
# a token ending in '+'/'-' followed by a space or end-of-string has \W on
# BOTH sides of that position, so \b silently fails to match there and the
# alternation falls through to the next (shorter) alternative — e.g. "A1+"
# would silently downgrade to "A1", which is outside RATING_PASS_SET even
# though A1+ is meant to pass. Use explicit [A-Z0-9] lookaround instead.
_RATING_TOKENS = ["AAA", "AA+", "AA-", "AA", "A1+", "A1", "BBB+", "BBB-", "BBB",
                  "BB+", "BB-", "BB", "A+", "A-", "A", "B+", "B-", "B", "C", "D"]
RATING_TOKEN_RE = re.compile(
    r"(?<![A-Z0-9])(" + "|".join(re.escape(t) for t in _RATING_TOKENS) + r")(?![A-Z0-9])"
)

SECTION_KEYWORDS = [
    (re.compile(r"MUTUAL\s+FUND", re.I), "MF"),
    (re.compile(r"\bSIF\b|SPECIALIZED\s+INVESTMENT\s+FUND", re.I), "SIF"),
    (re.compile(r"GOVERNMENT\s+SECURIT|BOND|DEBENTURE|\bNCD\b|\bSDL\b|G-?SEC", re.I), "DEBT"),
    (re.compile(r"\bINVIT\b|\bREIT\b", re.I), "INVIT"),
    (re.compile(r"EQUIT|\bSHARES?\b", re.I), "EQ"),
]
NAME_KEYWORDS = SECTION_KEYWORDS  # same patterns, applied to the row's own name text as fallback

NUM_RE = re.compile(r"-?\d[\d,]*\.?\d*")

CATEGORY_ALIASES = {"EQ": "EQ", "EQUITY": "EQ", "STOCK": "EQ", "SHARE": "EQ",
                    "MF": "MF", "MUTUALFUND": "MF",
                    "SIF": "SIF",
                    "DEBT": "DEBT", "BOND": "DEBT", "NCD": "DEBT", "GSEC": "DEBT",
                    "INVIT": "INVIT", "REIT": "INVIT"}

REQUIRES_PURCHASE_HISTORY = {"EQ", "MF", "SIF"}  # "stocks and MF/SIF in demat" per the checklist

# family-AUM minimum-ticket gate: floor, both conditions required (AND).
# Below either threshold -> the transfer does not proceed under this workflow.
THRESHOLD_PER_MEMBER_INR = 1_00_00_000    # > Rs 1 Cr per family member
THRESHOLD_FAMILY_INR = 2_50_00_000        # > Rs 2.5 Cr at family level


def _norm(s):
    return re.sub(r"[^a-z0-9]", "", str(s or "").lower())


def _to_str(v):
    """Safe stringify for any cell type openpyxl may hand back (int/float/
    datetime/None/str) — CSV rows are always already str, but XLSX cells keep
    their native type, and a bare number in e.g. the rating/family_member
    column must not crash the pipeline."""
    return "" if v is None else str(v).strip()


def _to_float(s):
    try:
        return float(str(s).replace(",", "").strip())
    except Exception:
        return None


NEGATIVE_TOKENS = {"NO", "N", "FALSE", "0", "UNLISTED", "NOT LISTED"}


def _parse_listed(raw):
    """True/False/None (not stated). A bare non-empty string is NOT automatically
    "listed" — must check for an explicit negative first (bool("No") is True in
    Python, which would otherwise treat an explicitly-unlisted instrument as
    listed)."""
    s = _to_str(raw).upper()
    if not s:
        return None
    if s in NEGATIVE_TOKENS:
        return False
    return True


def _extract_rating_token(raw):
    """Pull the bare rating notch (e.g. 'AA+') out of a free-text rating field
    that may carry an agency prefix ('CRISIL AA+') or outlook suffix
    ('AA+ (Stable)'). Falls back to the raw uppercased string if no recognized
    token is found, so an unrecognized rating still shows up verbatim in the
    exception reason rather than silently as ''."""
    s = _to_str(raw).upper()
    if not s:
        return ""
    m = RATING_TOKEN_RE.search(s)
    return m.group() if m else s


# ---------------------------------------------------------------------------
# Reference: current NIFTY 500 membership (bundled alongside this script)
# ---------------------------------------------------------------------------
def load_nifty500(path=NIFTY500_JSON):
    with open(path, encoding="utf-8") as f:
        payload = json.load(f)
    rows = payload["constituents"] if isinstance(payload, dict) else payload
    by_isin, by_name = {}, {}
    for r in rows:
        isin = str(r.get("isin") or "").strip().upper()
        if isin:
            by_isin[isin] = r
        by_name[_norm(r.get("company") or r.get("symbol"))] = r
    as_of = payload.get("as_of") if isinstance(payload, dict) else None
    return by_isin, by_name, as_of


def _match_nifty500(isin, name, by_isin, by_name):
    isin = str(isin or "").strip().upper()
    if isin and ISIN_RE.fullmatch(isin) and isin in by_isin:
        return by_isin[isin]
    key = _norm(name)
    if key and key in by_name:
        return by_name[key]
    # longest-prefix fallback (>=10 normalized chars) — statement spellings vary
    # ("Ltd" vs "Limited", "&" vs "and") and an exact-name match alone would
    # otherwise mislabel real NIFTY 500 constituents as "beyond NIFTY 500".
    best, best_n = None, 0
    for k, v in by_name.items():
        n = 0
        while n < min(len(key), len(k)) and key[n] == k[n]:
            n += 1
        if n > best_n:
            best, best_n = v, n
    return best if best_n >= 10 else None


# ---------------------------------------------------------------------------
# Input: pre-extracted CSV / XLSX (preferred path)
# ---------------------------------------------------------------------------
def load_pre_extracted(path):
    rows = []
    if path.lower().endswith((".xlsx", ".xls")):
        import openpyxl
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        ws = wb[wb.sheetnames[0]]
        head = None
        for r in ws.iter_rows(values_only=True):
            if head is None:
                head = [_norm(c) for c in r]
                continue
            if all(c is None for c in r):
                continue
            rows.append({h: v for h, v in zip(head, r)})
    else:
        with open(path, encoding="utf-8-sig", newline="") as f:
            for r in csv.DictReader(f):
                rows.append({_norm(k): v for k, v in r.items()})
    out = []
    for r in rows:
        raw_val = r.get("valueinr")
        if raw_val is None:
            raw_val = r.get("value")
        if raw_val is None:
            raw_val = r.get("marketvalue")
        value_inr = _to_float(raw_val)
        # a non-blank value that still failed to parse (stray currency symbol,
        # "N/A", etc.) must not silently become 0 — that would quietly corrupt
        # both the pre-approved value shown and the family-AUM total.
        value_parse_error = (_to_str(raw_val) != "" and value_inr is None)
        out.append({
            "category_hint": _to_str(r.get("type") or r.get("category")).upper(),
            "isin": _to_str(r.get("isin")).upper(),
            "name": _to_str(r.get("name")),
            "units": _to_float(r.get("units") or r.get("qty") or r.get("quantity")),
            "value_inr": value_inr,
            "value_parse_error": value_parse_error,
            "value_raw": _to_str(raw_val),
            "rating_raw": _to_str(r.get("rating")),
            "plan_raw": _to_str(r.get("plan")),
            "listed_raw": _to_str(r.get("listed")),
            "family_member": _to_str(r.get("familymember") or r.get("family")),
            "extraction_confidence": "high",  # human/ops-typed input
            "source_line": "",
        })
    return out


# ---------------------------------------------------------------------------
# Input: raw CAS/CAMS/Kfintech PDF (heuristic — see module docstring)
# ---------------------------------------------------------------------------
def extract_from_pdf(path):
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise SystemExit(
            "Raw PDF input needs PyMuPDF: run `pip install pymupdf`, or supply a "
            "pre-extracted CSV/XLSX instead (see this script's docstring)."
        )
    doc = fitz.open(path)
    out = []
    current_section = None
    for page_no, page in enumerate(doc, 1):
        text = page.get_text("text")
        for line in text.splitlines():
            line = line.strip()
            if not line:
                continue
            for pat, sect in SECTION_KEYWORDS:
                if pat.search(line):
                    current_section = sect
                    break
            m = ISIN_RE.search(line)
            if not m:
                continue
            isin = m.group()
            name_guess = line[:m.start()].strip(" -:|") or line
            nums = [_to_float(n) for n in NUM_RE.findall(line)]
            nums = [n for n in nums if n is not None]
            units_guess = nums[-2] if len(nums) >= 2 else None
            value_guess = nums[-1] if nums else None
            rating_m = RATING_TOKEN_RE.search(line)
            out.append({
                "category_hint": current_section or "",
                "isin": isin,
                "name": name_guess,
                "units": units_guess,
                "value_inr": value_guess,
                "value_parse_error": False,
                "value_raw": "" if value_guess is None else str(value_guess),
                "rating_raw": rating_m.group() if rating_m else "",
                "plan_raw": "DIRECT" if re.search(r"\bDIRECT\b", line, re.I)
                            else ("REGULAR" if re.search(r"\bREGULAR\b", line, re.I) else ""),
                "listed_raw": "LISTED" if re.search(r"\bLISTED\b|\bNSE\b|\bBSE\b", line, re.I) else "",
                "family_member": "",
                "extraction_confidence": "low",
                "source_line": f"p{page_no}: {line}",
            })
    doc.close()
    return out


# ---------------------------------------------------------------------------
# Category inference + checks
# ---------------------------------------------------------------------------
def guess_category(row):
    hint = CATEGORY_ALIASES.get(str(row["category_hint"]).strip().upper(), "")
    if hint:
        return hint
    # AMFI convention: mutual-fund-scheme ISINs start "INF" — check this before
    # falling back to name keywords, since a scheme literally named e.g. "HDFC
    # Equity Fund Direct Growth" would otherwise trip the EQ ("EQUIT") keyword
    # and get miscategorized as a stock.
    if (row.get("isin") or "").startswith("INF"):
        return "MF"
    for pat, sect in NAME_KEYWORDS:
        if pat.search(row["name"] or ""):
            return sect
    return "UNKNOWN"


def check_row(row, by_isin, by_name):
    """Returns (status, routed_to, reason) for one holding row."""
    # a value that failed to parse corrupts both this row's own numbers and the
    # family-AUM total — flag it ahead of the per-category checks, regardless
    # of category, rather than silently treating it as a zero-value holding.
    if row.get("value_parse_error"):
        return ("Exception", "RM",
                f"value_inr could not be parsed from '{row.get('value_raw')}' — fix and re-run")

    cat = row["category"]
    if cat == "EQ":
        hit = _match_nifty500(row["isin"], row["name"], by_isin, by_name)
        if hit:
            return "Pre-Approved", None, "in current NIFTY 500"
        return "Exception", "Yajash", "beyond NIFTY 500 — review and confirm"

    if cat in ("DEBT", "INVIT"):
        rating = _extract_rating_token(row["rating_raw"])
        listed = _parse_listed(row["listed_raw"])
        if rating in RATING_PASS_SET and listed is True:
            return "Pre-Approved", None, f"rating {rating}, listed"
        reasons = []
        if rating not in RATING_PASS_SET:
            reasons.append(f"rating '{rating or 'not found on statement'}' not confirmed > A+")
        if listed is False:
            reasons.append("explicitly marked NOT listed on the statement")
        elif listed is None:
            reasons.append("listed status not confirmed from statement")
        return "Exception", "Girdhar", "; ".join(reasons)

    if cat in ("MF", "SIF"):
        plan = row["plan_raw"].strip().upper()
        if plan == "DIRECT":
            return "Pre-Approved", None, "Direct plan, demat/ISIN holding"
        return "Exception", "RM", (f"plan is '{plan or 'not stated'}', not confirmed Direct — "
                                    f"no ISIN->plan master available; RM to check")

    return "Exception", "RM", "category could not be determined from statement — RM to classify"


# ---------------------------------------------------------------------------
# Report generation
# ---------------------------------------------------------------------------
def build_report_rows(raw_rows, by_isin, by_name):
    out = []
    isin_counts = {}
    for r in raw_rows:
        isin = r.get("isin") or ""
        if isin:
            isin_counts[isin] = isin_counts.get(isin, 0) + 1

    for r in raw_rows:
        cat = guess_category(r)
        r["category"] = cat
        status, routed_to, reason = check_row(r, by_isin, by_name)
        isin = r.get("isin") or ""
        if isin and isin_counts.get(isin, 0) > 1:
            # e.g. a CAS summary section + detail section both listing the same
            # ISIN — don't silently sum or drop either occurrence (that risk cuts
            # both ways), just surface it so a human resolves whether it's a
            # genuine double count before it corrupts the family-AUM total.
            reason = f"{reason} [DUPLICATE ISIN: appears {isin_counts[isin]}x in this input — verify not double-counted]"
        if cat in REQUIRES_PURCHASE_HISTORY:
            ph = ""
        elif cat == "UNKNOWN":
            ph = "category unresolved — classify before deciding"
        else:
            ph = "n/a"
        out.append({
            **r,
            "status": status,
            "routed_to": routed_to or "",
            "reason": reason,
            "rm_purchase_history_confirmed": ph,
        })
    return out


def compute_family_gate(rows):
    """Minimum-ticket gate: EACH family member's total value > Rs 1 Cr AND the
    combined family total > Rs 2.5 Cr (floor, AND logic). Returns a dict with
    per-member totals/pass, family total/pass, and an overall verdict — or
    "UNKNOWN" if no family_member tags were present on the input, since the
    per-member half of the gate can't be computed without them.

    Scope note [ASSUMPTION]: the family total here sums EVERY row in the
    statement (stocks + MF/SIF + debt/InvIT). If "family level" is meant to
    apply only to the sleeve actually moving into the PMS mandate rather than
    the whole existing portfolio, this total will read too high — confirm the
    intended scope."""
    member_totals = {}
    family_total = 0.0
    isin_counts = {}
    unparsed_count = 0
    for r in rows:
        if r.get("value_parse_error"):
            unparsed_count += 1
        v = r.get("value_inr") or 0
        family_total += v
        fm = _to_str(r.get("family_member"))
        if fm:
            member_totals[fm] = member_totals.get(fm, 0.0) + v
        isin = r.get("isin") or ""
        if isin:
            isin_counts[isin] = isin_counts.get(isin, 0) + 1

    notes = []
    if unparsed_count:
        notes.append(f"{unparsed_count} row(s) had a value that failed to parse and were "
                     f"counted as Rs 0 here — the true family total may be higher; see Exceptions.")
    dup_isins = [i for i, n in isin_counts.items() if n > 1]
    if dup_isins:
        notes.append(f"{len(dup_isins)} ISIN(s) appear more than once in this input — if that's a "
                     f"double count (e.g. summary + detail sections of one CAS), the family total "
                     f"below is inflated. See the DUPLICATE ISIN rows in the exceptions.")

    family_pass = family_total > THRESHOLD_FAMILY_INR
    if not member_totals:
        return {
            "member_totals": {}, "family_total": family_total, "family_pass": family_pass,
            "verdict": "UNKNOWN", "note": "no family_member column on the input — "
            "cannot evaluate the per-member >1 Cr half of the gate",
            "data_quality_notes": notes,
        }
    per_member_pass = {m: v > THRESHOLD_PER_MEMBER_INR for m, v in member_totals.items()}
    can_process = family_pass and all(per_member_pass.values())
    return {
        "member_totals": member_totals, "per_member_pass": per_member_pass,
        "family_total": family_total, "family_pass": family_pass,
        "verdict": "CAN PROCESS" if can_process else "CANNOT PROCESS (below minimum)",
        "data_quality_notes": notes,
    }


CSV_COLS = ["category", "isin", "name", "units", "value_inr", "value_raw", "rating_raw", "plan_raw",
            "listed_raw", "family_member", "status", "routed_to", "reason",
            "rm_purchase_history_confirmed", "extraction_confidence", "source_line"]


def load_prior_confirmations(csv_path):
    """If out_dir already has a transfer_report.csv from an earlier run (e.g. RM
    ticked rm_purchase_history_confirmed by hand, then the statement got re-run
    for an unrelated correction), carry those confirmations forward by (isin,
    name) — a rerun must never silently wipe a human sign-off."""
    prior = {}
    if not os.path.exists(csv_path):
        return prior
    try:
        with open(csv_path, encoding="utf-8-sig", newline="") as f:
            for r in csv.DictReader(f):
                val = (r.get("rm_purchase_history_confirmed") or "").strip()
                if val and val not in ("n/a", "category unresolved — classify before deciding"):
                    prior[(r.get("isin", "").strip().upper(), _norm(r.get("name")))] = val
    except Exception:
        return {}
    return prior


def apply_prior_confirmations(rows, prior):
    if not prior:
        return
    for r in rows:
        key = (r.get("isin", "").strip().upper(), _norm(r.get("name")))
        if key in prior and r.get("rm_purchase_history_confirmed") == "":
            r["rm_purchase_history_confirmed"] = prior[key]


def write_csv(rows, out_path):
    with open(out_path, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.DictWriter(f, fieldnames=CSV_COLS)
        w.writeheader()
        for r in rows:
            w.writerow({k: r.get(k, "") for k in CSV_COLS})


def _docx_table(doc, headers, rows, right_align_cols=None):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    right_align_cols = right_align_cols or set()
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = str(h)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            cell = t.rows[i].cells[j]
            cell.text = str(v)
            if j in right_align_cols:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()
    return t


def write_docx(rows, out_path, client_name, source_path, as_of, nifty_as_of, family_gate):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    approved = [r for r in rows if r["status"] == "Pre-Approved"]
    exceptions = [r for r in rows if r["status"] != "Pre-Approved"]
    low_conf = any(r["extraction_confidence"] == "low" for r in rows)

    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Transfer-In Review — Transfer Report")
    tr.bold = True
    tr.font.size = Pt(20)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Client: {client_name}  |  As of: {as_of}  |  NIFTY 500 list as of: {nifty_as_of or 'unknown'}")
    doc.add_paragraph()

    doc.add_heading("Family AUM gate (> Rs 1 Cr per member AND > Rs 2.5 Cr family total)", level=2)
    gp = doc.add_paragraph()
    gr = gp.add_run(f"Verdict: {family_gate['verdict']}")
    gr.bold = True
    gr.font.size = Pt(12)
    gr.font.color.rgb = (RGBColor(0x2E, 0x6E, 0x62) if family_gate["verdict"] == "CAN PROCESS"
                         else RGBColor(0xA3, 0x4A, 0x28))
    if family_gate.get("note"):
        doc.add_paragraph(family_gate["note"]).runs[0].italic = True
    if family_gate.get("member_totals"):
        _docx_table(
            doc, ["Family Member", "Total Value (INR)", "> Rs 1 Cr?"],
            [[m, f"{v:,.0f}", "PASS" if family_gate["per_member_pass"][m] else "FAIL"]
             for m, v in family_gate["member_totals"].items()],
            right_align_cols={1},
        )
    doc.add_paragraph(
        f"Family total: Rs {family_gate['family_total']:,.0f} "
        f"({'PASS' if family_gate['family_pass'] else 'FAIL'} vs > Rs 2.5 Cr)"
    )
    for note in family_gate.get("data_quality_notes", []):
        p = doc.add_paragraph()
        r = p.add_run(f"Data quality: {note}")
        r.italic = True
        r.font.color.rgb = RGBColor(0xA3, 0x4A, 0x28)
        r.font.size = Pt(9)
    doc.add_paragraph()

    if low_conf:
        p = doc.add_paragraph()
        run = p.add_run(
            "CAUTION: this statement was parsed from a raw PDF via heuristic line/ISIN "
            "scanning (no verified table layout for this statement format). Every row below "
            "marked extraction_confidence=low must be checked against the source pages before "
            "this report is treated as final."
        )
        run.bold = True
        run.font.color.rgb = RGBColor(0xA3, 0x4A, 0x28)
        run.font.size = Pt(10)
        doc.add_paragraph()

    for cat, label in [("EQ", "Stocks"), ("MF", "Mutual Funds"), ("SIF", "SIF"),
                        ("DEBT", "Bonds / Debentures"), ("INVIT", "InvIT / REIT")]:
        cat_rows = [r for r in approved if r["category"] == cat]
        if not cat_rows:
            continue
        doc.add_heading(f"Pre-Approved — {label}", level=2)
        _docx_table(
            doc, ["Name", "ISIN", "Value (INR)", "Note"],
            [[r["name"], r["isin"], f"{r['value_inr']:,.0f}" if r["value_inr"] else "", r["reason"]]
             for r in cat_rows],
            right_align_cols={2},
        )

    if exceptions:
        doc.add_heading("Exceptions — routed for review", level=2)
        _docx_table(
            doc, ["Name", "ISIN", "Category", "Value (INR)", "Routed To", "Reason"],
            [[r["name"], r["isin"], r["category"], f"{r['value_inr']:,.0f}" if r["value_inr"] else "",
              r["routed_to"], r["reason"]] for r in exceptions],
            right_align_cols={3},
        )

    doc.add_heading("Action items", level=2)
    for owner, note in [
        ("Yajash", "Review every Stocks exception above (beyond NIFTY 500) and confirm in writing."),
        ("Girdhar", "Confirm listed status + credit rating for every Debt/InvIT exception above."),
        ("RM", "Confirm purchase history for every Stock and every demat-mode MF/SIF row "
               "(pre-approved AND exceptions) before any transfer proceeds — see "
               "rm_purchase_history_confirmed column in transfer_report.csv."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{owner}: ")
        r.bold = True
        p.add_run(note)

    p = doc.add_paragraph()
    r = p.add_run(f"Source statement: {os.path.basename(source_path)}")
    r.italic = True
    r.font.size = Pt(8.5)

    doc.save(out_path)


def run(statement_path, client_name, out_dir, as_of=None):
    as_of = as_of or dt.date.today().isoformat()
    os.makedirs(out_dir, exist_ok=True)

    if statement_path.lower().endswith(".pdf"):
        raw_rows = extract_from_pdf(statement_path)
    else:
        raw_rows = load_pre_extracted(statement_path)

    by_isin, by_name, nifty_as_of = load_nifty500()
    rows = build_report_rows(raw_rows, by_isin, by_name)
    family_gate = compute_family_gate(rows)

    csv_path = os.path.join(out_dir, "transfer_report.csv")
    docx_path = os.path.join(out_dir, f"Transfer_Report_{_norm(client_name) or 'client'}.docx")
    prior = load_prior_confirmations(csv_path)
    if prior:
        apply_prior_confirmations(rows, prior)
        print(f"carried forward {len(prior)} prior RM purchase-history confirmation(s) from the existing {csv_path}")
    write_csv(rows, csv_path)
    write_docx(rows, docx_path, client_name, statement_path, as_of, nifty_as_of, family_gate)

    summary = {
        "client": client_name, "as_of": as_of, "source": statement_path,
        "nifty500_list_as_of": nifty_as_of,
        "n_rows": len(rows),
        "n_pre_approved": sum(1 for r in rows if r["status"] == "Pre-Approved"),
        "n_exceptions": sum(1 for r in rows if r["status"] != "Pre-Approved"),
        "routed": {name: sum(1 for r in rows if r["routed_to"] == name) for name in ("Yajash", "Girdhar", "RM")},
        "low_confidence_rows": sum(1 for r in rows if r["extraction_confidence"] == "low"),
        "family_gate": family_gate,
    }
    json.dump(summary, open(os.path.join(out_dir, "summary.json"), "w", encoding="utf-8"), indent=2)

    print(f"transfer_in_review: {summary['n_pre_approved']} pre-approved, "
          f"{summary['n_exceptions']} exceptions -> {csv_path}")
    print(f"family AUM gate: {family_gate['verdict']}")
    if summary["low_confidence_rows"]:
        print(f"WARNING: {summary['low_confidence_rows']} rows extracted from PDF at low confidence — "
              f"verify against source pages before treating this report as final.")
    print(f"docx -> {docx_path}")
    return summary


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--statement", required=True, help="CAS/CAMS/Kfintech PDF, or a pre-extracted CSV/XLSX")
    ap.add_argument("--client", required=True)
    ap.add_argument("--out", default="")
    a = ap.parse_args()
    out_dir = a.out or os.path.join(HERE, "out", f"{_norm(a.client)}_{dt.date.today().isoformat()}")
    run(a.statement, a.client, out_dir)
